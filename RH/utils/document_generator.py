import io
import tempfile
import boto3
from botocore.exceptions import BotoCoreError, NoCredentialsError
from datetime import datetime, timedelta
from django.conf import settings
from django.core.files.storage import default_storage, FileSystemStorage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

class DocumentGenerator:
    """
    GENERADOR DE DOCUMENTOS SUPER COMPLETO
    ======================================
    Características:
    1. Detecta automáticamente operadores vs administrativos
    2. Usa carpetas separadas: /administrativos/ y /operadores/
    3. Formatea fechas correctamente: "12 de Diciembre del 2025"
    4. Maneja todas las variantes de variables (con/sin acentos, mayúsculas, etc.)
    5. Diagnóstico completo para troubleshooting
    6. Genera documentos específicos según tipo de empleado
    """
    
    def __init__(self, empleado):
        self.empleado = empleado
        self.base_dir = 'plantillas' # Relativo a MEDIA dentro del storage
        self.context = self._build_context()
        print(f"\n📋 DocumentGenerator iniciado para: {self.empleado.nombre_completo}")
        print(f"   📌 Tipo: {'OPERADOR' if self.es_operador else 'ADMINISTRATIVO'}")
    
    # ========== PROPIEDADES PRINCIPALES ==========
    
    @property
    def es_operador(self):
        """Determina si el empleado es operador"""
        if not self.empleado.puesto or not self.empleado.puesto.nombre:
            return False
        
        puesto = self.empleado.puesto.nombre.lower()
        if 'mensajero' in puesto:
            return False
            
        palabras_operador = ['operador', 'chofer', 'conductor', 'motorista', 'driver', 'operadora']
        return any(palabra in puesto for palabra in palabras_operador)
    
    @property
    def folder_tipo(self):
        """Retorna la carpeta de plantillas según tipo"""
        return 'operadores' if self.es_operador else 'administrativos'
    
    @property
    def folder_path(self):
        """Ruta relativa a la carpeta de plantillas en Storage"""
        return f"{self.base_dir}/{self.folder_tipo}"
    
    # ========== MAPEO DE PLANTILLAS ==========
    
    def _get_template_mapping(self):
        """
        MAPEO COMPLETO DE PLANTILLAS
        -----------------------------
        Cada tipo de empleado tiene sus propias plantillas en carpetas separadas.
        """
        if self.es_operador:
            # ========== PLANTILLAS PARA OPERADORES ==========
            return {
                # Documentos ESPECÍFICOS para operadores (4 documentos)
                'carta_renuncia': 'carta renuncia op.docx',
                'contrato_indeterminado': 'contrato indeterminado op.docx',
                'contrato_determinado': 'contrato determinado OP.docx',
                'formato_desfase': 'formato de desfase op.docx',
            }
        else:
            # ========== PLANTILLAS PARA ADMINISTRATIVOS ==========
            return {
                # Documentos para administrativos (NO tienen 'op' en el nombre)
                'contrato_indeterminado': 'CONTRATO INDETERMINADO_ALEXIS EDUARDO AGUNDIZ.docx',
                'contrato_determinado': 'CONTRATO.docx',
                
                # Documentos comunes
                'aviso_privacidad': 'AVISO DE PRIVACIDAD.docx',
                'convenio_confidencialidad': 'CONVENIO DE CONFIDENCIALIDAD EMPLEADOS.docx',
                'autorizacion_correo': 'AUTORIZACION CORREO PARA RECIBO SAT.docx',
                'carta_adhesion': 'CARTA ADHESION PLAN PREV. SOC.docx',
                'descripcion_puesto': 'DESCRIPCION DE PUESTO.docx',
                
                # Estos NO existen para administrativos (retornar None)
                'carta_renuncia': None,
                'formato_desfase': None,
            }
    
    # ========== CONSTRUCCIÓN DEL CONTEXTO ==========
    
    def _build_context(self):
        """
        CONSTRUYE EL CONTEXTO CON TODAS LAS VARIABLES POSIBLES
        ------------------------------------------------------
        Incluye múltiples formatos para cada dato.
        """
        hoy = datetime.now()
        fecha_contratacion = self.empleado.fecha_contratacion or hoy
        
        # ========== FECHAS IMPORTANTES ==========
        fecha_fin_contrato = fecha_contratacion + timedelta(days=90)
        
        # ========== SALARIO ==========
        salario_actual = self.empleado.salario_actual
        if salario_actual and salario_actual.sueldo_diario:
            try:
                sueldo_diario = float(salario_actual.sueldo_diario)
            except (ValueError, TypeError):
                sueldo_diario = 0.0
        else:
            sueldo_diario = 0.0
        
        sueldo_mensual = sueldo_diario * 30
        sueldo_quincenal = sueldo_diario * 15
        
        # ========== EDAD ==========
        edad = self._calcular_edad()
        
        # ========== SEXO/GÉNERO ==========
        sexo = self._determinar_sexo()
        
        # ========== DIRECCIÓN COMPLETA ==========
        direccion_completa = self._construir_direccion_completa()
        
        # ========== CONSTRUIR CONTEXTO COMPLETO ==========
        context = {
            # ========== INFORMACIÓN BÁSICA (múltiples formatos) ==========
            'nombre_completo': f"{self.empleado.nombre} {self.empleado.apellido}".strip(),
            'NOMBRE_COMPLETO': f"{self.empleado.nombre} {self.empleado.apellido}".strip().upper(),
            'nombre': self.empleado.nombre or '',
            'apellido': self.empleado.apellido or '',
            'apellido_paterno': self.empleado.apellido or '',
            'apellido_materno': '',  # Si tienes este campo separado, ajústalo
            
            # ========== EDAD (múltiples formatos) ==========
            'edad': edad,
            'EDAD': edad,
            'edad_actual': edad,
            'edad_completa': f"{edad} años" if edad else '',
            'edad_años': edad,
            
            # ========== SEXO/GÉNERO ==========
            'sexo': sexo,
            'genero': sexo,
            'SEXO': sexo.upper(),
            'GENERO': sexo.upper(),
            
            # ========== DATOS PERSONALES ==========
            'curp': self.empleado.curp or '',
            'CURP': (self.empleado.curp or '').upper(),
            'rfc': self.empleado.rfc or '',
            'RFC': (self.empleado.rfc or '').upper(),
            'nss': self.empleado.nss or '',
            'NSS': self.empleado.nss or '',
            'numero_seguro_social': self.empleado.nss or '',
            
            # ========== NACIONALIDAD ==========
            'nacionalidad': self.empleado.nacionalidad or 'Mexicana',
            'nacionalidad_completa': f"{self.empleado.nacionalidad or 'Mexicana'}",
            'NACIONALIDAD': (self.empleado.nacionalidad or 'Mexicana').upper(),
            
            # ========== ESTADO CIVIL ==========
            'estado_civil': self.empleado.get_estado_civil_display() if self.empleado.estado_civil else '',
            'estado_civil_minusculas': (self.empleado.get_estado_civil_display().lower() 
                                       if self.empleado.estado_civil else ''),
            'ESTADO_CIVIL': (self.empleado.get_estado_civil_display().upper() 
                            if self.empleado.estado_civil else ''),
            
            # ========== DIRECCIÓN (MÚLTIPLES FORMATOS - ESTO ES CLAVE) ==========
            'direccion_completa': direccion_completa,
            'dirección_completa': direccion_completa,  # Con acento
            'DIRECCION_COMPLETA': direccion_completa.upper(),
            'DIRECCIÓN_COMPLETA': direccion_completa.upper(),  # Con acento y mayúsculas
            'Direccion_Completa': direccion_completa.title(),
            'Dirección_Completa': direccion_completa.title(),  # Con acento
            
            # Componentes individuales de dirección
            'direccion': self.empleado.direccion or '',
            'dirección': self.empleado.direccion or '',  # Con acento
            'calle': self.empleado.direccion or '',
            'domicilio': self.empleado.direccion or '',
            
            'colonia': self.empleado.colonia or '',
            'COLONIA': (self.empleado.colonia or '').upper(),
            
            'codigo_postal': self.empleado.codigo_postal or '',
            'cp': self.empleado.codigo_postal or '',
            'CODIGO_POSTAL': self.empleado.codigo_postal or '',
            'CP': self.empleado.codigo_postal or '',
            
            'ciudad': self.empleado.ciudad or '',
            'CIUDAD': (self.empleado.ciudad or '').upper(),
            
            'estado': self.empleado.estado or '',
            'ESTADO': (self.empleado.estado or '').upper(),
            
            'pais': self.empleado.pais or 'México',
            'PAIS': (self.empleado.pais or 'México').upper(),
            
            # ========== CONTACTO ==========
            'email': self.empleado.email or '',
            'correo': self.empleado.email or '',
            'correo_electronico': self.empleado.email or '',
            'EMAIL': (self.empleado.email or '').upper(),
            
            'telefono': self.empleado.telefono_personal or '',
            'telefono_personal': self.empleado.telefono_personal or '',
            'celular': self.empleado.telefono_personal or '',
            'TELEFONO': self.empleado.telefono_personal or '',
            
            # ========== DATOS LABORALES ==========
            'puesto': self.empleado.puesto.nombre if self.empleado.puesto else '',
            'PUESTO': (self.empleado.puesto.nombre if self.empleado.puesto else '').upper(),
            'cargo': self.empleado.puesto.nombre if self.empleado.puesto else '',
            'CARGO': (self.empleado.puesto.nombre if self.empleado.puesto else '').upper(),
            
            'departamento': self.empleado.departamento.nombre if self.empleado.departamento else '',
            'DEPARTAMENTO': (self.empleado.departamento.nombre if self.empleado.departamento else '').upper(),
            'area': self.empleado.departamento.nombre if self.empleado.departamento else '',
            'AREA': (self.empleado.departamento.nombre if self.empleado.departamento else '').upper(),
            
            # ========== TIPO DE EMPLEADO ==========
            'es_operador': 'Sí' if self.es_operador else 'No',
            'ES_OPERADOR': 'SÍ' if self.es_operador else 'NO',
            'tipo_empleado': 'Operador' if self.es_operador else 'Personal Administrativo',
            'TIPO_EMPLEADO': 'OPERADOR' if self.es_operador else 'PERSONAL ADMINISTRATIVO',
            
            # ========== FECHAS - FORMATO "12 de Diciembre del 2025" ==========
            # Fecha actual (hoy)
            'fecha_hoy': self._formatear_fecha_exacta(hoy),
            'fecha_actual': self._formatear_fecha_exacta(hoy),
            'fecha': self._formatear_fecha_exacta(hoy),
            'fecha_exacta': self._formatear_fecha_exacta(hoy),
            'fecha_generacion': self._formatear_fecha_exacta(hoy),
            'fecha_impresion': self._formatear_fecha_exacta(hoy),
            'FECHA_HOY': self._formatear_fecha_exacta(hoy).upper(),
            
            # Fecha de contratación
            'fecha_contratacion': self._formatear_fecha_exacta(fecha_contratacion),
            'fecha_ingreso': self._formatear_fecha_exacta(fecha_contratacion),
            'fecha_inicio': self._formatear_fecha_exacta(fecha_contratacion),
            'fecha_inicio_contrato': self._formatear_fecha_exacta(fecha_contratacion),
            'FECHA_CONTRATACION': self._formatear_fecha_exacta(fecha_contratacion).upper(),
            
            # Fecha de fin de contrato
            'fecha_fin_contrato': self._formatear_fecha_exacta(fecha_fin_contrato),
            'fecha_termino': self._formatear_fecha_exacta(fecha_fin_contrato),
            'fecha_finalizacion': self._formatear_fecha_exacta(fecha_fin_contrato),
            'FECHA_FIN_CONTRATO': self._formatear_fecha_exacta(fecha_fin_contrato).upper(),
            
            # ========== FECHAS - FORMATO ALTERNATIVOS ==========
            'fecha_hoy_simple': self._formatear_fecha_simple(hoy),
            'fecha_contratacion_simple': self._formatear_fecha_simple(fecha_contratacion),
            'fecha_fin_simple': self._formatear_fecha_simple(fecha_fin_contrato),
            
            'fecha_hoy_numerica': hoy.strftime('%d/%m/%Y'),
            'fecha_contratacion_numerica': fecha_contratacion.strftime('%d/%m/%Y'),
            'fecha_fin_numerica': fecha_fin_contrato.strftime('%d/%m/%Y'),
            
            # ========== SALARIO - FORMATO NUMÉRICO ==========
            'salario_diario': f"${sueldo_diario:,.2f}" if sueldo_diario else "$0.00",
            'salario_diario_numero': f"{sueldo_diario:,.2f}" if sueldo_diario else "0.00",
            'sueldo_diario': f"${sueldo_diario:,.2f}" if sueldo_diario else "$0.00",
            'SUELDO_DIARIO': f"${sueldo_diario:,.2f}".upper() if sueldo_diario else "$0.00",
            
            'salario_mensual': f"${sueldo_mensual:,.2f}" if sueldo_mensual else "$0.00",
            'sueldo_mensual': f"${sueldo_mensual:,.2f}" if sueldo_mensual else "$0.00",
            'SUELDO_MENSUAL': f"${sueldo_mensual:,.2f}".upper() if sueldo_mensual else "$0.00",
            
            'salario_quincenal': f"${sueldo_quincenal:,.2f}" if sueldo_quincenal else "$0.00",
            'sueldo_quincenal': f"${sueldo_quincenal:,.2f}" if sueldo_quincenal else "$0.00",
            
            # ========== SALARIO - EN LETRAS ==========
            'salario_diario_letras': self._numero_a_letras(sueldo_diario),
            'salario_diario_en_letras': self._numero_a_letras(sueldo_diario),
            'sueldo_diario_letras': self._numero_a_letras(sueldo_diario),
            'sueldo_diario_en_letras': self._numero_a_letras(sueldo_diario),
            
            'salario_mensual_letras': self._numero_a_letras(sueldo_mensual),
            'salario_mensual_en_letras': self._numero_a_letras(sueldo_mensual),
            'sueldo_mensual_letras': self._numero_a_letras(sueldo_mensual),
            'sueldo_mensual_en_letras': self._numero_a_letras(sueldo_mensual),
            
            'salario_quincenal_letras': self._numero_a_letras(sueldo_quincenal),
            'sueldo_quincenal_letras': self._numero_a_letras(sueldo_quincenal),
            
            # ========== INFORMACIÓN DE LA EMPRESA ==========
            'nombre_empresa': 'TRANSPORTES MIGMAR S.A. DE C.V.',
            'empresa': 'TRANSPORTES MIGMAR S.A. DE C.V.',
            'razon_social': 'TRANSPORTES MIGMAR S.A. DE C.V.',
            'NOMBRE_EMPRESA': 'TRANSPORTES MIGMAR S.A. DE C.V.',
            'EMPRESA': 'TRANSPORTES MIGMAR S.A. DE C.V.',
            
            'direccion_empresa': 'AVENIDA 1 (UNO), NÚMERO 130, COLONIA CENTRAL DE CARGA',
            'domicilio_empresa': 'AVENIDA 1 (UNO), NÚMERO 130, COLONIA CENTRAL DE CARGA',
            'DIRECCION_EMPRESA': 'AVENIDA 1 (UNO), NÚMERO 130, COLONIA CENTRAL DE CARGA',
            
            'ciudad_empresa': 'San Nicolás de los Garza, N.L.',
            'ubicacion_empresa': 'San Nicolás de los Garza, N.L.',
            'CIUDAD_EMPRESA': 'SAN NICOLÁS DE LOS GARZA, N.L.',
            
            # ========== INFORMACIÓN OPERATIVA (para operadores) ==========
            'tipo_licencia': self._obtener_tipo_licencia(),
            'numero_licencia': self._obtener_numero_licencia(),
            'fecha_vencimiento_licencia': self._obtener_fecha_vencimiento_licencia(),
            
            'tipo_viaje': ', '.join([tv.nombre for tv in self.empleado.tipo_viaje.all()]),
            'tipo_carga': ', '.join([tc.nombre for tc in self.empleado.tipo_carga.all()]),
            'division_operativa': ', '.join([do.nombre for do in self.empleado.division_operativa.all()]),
            
            # ========== REFERENCIAS ==========
            'nombre_referencia_1': self.empleado.nombre_referencia_1 or '',
            'telefono_referencia_1': self.empleado.telefono_referencia_1 or '',
            'relacion_referencia_1': self.empleado.relacion_referencia_1 or '',
            
            'nombre_referencia_2': self.empleado.nombre_referencia_2 or '',
            'telefono_referencia_2': self.empleado.telefono_referencia_2 or '',
            'relacion_referencia_2': self.empleado.relacion_referencia_2 or '',
            
            # ========== NÚMERO DE EMPLEADO ==========
            'numero_empleado': self.empleado.numero_empleado or '',
            'num_empleado': self.empleado.numero_empleado or '',
            'NÚMERO_EMPLEADO': (self.empleado.numero_empleado or '').upper(),
        }
        
        return context
    
    # ========== MÉTODOS AUXILIARES DE FORMATO ==========
    
    def _formatear_fecha_exacta(self, fecha):
        """Formatea fecha a '12 de Diciembre del 2025'"""
        if not fecha:
            return ''
        
        meses = {
            1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
            5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
            9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
        }
        
        return f"{fecha.day} de {meses.get(fecha.month, '')} del {fecha.year}"
    
    def _formatear_fecha_simple(self, fecha):
        """Formatea fecha a '12/diciembre/2025'"""
        if not fecha:
            return ''
        
        meses = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        
        return f"{fecha.day}/{meses.get(fecha.month, '')}/{fecha.year}"
    
    def _calcular_edad(self):
        """Calcula la edad del empleado"""
        if not self.empleado.fecha_nacimiento:
            return ''
        
        hoy = datetime.now().date()
        nacimiento = self.empleado.fecha_nacimiento
        
        edad = hoy.year - nacimiento.year
        if (hoy.month, hoy.day) < (nacimiento.month, nacimiento.day):
            edad -= 1
        
        return str(edad)
    
    def _determinar_sexo(self):
        """Determina el sexo basado en CURP"""
        if not self.empleado.curp or len(self.empleado.curp) < 11:
            return 'masculino'  # valor por defecto
        
        sexo_char = self.empleado.curp[10].upper()
        if sexo_char == 'H':
            return 'masculino'
        elif sexo_char == 'M':
            return 'femenino'
        else:
            return 'masculino'
    
    def _construir_direccion_completa(self):
        """Construye la dirección completa en formato legible"""
        partes = []
        
        if self.empleado.direccion:
            partes.append(self.empleado.direccion)
        
        if self.empleado.colonia:
            partes.append(f"Col. {self.empleado.colonia}")
        
        if self.empleado.codigo_postal:
            partes.append(f"C.P. {self.empleado.codigo_postal}")
        
        if self.empleado.ciudad:
            partes.append(self.empleado.ciudad)
        
        if self.empleado.estado:
            partes.append(self.empleado.estado)
        
        if self.empleado.pais and self.empleado.pais != 'México':
            partes.append(self.empleado.pais)
        
        return ', '.join(partes) if partes else 'Dirección no proporcionada'
    
    def _numero_a_letras(self, numero):
        """Convierte número a letras en español (versión mejorada)"""
        try:
            # Limpiar y convertir
            if isinstance(numero, str):
                numero = numero.replace('$', '').replace(',', '').strip()
            
            numero_float = float(numero)
            numero_float = round(numero_float, 2)
            
            # Separar parte entera y decimal
            entero = int(numero_float)
            decimal = int(round((numero_float - entero) * 100))
            
            # Listas de números en letras
            unidades = ['', 'UN', 'DOS', 'TRES', 'CUATRO', 'CINCO', 
                       'SEIS', 'SIETE', 'OCHO', 'NUEVE']
            
            decenas = ['', 'DIEZ', 'VEINTE', 'TREINTA', 'CUARENTA', 'CINCUENTA',
                      'SESENTA', 'SETENTA', 'OCHENTA', 'NOVENTA']
            
            # Números especiales (11-29)
            especiales = {
                11: 'ONCE', 12: 'DOCE', 13: 'TRECE', 14: 'CATORCE', 15: 'QUINCE',
                16: 'DIECISÉIS', 17: 'DIECISIETE', 18: 'DIECIOCHO', 19: 'DIECINUEVE',
                20: 'VEINTE', 21: 'VEINTIÚN', 22: 'VEINTIDÓS', 23: 'VEINTITRÉS',
                24: 'VEINTICUATRO', 25: 'VEINTICINCO', 26: 'VEINTISÉIS', 
                27: 'VEINTISIETE', 28: 'VEINTIOCHO', 29: 'VEINTINUEVE'
            }
            
            # Función para convertir números pequeños
            def convertir_menor_100(num):
                if num == 0:
                    return ''
                elif num in especiales:
                    return especiales[num]
                elif num < 10:
                    return unidades[num]
                elif num < 100:
                    decena = num // 10
                    unidad = num % 10
                    
                    if unidad == 0:
                        return decenas[decena]
                    elif decena == 2:  # Veinti-algo
                        return f"VEINTI{unidades[unidad].lower()}"
                    else:
                        return f"{decenas[decena]} Y {unidades[unidad].lower()}"
                else:
                    return str(num)
            
            # Convertir parte entera
            if entero == 0:
                texto = 'CERO'
            elif entero < 100:
                texto = convertir_menor_100(entero)
            elif entero < 1000:
                centenas = entero // 100
                resto = entero % 100
                
                if centenas == 1:
                    if resto == 0:
                        texto = 'CIEN'
                    else:
                        texto = f"CIENTO {convertir_menor_100(resto).lower()}"
                elif centenas == 5:
                    texto = f"QUINIENTOS {convertir_menor_100(resto).lower()}"
                elif centenas == 7:
                    texto = f"SETECIENTOS {convertir_menor_100(resto).lower()}"
                elif centenas == 9:
                    texto = f"NOVECIENTOS {convertir_menor_100(resto).lower()}"
                else:
                    texto = f"{unidades[centenas]}CIENTOS {convertir_menor_100(resto).lower()}"
            else:
                # Para números grandes, simplificar
                texto = str(entero)
            
            texto = texto.strip().upper()
            
            # Agregar decimales
            if decimal > 0:
                return f"{texto} PESOS {decimal:02d}/100 M.N."
            else:
                return f"{texto} PESOS 00/100 M.N."
                
        except (ValueError, TypeError, AttributeError):
            return "CERO PESOS 00/100 M.N."
    
    # ========== MÉTODOS PARA INFORMACIÓN OPERATIVA ==========
    
    def _obtener_tipo_licencia(self):
        """Obtiene tipo de licencia para operadores"""
        if not self.es_operador:
            return ''
        
        for doc in self.empleado.documentos_operador.all():
            if 'licencia' in doc.tipo_documento.nombre.lower():
                return doc.tipo_documento.nombre
        return 'Licencia Federal Tipo C'
    
    def _obtener_numero_licencia(self):
        """Obtiene número de licencia para operadores"""
        if not self.es_operador:
            return ''
        
        for doc in self.empleado.documentos_operador.all():
            if 'licencia' in doc.tipo_documento.nombre.lower() and doc.numero_documento:
                return doc.numero_documento
        return ''
    
    def _obtener_fecha_vencimiento_licencia(self):
        """Obtiene fecha de vencimiento de licencia"""
        if not self.es_operador:
            return ''
        
        for doc in self.empleado.documentos_operador.all():
            if 'licencia' in doc.tipo_documento.nombre.lower() and doc.fecha_vencimiento:
                return self._formatear_fecha_exacta(doc.fecha_vencimiento)
        return ''
    
    # ========== MÉTODOS DE PROCESAMIENTO DE DOCUMENTOS ==========
    
    def _reemplazar_variables_inteligente(self, texto):
        """
        Reemplaza variables de manera INTELIGENTE
        Busca todas las variantes posibles (con/sin acentos, mayúsculas, etc.)
        """
        if not texto or ('{' not in texto and '[' not in texto):
            return texto
        
        # Copiar el texto para trabajar
        resultado = texto
        
        # PRIMERA PASADA: reemplazo directo
        for key, value in self.context.items():
            # Reemplazar {variable}
            if f"{{{key}}}" in resultado:
                resultado = resultado.replace(f"{{{key}}}", str(value))
            
            # Reemplazar {{variable}}
            if f"{{{{{key}}}}}" in resultado:
                resultado = resultado.replace(f"{{{{{key}}}}}", str(value))
            
            # Reemplazar [variable]
            if f"[{key}]" in resultado:
                resultado = resultado.replace(f"[{key}]", str(value))
        
        # SEGUNDA PASADA: buscar variables con acentos/minúsculas
        if '{' in resultado:
            import re
            variables_sin_reemplazar = re.findall(r'\{([^}]+)\}', resultado)
            
            for var_original in variables_sin_reemplazar:
                # Crear múltiples variantes para buscar
                variantes = [
                    var_original,  # Original
                    var_original.lower(),  # Minúsculas
                    var_original.upper(),  # Mayúsculas
                    var_original.title(),  # Tipo título
                    # Quitar acentos
                    var_original.replace('á', 'a').replace('é', 'e').replace('í', 'i')
                             .replace('ó', 'o').replace('ú', 'u'),
                    var_original.replace('Á', 'A').replace('É', 'E').replace('Í', 'I')
                             .replace('Ó', 'O').replace('Ú', 'U'),
                    # Reemplazar espacios por guiones bajos y viceversa
                    var_original.replace(' ', '_'),
                    var_original.replace('_', ' '),
                ]
                
                # Buscar en el contexto
                for variante in variantes:
                    if variante in self.context:
                        resultado = resultado.replace(
                            f"{{{var_original}}}", 
                            str(self.context[variante])
                        )
                        break
        
        return resultado
    
    def _procesar_documento(self, doc):
        """Procesa todo el documento reemplazando variables"""
        # Procesar párrafos
        for paragraph in doc.paragraphs:
            self._procesar_parrafo(paragraph)
        
        # Procesar tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._procesar_parrafo(paragraph)
        
        # Procesar encabezados y pies de página
        for section in doc.sections:
            # Encabezado
            for paragraph in section.header.paragraphs:
                self._procesar_parrafo(paragraph)
            
            # Pie de página
            for paragraph in section.footer.paragraphs:
                self._procesar_parrafo(paragraph)
    
    def _procesar_parrafo(self, paragraph):
        """Procesa un párrafo manteniendo formato"""
        texto_original = paragraph.text
        
        if not texto_original or ('{' not in texto_original and '[' not in texto_original):
            return
        
        # Reemplazar variables
        nuevo_texto = self._reemplazar_variables_inteligente(texto_original)
        
        if nuevo_texto != texto_original:
            self._actualizar_texto_con_formato(paragraph, nuevo_texto)
    
    def _actualizar_texto_con_formato(self, paragraph, nuevo_texto):
        """Actualiza el texto manteniendo el formato original"""
        if not paragraph.runs:
            paragraph.text = nuevo_texto
            return
        
        # Guardar formato del primer run (asumimos formato consistente)
        primer_run = paragraph.runs[0]
        formato = {
            'name': primer_run.font.name if hasattr(primer_run.font, 'name') else None,
            'size': primer_run.font.size,
            'bold': primer_run.font.bold,
            'italic': primer_run.font.italic,
            'underline': primer_run.font.underline,
            'color': primer_run.font.color.rgb if hasattr(primer_run.font.color, 'rgb') else None,
        }
        
        # Limpiar todos los runs
        for run in paragraph.runs:
            run.text = ""
        
        # Poner el nuevo texto en el primer run
        paragraph.runs[0].text = nuevo_texto
        
        # Restaurar formato
        if formato['name']:
            paragraph.runs[0].font.name = formato['name']
        if formato['size']:
            paragraph.runs[0].font.size = formato['size']
        
        paragraph.runs[0].font.bold = formato['bold']
        paragraph.runs[0].font.italic = formato['italic']
        paragraph.runs[0].font.underline = formato['underline']
        
        if formato['color']:
            paragraph.runs[0].font.color.rgb = formato['color']
    
    # ========== MÉTODOS PRINCIPALES DE GENERACIÓN ==========
    
    def get_template_path(self, doc_type):
        """
        Obtiene la ruta RELATIVA de la plantilla
        Retorna None si no existe para este tipo de empleado
        """
        mapping = self._get_template_mapping()
        
        if doc_type not in mapping:
            print(f"❌ Tipo de documento no válido: {doc_type}")
            return None
        
        template_name = mapping[doc_type]
        
        # Si es None, significa que no existe para este tipo
        if template_name is None:
            print(f"⚠ Documento no disponible: {doc_type} para {'operadores' if self.es_operador else 'administrativos'}")
            return None
        
        # Construir ruta relativa para storage
        return f"{self.base_dir}/{self.folder_tipo}/{template_name}"
        
    def _descargar_plantilla_s3(self, template_path):
        """Descarga la plantilla desde S3 a un buffer en memoria usando boto3"""
        try:
            # Validar si tenemos las credenciales de AWS antes de intentar conectar
            aws_bucket = getattr(settings, 'AWS_STORAGE_BUCKET_NAME', None)
            aws_key = getattr(settings, 'AWS_ACCESS_KEY_ID', None)
            
            if not aws_bucket or not aws_key:
                print(f"   ☁ S3 no configurado (falta AWS_STORAGE_BUCKET_NAME o AWS_ACCESS_KEY_ID). Omitiendo S3...")
                return None
                
            print(f"   ☁ Intentando descargar de S3: {template_path}...")
            s3_client = boto3.client(
                's3',
                aws_access_key_id=aws_key,
                aws_secret_access_key=getattr(settings, 'AWS_SECRET_ACCESS_KEY', None),
                region_name=getattr(settings, 'AWS_S3_REGION_NAME', 'us-east-1')
            )
            
            # Obtener location (ej: 'media'), pero si no está definido, usar string vacío
            aws_location = getattr(settings, 'AWS_MEDIA_LOCATION', '')
            if aws_location:
                full_s3_key = f"{aws_location}/{template_path}"
            else:
                full_s3_key = template_path
                
            archivo_obj = io.BytesIO()
            
            s3_client.download_fileobj(
                aws_bucket,
                full_s3_key,
                archivo_obj
            )
            archivo_obj.seek(0)
            print(f"   ✅ Éxito descargando de S3: {template_path}")
            return archivo_obj
        except (BotoCoreError, NoCredentialsError, Exception) as e:
            print(f"   ❌ Falló descarga de S3: {e}")
            return None

    def generate_document(self, doc_type, output_name):
        """Genera un documento específico y retorna un BytesIO con el contenido"""
        template_path = self.get_template_path(doc_type)
        
        if not template_path:
            return self._crear_documento_error(doc_type, "Plantilla no encontrada")
        
        try:
            print(f"   📄 Generando: {doc_type}...")
            f = None
            
            # 1. Intentar descargar de S3 mediante boto3 directamente
            buffer_s3 = self._descargar_plantilla_s3(template_path)
            
            if buffer_s3:
                f = buffer_s3
            else:
                # 2. Si falla S3, usar FileSystemStorage local estricto
                local_fs = FileSystemStorage(location=settings.MEDIA_ROOT)
                if local_fs.exists(template_path):
                    print(f"   💡 Usando plantilla local de media/plantillas/: {template_path}")
                    f = local_fs.open(template_path, 'rb')
                else:
                    print(f"❌ Plantilla no encontrada ni en S3 ni en local: {template_path}")
                    return self._crear_documento_error(doc_type, f"Plantilla no existe en S3 ni en local: {template_path}")
            
            # Procesar el documento con docx
            doc = Document(f)
            self._procesar_documento(doc)
            
            if hasattr(f, 'close') and not isinstance(f, io.BytesIO):
                f.close()
            
            # Guardar en buffer de memoria
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            print(f"   ✅ Generado exitosamente en memoria")
            return output
            
        except Exception as e:
            print(f"   ❌ Error generando {doc_type}: {str(e)}")
            return self._crear_documento_error(doc_type, str(e))
    
    def _crear_documento_error(self, doc_type, error_msg):
        """Crea documento de error en memoria cuando falla la generación"""
        try:
            doc = Document()
            
            # Título en rojo
            title = doc.add_heading(f'ERROR AL GENERAR {doc_type.upper()}', 0)
            for run in title.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
            
            # Información
            doc.add_heading('Información:', level=1)
            doc.add_paragraph(f"Empleado: {self.empleado.nombre_completo}")
            doc.add_paragraph(f"Tipo: {'Operador' if self.es_operador else 'Administrativo'}")
            doc.add_paragraph(f"Puesto: {self.empleado.puesto.nombre if self.empleado.puesto else 'No asignado'}")
            doc.add_paragraph(f"Error: {error_msg}")
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            
            # Solución sugerida
            doc.add_heading('Solución sugerida:', level=1)
            doc.add_paragraph(f"Verifica que exista la plantilla en el bucket/storage:")
            doc.add_paragraph(f"  media/plantillas/{self.folder_tipo}/")
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output
            
        except Exception as e:
            print(f"❌ Error crítico creando documento de error: {e}")
            return None
    
    # ========== MÉTODOS PÚBLICOS PRINCIPALES ==========
    
    def generate_all_documents(self):
        """
        Genera TODOS los documentos apropiados para el tipo de empleado
        """
        print(f"\n🎯 GENERANDO DOCUMENTOS PARA: {self.empleado.nombre_completo}")
        print(f"   📁 Carpeta: {self.folder_tipo}/")
        
        documentos = {}
        mapping = self._get_template_mapping()
        
        # Mapeo de nombres de salida
        output_names = {
            'contrato_indeterminado': 'CONTRATO_INDETERMINADO',
            'contrato_determinado': 'CONTRATO_DETERMINADO',
            'carta_renuncia': 'CARTA_RENUNCIA',
            'formato_desfase': 'FORMATO_DESFASE',
            'aviso_privacidad': 'AVISO_PRIVACIDAD',
            'convenio_confidencialidad': 'CONVENIO_CONFIDENCIALIDAD',
            'autorizacion_correo': 'AUTORIZACION_CORREO',
            'carta_adhesion': 'CARTA_ADHESION',
            'descripcion_puesto': 'DESCRIPCION_PUESTO',
        }
        
        for doc_type in mapping.keys():
            # Solo generar si existe plantilla para este tipo
            if mapping[doc_type] is not None:
                output_name = output_names.get(doc_type, doc_type.upper())
                doc_path = self.generate_document(doc_type, output_name)
                
                if doc_path:
                    documentos[doc_type] = doc_path
        
        total = len(documentos)
        print(f"\n📊 RESUMEN: {total} documentos generados exitosamente")
        
        return documentos
    
    def create_zip_file(self):
        """Crea archivo ZIP con todos los documentos generados en memoria"""
        import zipfile
        
        documentos = self.generate_all_documents()
        
        if not documentos:
            print("❌ No se generaron documentos para crear ZIP")
            return None
        
        # Mapeo de nombres amigables para los archivos
        nombres_amigables = {
            'contrato_indeterminado': 'CONTRATO INDETERMINADO',
            'contrato_determinado': 'CONTRATO DETERMINADO',
            'carta_renuncia': 'CARTA DE RENUNCIA',
            'formato_desfase': 'FORMATO DE DESFASE',
            'aviso_privacidad': 'AVISO DE PRIVACIDAD',
            'convenio_confidencialidad': 'CONVENIO DE CONFIDENCIALIDAD',
            'autorizacion_correo': 'AUTORIZACION DE CORREO',
            'carta_adhesion': 'CARTA DE ADHESION',
            'descripcion_puesto': 'DESCRIPCION DE PUESTO',
        }
        
        # Crear buffer para el ZIP
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            print(f"\n📦 Preparando archivos para ZIP...")
            for doc_type, doc_content in documentos.items():
                if doc_content:
                    nombre = nombres_amigables.get(doc_type, doc_type.upper())
                    filename = f"{nombre} - {self.empleado.nombre_completo}.docx"
                    
                    # Limpiar nombre para ZIP
                    filename = filename.replace('/', '_').replace('\\', '_')
                    
                    # Agregar al ZIP
                    zip_file.writestr(filename, doc_content.getvalue())
                    print(f"   📎 Agregado: {filename}")
        
        zip_buffer.seek(0)
        return zip_buffer
    
    # ========== MÉTODOS DE DIAGNÓSTICO ==========
    
    def diagnosticar_completo(self):
        """
        Diagnóstico COMPLETO del sistema para Storage/S3
        """
        print(f"\n{'='*80}")
        print(f"DIAGNÓSTICO COMPLETO - {self.empleado.nombre_completo}")
        print(f"{'='*80}")
        
        # 1. Información del empleado
        print(f"\n📋 INFORMACIÓN DEL EMPLEADO:")
        print(f"   • Nombre: {self.empleado.nombre_completo}")
        print(f"   • Puesto: {self.empleado.puesto.nombre if self.empleado.puesto else 'N/A'}")
        print(f"   • Tipo detectado: {'OPERADOR' if self.es_operador else 'ADMINISTRATIVO'}")
        
        # 2. Estructura de carpetas en Storage
        print(f"\n📁 ESTRUCTURA DE CARPETAS (STORAGE):")
        print(f"   • Base: {self.base_dir}")
        print(f"   • Carpeta usada: {self.folder_tipo}/")
        print(f"   • Path Storage: {self.folder_path}")
        
        # 3. Verificando plantillas en Storage
        print(f"\n📄 VERIFICANDO PLANTILLAS DISPONIBLES EN STORAGE:")
        try:
            if default_storage.exists(self.folder_path):
                print(f"   ✅ Carpeta '{self.folder_path}' EXISTE.")
                dirs, files = default_storage.listdir(self.folder_path)
                if files:
                    print(f"   Archivos encontrados ({len(files)}):")
                    for f in sorted(files):
                        print(f"     - {f}")
                else:
                    print(f"   ⚠ La carpeta está vacía.")
            else:
                print(f"   ❌ ERROR: La carpeta '{self.folder_path}' NO EXISTE.")
        except Exception as e:
            print(f"   ❌ Error al verificar storage: {e}")
        
        # 4. Plantillas mapeadas
        print(f"\n🗺  MAPEO DE PLANTILLAS para {self.folder_tipo}:")
        mapping = self._get_template_mapping()
        
        for doc_type, template_name in mapping.items():
            if template_name:
                template_path = f"{self.folder_path}/{template_name}"
                existe = default_storage.exists(template_path)
                status = "✅" if existe else "❌"
                print(f"   {status} {doc_type:25} -> {template_name:40} {'(EXISTE)' if existe else '(NO EXISTE)'}")
            else:
                print(f"   - {doc_type:25} -> {'NO DISPONIBLE PARA ESTE TIPO':40}")
        
        # 5. Variables (ejemplos)
        print(f"\n🔤 EJEMPLO DE VARIABLES:")
        print(f"   • nombre_completo: {self.context.get('nombre_completo', 'N/A')}")
        print(f"   • fecha_hoy: {self.context.get('fecha_hoy', 'N/A')}")
        
        print(f"\n{'='*80}")
        return True
    
    def mostrar_variables_plantilla(self, template_path):
        """
        Analiza una plantilla en storage y muestra las variables que contiene
        """
        if not default_storage.exists(template_path):
            print(f"❌ Plantilla no encontrada en storage: {template_path}")
            return
        
        try:
            with default_storage.open(template_path, 'rb') as f:
                doc = Document(f)
            
            todas_variables = set()
            import re
            
            # Extraer todas las variables de la plantilla
            for paragraph in doc.paragraphs:
                variables = re.findall(r'\{([^}]+)\}', paragraph.text)
                todas_variables.update(variables)
            
            print(f"\n🔍 ANALIZANDO PLANTILLA: {template_path}")
            print(f"   Encontradas {len(todas_variables)} variables:")
            
            for var in sorted(todas_variables):
                # Verificar si existe en contexto
                existe = False
                valor = 'N/A'
                
                if var in self.context:
                    existe = True
                    valor = str(self.context[var])
                
                status = "✅" if existe else "❌"
                print(f"   {status} {{{var:30}}} -> {valor[:40]}...")
            
        except Exception as e:
            print(f"❌ Error analizando plantilla: {e}")
    
    # ========== MÉTODOS DE CONVENIENCIA ==========
    
    def get_documentos_disponibles(self):
        """Lista de documentos disponibles para este tipo de empleado"""
        mapping = self._get_template_mapping()
        disponibles = []
        
        for doc_type, template_name in mapping.items():
            if template_name:  # Solo si no es None
                disponibles.append(doc_type)
        
        return disponibles
    
    def get_nombre_carpeta(self):
        """Nombre de la carpeta de plantillas en uso"""
        return self.folder_tipo
    
    def verificar_plantilla(self, doc_type):
        """Verifica si una plantilla específica existe (S3 compatible)"""
        template_path = self.get_template_path(doc_type)
        
        if template_path and default_storage.exists(template_path):
            print(f"✅ Plantilla {doc_type} existe: {template_path}")
            return True
        else:
            print(f"❌ Plantilla {doc_type} NO existe o no está disponible")
            return False
    
    def test_generacion(self):
        """Prueba rápida de generación de un documento (S3 compatible)"""
        print(f"\n🧪 PRUEBA RÁPIDA DE GENERACIÓN")
        
        # Probar con contrato indeterminado
        doc_type = 'contrato_indeterminado'
        print(f"   Probando: {doc_type}")
        
        output = self.generate_document(doc_type, 'TEST')
        
        if output and isinstance(output, io.BytesIO) and output.getbuffer().nbytes > 0:
            print(f"   ✅ Prueba exitosa: Documento generado en memoria ({output.getbuffer().nbytes} bytes)")
            return True
        else:
            print(f"   ❌ Prueba fallida")
            return False