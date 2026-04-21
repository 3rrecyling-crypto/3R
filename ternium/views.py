import io
import os
import zipfile
import datetime
import decimal
import boto3
import urllib.request
from io import BytesIO
from openpyxl.drawing.image import Image as OpenXLImage
import time
from . import api_views  
from django.contrib.auth.decorators import permission_required
from .models import ControlManifiestoTrane
from pypdf import PdfReader
from django.views.decorators.http import require_POST  

# --- Django Core & Utils ---
from django.conf import settings
from django.core.paginator import Paginator
from django.core.mail import send_mail
from django.core.files.storage import default_storage
from django.core.exceptions import ValidationError, PermissionDenied
from django.db import IntegrityError, transaction
from django.shortcuts import redirect

# --- CORRECCIÓN AQUÍ: Abs se movió a functions ---
from django.db.models import (
    Count, Sum, F, Avg, Q, FloatField, Case, When, Value, Max, Min, ExpressionWrapper
)
from django.db.models.functions import TruncMonth, Coalesce, Abs

from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse
from django.urls import reverse_lazy
from django.contrib import messages
from django.utils import timezone
from django.utils.decorators import method_decorator

# --- Django Views & Decorators ---
from django.views.generic import ListView, CreateView, UpdateView, DetailView, View, DeleteView
from django.views.decorators.http import require_POST
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.views.decorators.clickjacking import xframe_options_sameorigin, xframe_options_exempt
from django.forms import inlineformset_factory

# --- Third Party (AWS & Excel) ---
from botocore.exceptions import BotoCoreError, NoCredentialsError
from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.styles import Alignment, PatternFill
from openpyxl.utils import get_column_letter

# --- LOCAL IMPORTS: MODELS ---
from .models import (
    Empresa, 
    Lugar, 
    Remision, 
    EntradaMaquila, 
    LineaTransporte,
    Operador, 
    Material, 
    Unidad, 
    Contenedor, 
    DetalleRemision,
    InventarioPatio, 
    Descarga, 
    RegistroLogistico,
    EvidenciaRemision,
    HistorialRemision,
    Plastico,             
    EvidenciaPlastico,
    HistorialPlastico,
    ControlTarima         
)

# --- LOCAL IMPORTS: FORMS ---
from .forms import (
    EmpresaForm, 
    LugarForm, 
    RemisionForm, 
    DetalleRemisionForm,
    EntradaMaquilaForm, 
    LineaTransporteForm, 
    OperadorForm,
    MaterialForm, 
    UnidadForm, 
    ContenedorForm, 
    DescargaForm,
    RegistroLogisticoForm,
    ImportarEvidenciasZipForm,
    EmpresaOrigenesForm,
    PlasticoForm,
    ControlTarimaForm     
)

from .models import PrecioMedline # <--- Asegúrate de importar PrecioMedline
# ==============================================================================
# === NUEVAS FUNCIONES AUXILIARES PARA GESTIONAR ARCHIVOS EN S3 MANUALMENTE ===
# ==============================================================================

def _subir_archivo_a_s3(archivo_obj, s3_ruta_relativa):
    """
    Sube un archivo a S3.
    """
    try:
        # ---> AÑADE ESTAS DOS LÍNEAS <---
        if hasattr(archivo_obj, 'seek'):
            archivo_obj.seek(0) # Rebobina el archivo al byte 0
            
        s3_client = boto3.client(
            's3',
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_S3_REGION_NAME
        )
        
        full_s3_path = f"{settings.AWS_MEDIA_LOCATION}/{s3_ruta_relativa}"

        s3_client.upload_fileobj(
            archivo_obj,
            settings.AWS_STORAGE_BUCKET_NAME,
            full_s3_path
        )
        return s3_ruta_relativa
        
    except (BotoCoreError, NoCredentialsError, Exception) as e:
        print(f"Error al subir el archivo a S3: {e}")
        return None

def _eliminar_archivo_de_s3(ruta_completa_s3):
    """
    Elimina un archivo de S3.
    - `ruta_completa_s3` es la ruta que Django provee (ej: 'media/entradas/foto.jpg'),
      que es lo que Boto3 necesita como 'Key'.
    """
    if not ruta_completa_s3:
        return
    try:
        s3_client = boto3.client(
            's3',
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_S3_REGION_NAME
        )
        s3_client.delete_object(
            Bucket=settings.AWS_STORAGE_BUCKET_NAME,
            Key=str(ruta_completa_s3)
        )
    except (BotoCoreError, NoCredentialsError, Exception) as e:
        print(f"Error al eliminar archivo antiguo de S3: {e}")
        
def _update_inventory_from_remision(remision, revert=False):
    """
    Ajusta el inventario en los patios basado en una remisión.
    CORREGIDO: Asume que los pesos ya están en KG (No multiplica por 1000).
    """
    
    # 1. Si sale de un patio (Resta del inventario)
    if remision.origen and remision.origen.es_patio:
        for detalle in remision.detalles.all():
            if detalle.peso_ld > 0:
                inventario, _ = InventarioPatio.objects.get_or_create(
                    patio=remision.origen, material=detalle.material
                )
                
                # --- CAMBIO: Usar directo el peso_ld ---
                cantidad_kg = decimal.Decimal(detalle.peso_ld) 
                
                current = decimal.Decimal(inventario.cantidad)
                # Si revert=True (al borrar/editar), devolvemos el material (suma)
                # Si es normal, lo restamos.
                new_inv = current + cantidad_kg if revert else current - cantidad_kg
                
                inventario.cantidad = new_inv
                inventario.save()

    # 2. Si llega a un patio (Suma al inventario)
    if remision.destino and remision.destino.es_patio:
        for detalle in remision.detalles.all():
            if detalle.peso_dlv > 0:
                inventario, _ = InventarioPatio.objects.get_or_create(
                    patio=remision.destino, material=detalle.material
                )
                
                # --- CAMBIO: Usar directo el peso_dlv ---
                cantidad_kg = decimal.Decimal(detalle.peso_dlv)
                
                current = decimal.Decimal(inventario.cantidad)
                # Si revert=True (al borrar/editar), quitamos lo que habíamos sumado (resta)
                # Si es normal, lo sumamos.
                new_inv = current - cantidad_kg if revert else current + cantidad_kg
                
                inventario.cantidad = new_inv
                inventario.save()

    # 3. NUEVO: Si hay rechazo y se especificó a qué patio va (Suma al inventario)
    for detalle in remision.detalles.all():
        if detalle.peso_rechazado > 0 and detalle.patio_rechazo and detalle.patio_rechazo.es_patio:
            inventario_rechazo, _ = InventarioPatio.objects.get_or_create(
                patio=detalle.patio_rechazo, material=detalle.material
            )
            
            cantidad_rechazada_kg = decimal.Decimal(detalle.peso_rechazado)
            
            current_rechazo = decimal.Decimal(inventario_rechazo.cantidad)
            # Si revert=True (al borrar/editar), quitamos el rechazo que habíamos sumado (resta)
            # Si es normal, sumamos el rechazo al inventario del patio.
            new_inv_rechazo = current_rechazo - cantidad_rechazada_kg if revert else current_rechazo + cantidad_rechazada_kg
            
            inventario_rechazo.cantidad = new_inv_rechazo
            inventario_rechazo.save()
            
def asignar_folio_medline(remision):
    """
    Evalúa si la remisión es de MEDLINE y contiene Cartón o Archivo.
    SOLO aplica para remisiones con fecha del 1 de Abril de 2026 en adelante.
    """
    import datetime

    # 1. Obtener la fecha base (si es datetime, la convertimos a date)
    fecha_base = remision.fecha or timezone.now().date()
    if isinstance(fecha_base, datetime.datetime):
        fecha_base = fecha_base.date()

    # 2. REGLA DE FECHA: Si es anterior al 1 de abril de 2026, NO hacemos nada.
    if fecha_base < datetime.date(2026, 4, 1):
        return

    # 3. Si ya tiene folio, no lo reasignamos para no alterar históricos ya generados
    # (Si quieres forzar a que se sobreescriban, puedes quitar estas dos líneas)
    if remision.folio_medline:
        return

    # 4. Validar que el origen sea MEDLINE
    if not remision.origen or 'MEDLINE' not in remision.origen.nombre.upper():
        return

    # 5. Validar que al menos un material sea Cartón o Archivo Muerto
    aplica_material = False
    for detalle in remision.detalles.all():
        if detalle.material:
            mat_nom = detalle.material.nombre.upper()
            if "CARTON" in mat_nom or "CARTÓN" in mat_nom or "ARCHIVO" in mat_nom:
                aplica_material = True
                break
                
    if not aplica_material:
        return

    # 6. Generar prefijo 3R-AÑO-MES-
    year = fecha_base.year
    month = f"{fecha_base.month:02d}"
    prefix = f"3R-{year}-{month}-"

    # 7. Buscar el último consecutivo de este mes con bloqueo para evitar race conditions
    with transaction.atomic():
        todos_mes = list(
            Remision.objects
            .select_for_update()
            .filter(folio_medline__startswith=prefix)
        )
        # Ordenar numéricamente para evitar error con strings ≥ 1000
        nums = []
        for r in todos_mes:
            try:
                nums.append(int(r.folio_medline.split('-')[-1]))
            except (ValueError, IndexError, AttributeError):
                pass
        next_num = (max(nums) + 1) if nums else 1

        # 8. Asignar y guardar permanentemente dentro de la misma transacción
        remision.folio_medline = f"{prefix}{next_num:03d}"
        remision.save(update_fields=['folio_medline'])


@login_required
def home_bienvenida(request):
    """
    Vista simple que carga la landing page.
    Accesible para todos los usuarios logueados.
    """
    return render(request, 'ternium/bienvenida.html')


from django.db.models import Count, Sum, Q, Max
from django.http import HttpResponse
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from datetime import datetime
# Tus modelos...
from .models import (
    EntradaMaquila, RegistroLogistico, Lugar, InventarioPatio, 
    Remision, Plastico, ControlTarima
)

@login_required
def dashboard_operaciones_view(request):
    if not request.user.has_perm('ternium.acceso_dashboard_patio'):
         return render(request, 'ternium/home.html', {'has_permission': False})

    # KPIs
    total_entradas = EntradaMaquila.objects.count()
    total_registros_logistica = RegistroLogistico.objects.count()
    
    # Patios
    patios_activos = Lugar.objects.filter(es_patio=True).order_by('nombre')
    patios_data = []
    for patio in patios_activos:
        inventario = InventarioPatio.objects.filter(patio=patio, cantidad__gt=0).select_related('material')
        total_kg = inventario.aggregate(total=Sum('cantidad'))['total'] or 0
        materiales_list = []
        for item in inventario:
            # Barra visual basada en 40 toneladas
            porcentaje = min((float(item.cantidad) / 40000) * 100, 100)
            materiales_list.append({'nombre': item.material.nombre, 'cantidad': item.cantidad, 'porcentaje': porcentaje})
        
        ultima = InventarioPatio.objects.filter(patio=patio).aggregate(Max('ultima_actualizacion'))['ultima_actualizacion__max']
        patios_data.append({'nombre': patio.nombre, 'total_kg': total_kg, 'materiales': materiales_list, 'ultima_actualizacion': ultima})

    # Listas para Filtros del Dashboard
    operaciones_list = Lugar.objects.filter(tipo__in=['ORIGEN', 'DESTINO', 'AMBOS']).distinct().order_by('nombre')

    # --- CAMBIO SOLICITADO: EMPRESAS CON > 10 REMISIONES ---
    empresas_reporte = Empresa.objects.annotate(
        num_remisiones=Count('remisiones')
    ).filter(
        num_remisiones__gt=10  # <--- FILTRO APLICADO
    ).order_by('nombre')
    # -------------------------------------------------------

    # --- NUEVO: REGISTROS DE SALIDAS PARA EL MODAL DE TRANE ---
    # Obtenemos todos los registros ordenados por fecha y ID descendente
    registros_salidas = ControlManifiestoTrane.objects.all().order_by('-fecha_captura', '-id')
    # ----------------------------------------------------------

    context = {
        'has_permission': True,
        'total_entradas': total_entradas,
        'total_registros_logistica': total_registros_logistica,
        'patios_inventario': patios_data,
        'operaciones_list': operaciones_list,
        'empresas_reporte': empresas_reporte, 
        'registros_salidas': registros_salidas,  # <--- AÑADIDO AL CONTEXTO
    }
    
    # ¡MUY IMPORTANTE! 
    # Si esta vista carga el Dashboard de TRANE, el HTML debe ser 'dashboard_trane.html'
    # Si tuvieras otra vista separada solo para TRANE, entonces debes pegar la variable 'registros_salidas' en esa otra vista.
    return render(request, 'ternium/home.html', context)

@login_required
@permission_required('ternium.view_ternium_module', raise_exception=True)
def home_portal_view(request):
    return render(request, 'ternium/home_portal.html')


# --- VISTAS DE ENTRADA MAQUILA ---
@method_decorator(login_required, name='dispatch')
class EntradaMaquilaListView(PermissionRequiredMixin, ListView):
    permission_required = 'ternium.view_ternium_module'
    model = EntradaMaquila
    template_name = 'ternium/lista_entradas.html'
    context_object_name = 'entradas'
    paginate_by = 10
    
    def get_queryset(self):
        # 1. Obtener QuerySet base
        queryset = super().get_queryset()
        
        # 2. Capturar parámetros de la URL
        q = self.request.GET.get('q')
        transporte = self.request.GET.get('transporte')
        calidad = self.request.GET.get('calidad')
        status = self.request.GET.get('status')
        fecha_inicio = self.request.GET.get('fecha_inicio')
        fecha_fin = self.request.GET.get('fecha_fin')
        alerta = self.request.GET.get('alerta')

        # 3. Aplicar Filtros Dinámicos
        
        # Búsqueda general (ID Remito o Boleta)
        if q:
            queryset = queryset.filter(
                Q(c_id_remito__icontains=q) | 
                Q(num_boleta_remision__icontains=q)
            )

        # Filtro por Transporte
        if transporte:
            queryset = queryset.filter(transporte__icontains=transporte)

        # Filtro por Calidad (Match exacto)
        if calidad:
            queryset = queryset.filter(calidad=calidad)

        # Filtro por Estatus
        if status:
            queryset = queryset.filter(status=status)

        # Filtro por Rango de Fechas
        if fecha_inicio:
            queryset = queryset.filter(fecha_ingreso__gte=fecha_inicio)
        if fecha_fin:
            queryset = queryset.filter(fecha_ingreso__lte=fecha_fin)

        # Filtro por Alerta de Discrepancia
        if alerta == 'SI':
            queryset = queryset.filter(alerta=True)
        elif alerta == 'NO':
            queryset = queryset.filter(alerta=False)

        # 4. Ordenamiento por defecto
        return queryset.order_by('-fecha_ingreso', '-creado_en')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Pasamos los parámetros actuales para mantener los filtros en el HTML
        context['filtros'] = self.request.GET
        
        # Obtenemos lista única de calidades para el dropdown
        context['calidades_list'] = EntradaMaquila.objects.exclude(calidad__isnull=True).exclude(calidad='').values_list('calidad', flat=True).distinct().order_by('calidad')
        
        return context


@login_required
def detalle_entrada(request, pk):
    entrada = get_object_or_404(EntradaMaquila, pk=pk)
    return render(request, 'ternium/detalle_entrada.html', {'entrada': entrada})

from .models import EntradaMaquila
from .forms import EntradaMaquilaForm
@login_required
def crear_entrada(request):
    if request.method == 'POST':
        form = EntradaMaquilaForm(request.POST, request.FILES)
        if form.is_valid():
            # Paso 1: Guardar el formulario sin commit
            entrada = form.save(commit=False)
            remito_id = form.cleaned_data.get('c_id_remito', 'sin_remito').strip()

            # Paso 2: Subir archivos manualmente a S3
            archivos_a_subir = {
                'foto_frontal': '1', 'foto_superior_cargada': '2', 'foto_trasera': '3',
                'foto_superior_vacia': '4', 'documento_remision_clientes': '5'
            }

            error_subida = False
            for campo, sufijo in archivos_a_subir.items():
                if campo in request.FILES:
                    archivo = request.FILES[campo]
                    _nombre_base, extension = os.path.splitext(archivo.name)
                    s3_path = f"entradas_maquila/{remito_id}/{remito_id}-{sufijo}{extension}"
                    
                    ruta_guardada = _subir_archivo_a_s3(archivo, s3_path)
                    if ruta_guardada:
                        setattr(entrada, campo, ruta_guardada)
                    else:
                        messages.error(request, f"No se pudo subir el archivo para '{campo}'.")
                        error_subida = True
                        break # Detenemos si hay error
            
            if error_subida:
                empresas_disponibles = Empresa.objects.all().order_by('nombre')
                return render(request, 'ternium/formulario_entrada.html', {
                    'form': form, 
                    'titulo': 'Nuevo Registro de Entrada',
                    'empresas_list': empresas_disponibles
                })

            # Paso 3: Guardar el objeto inicial
            entrada.save()

            # --- FIX: FORZAR ACTUALIZACIÓN DE ESTATUS ---
            # Recargamos desde la BD para asegurar que los FileFields se instancien correctamente
            # y guardamos de nuevo para que el modelo ejecute la lógica de _is_terminado()
            entrada.refresh_from_db()
            if entrada.status == 'PENDIENTE':
                entrada.save() 
            # --------------------------------------------

            messages.success(request, 'Entrada registrada y archivos subidos exitosamente.')
            
            # CAMBIO AQUÍ: Redirigir a la lista en lugar del detalle
            return redirect('lista_entradas')

    else:
        form = EntradaMaquilaForm()
        
    empresas_disponibles = Empresa.objects.all().order_by('nombre')

    context = {
        'form': form, 
        'titulo': 'Nuevo Registro de Entrada',
        'empresas_list': empresas_disponibles
    }
    return render(request, 'ternium/formulario_entrada.html', context)
@login_required
def editar_entrada(request, pk):
    entrada_original = get_object_or_404(EntradaMaquila, pk=pk)
    
    if entrada_original.status == 'CANCELADO':
        messages.error(request, "No se puede editar una entrada CANCELADA.")
        return redirect('lista_entradas')

    if request.method == 'POST':
        form = EntradaMaquilaForm(request.POST, request.FILES, instance=entrada_original)
        if form.is_valid():
            entrada = form.save(commit=False)
            remito_id = form.cleaned_data.get('c_id_remito', 'sin_remito').strip()

            archivos_a_subir = {
                'foto_frontal': '1', 'foto_superior_cargada': '2', 'foto_trasera': '3',
                'foto_superior_vacia': '4', 'documento_remision_clientes': '5'
            }

            for campo, sufijo in archivos_a_subir.items():
                if campo in request.FILES:
                    # 1. Eliminar el archivo antiguo ANTES de subir el nuevo
                    ruta_antigua = getattr(entrada_original, campo)
                    if ruta_antigua and hasattr(ruta_antigua, 'name'):
                        _eliminar_archivo_de_s3(ruta_antigua.name)

                    # 2. Subir el archivo nuevo
                    archivo = request.FILES[campo]
                    _nombre_base, extension = os.path.splitext(archivo.name)
                    s3_path = f"entradas_maquila/{remito_id}/{remito_id}-{sufijo}{extension}"
                    
                    ruta_guardada = _subir_archivo_a_s3(archivo, s3_path)
                    if ruta_guardada:
                        setattr(entrada, campo, ruta_guardada)
                    else:
                        messages.error(request, f"No se pudo actualizar el archivo para '{campo}'.")
                        context = {'form': form, 'object': entrada_original, 'titulo': f'Editar Entrada: {remito_id}'}
                        return render(request, 'ternium/formulario_entrada.html', context)

            entrada.save()

            # --- FIX: FORZAR ACTUALIZACIÓN DE ESTATUS ---
            entrada.refresh_from_db()
            if entrada.status == 'PENDIENTE':
                entrada.save()
            # --------------------------------------------

            messages.success(request, 'Entrada actualizada correctamente.')
            
            # CAMBIO AQUÍ: Redirigir a la lista en lugar del detalle
            return redirect('lista_entradas')
    else:
        form = EntradaMaquilaForm(instance=entrada_original)

    context = {
        'form': form, 'object': entrada_original, 'titulo': f'Editar Entrada: {entrada_original.c_id_remito}'
    }
    return render(request, 'ternium/formulario_entrada.html', context)

@login_required
@require_POST # Asegura que esta vista solo acepte peticiones POST por seguridad
def eliminar_entrada(request, pk):
    """
    Vista para eliminar una EntradaMaquila.
    El borrado de archivos en S3 se maneja automáticamente por la señal post_delete.
    """
    entrada = get_object_or_404(EntradaMaquila, pk=pk)
    
    if entrada.status == 'AUDITADO':
        messages.error(request, "No se puede eliminar una entrada que ya ha sido auditada.")
    else:
        # Al hacer .delete(), la señal que ya tienes se activará 
        # y borrará los archivos de S3 automáticamente.
        entrada.delete()
        messages.success(request, 'La entrada ha sido eliminada exitosamente.')
        
    return redirect('lista_entradas')


@login_required
@require_POST
def auditar_entrada(request, pk):
    entrada = get_object_or_404(EntradaMaquila, pk=pk)
    if entrada.status != 'TERMINADO':
        messages.error(request, 'Esta entrada no cumple los requisitos para ser auditada.')
    else:
        entrada.status = 'AUDITADO'
        entrada.auditado_por = request.user
        entrada.auditado_en = timezone.now()
        entrada.save(update_fields=['status', 'auditado_por', 'auditado_en'])
        messages.success(request, f'La entrada {entrada.c_id_remito} ha sido auditada.')
    return redirect('detalle_entrada', pk=pk)


@login_required
def export_entradas_to_excel(request):
    # 1. Definición de Estilos
    # Lógica de colores para 3R:
    # Rojo = Negativo (Pérdida / Faltante)
    # Verde = Positivo (Ganancia / Sobrante)
    red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid") 
    green_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid") 
    center_align = Alignment(horizontal='center', vertical='center')
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Entradas Maquila"
    
    # 2. Encabezados
    headers = [
        "C_ID_REMITO", 
        "PESO-REMISION", 
        "Num Boleta/Remision Bascula", 
        "Peso Tara", 
        "Peso Bruto", 
        "Peso Neto", 
        "Fecha de Ingreso", 
        "Calidad", 
        "Dif Ton (Neto - Rem)", # Etiqueta corregida para reflejar la fórmula real
        "TRANSPORTE", 
        "FECHA ENTREGA A TERNIUM", 
        "PORCENTAJE DIFERENCIA" # Etiqueta neutral
    ]
    ws.append(headers)
    ws.freeze_panes = 'A2'
    
    # 3. Obtener datos
    entradas = EntradaMaquila.objects.all().order_by('-fecha_ingreso')
    
    for entrada in entradas:
        # Cálculos seguros (evitar None)
        p_rem = entrada.peso_remision or 0
        p_net = entrada.peso_neto or 0
        
        # --- LÓGICA DE NEGOCIO ---
        # Fórmula: Peso Real (Neto) - Peso Papel (Remisión)
        # - Si el resultado es Negativo: Faltó material (Pérdida)
        # - Si el resultado es Positivo: Sobró material (Ganancia)
        diff_ton = p_net - p_rem
        
        # Cálculo del Porcentaje para Excel
        pct_merma = 0
        if p_rem > 0:
            pct_merma = diff_ton / p_rem

        row_data = [
            entrada.c_id_remito, 
            entrada.peso_remision, 
            entrada.num_boleta_remision,
            entrada.peso_tara, 
            entrada.peso_bruto, 
            entrada.peso_neto, 
            entrada.fecha_ingreso,
            entrada.calidad, 
            diff_ton, # Valor calculado (Neto - Rem)
            entrada.transporte, 
            entrada.fecha_entrega_ternium,
            pct_merma # Valor decimal para porcentaje
        ]
        ws.append(row_data)

        # 4. Aplicar estilos a la fila actual
        current_row = ws.max_row
        for col_idx, cell in enumerate(ws[current_row], start=1):
            cell.alignment = center_align
            
            # Columna "Dif Ton" (Columna 9)
            if col_idx == 9:
                cell.number_format = '0.000'

            # Columna "PORCENTAJE DIFERENCIA" (Columna 12)
            if col_idx == 12:
                cell.number_format = '0.00%' # Excel lo mostrará como %
                
                # Colorear según el resultado
                if pct_merma < 0: 
                    # Negativo = Pérdida -> Rojo
                    cell.fill = red_fill
                elif pct_merma > 0:
                    # Positivo = Ganancia -> Verde
                    cell.fill = green_fill

    # 5. Ajustar ancho de columnas automáticamente
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    length = len(str(cell.value))
                    if length > max_length:
                        max_length = length
            except:
                pass
        ws.column_dimensions[column].width = max_length + 4

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Entradas_Maquila_{datetime.date.today()}.xlsx"'
    wb.save(response)
    return response


@method_decorator(login_required, name='dispatch')
class DescargarZipMaquilaView(View):
    """
    Vista CORREGIDA para descargar un ZIP con todos los archivos de una EntradaMaquila.
    Usa boto3 para descargar explícitamente cada archivo desde S3.
    """
    def get(self, request, *args, **kwargs):
        entrada = get_object_or_404(EntradaMaquila, pk=self.kwargs['pk'])
        
        try:
            s3_client = boto3.client(
                's3',
                aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
                region_name=settings.AWS_S3_REGION_NAME
            )
        except (BotoCoreError, NoCredentialsError) as e:
            messages.error(request, f"Error de configuración con S3: {e}")
            return redirect('detalle_entrada', pk=entrada.pk)

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            file_fields = [
                'foto_frontal', 'foto_superior_cargada', 'foto_trasera',
                'foto_superior_vacia', 'documento_remision_clientes'
            ]
            
            archivos_agregados = 0
            for field_name in file_fields:
                file_field = getattr(entrada, field_name)
                # Nos aseguramos que el campo tenga un nombre de archivo guardado
                if file_field and file_field.name:
                    # La ruta completa en S3 (la "Key") incluye el prefijo 'media/'
                    s3_key = f"{settings.AWS_MEDIA_LOCATION}/{file_field.name}"
                    file_content_buffer = io.BytesIO()
                    
                    try:
                        # Descargamos el archivo desde S3 a un buffer en memoria
                        s3_client.download_fileobj(settings.AWS_STORAGE_BUCKET_NAME, s3_key, file_content_buffer)
                        file_content_buffer.seek(0) # Rebobinamos el buffer para leerlo
                        
                        # Usamos el nombre base del archivo para guardarlo en el ZIP
                        filename_in_zip = os.path.basename(file_field.name)
                        zip_file.writestr(filename_in_zip, file_content_buffer.read())
                        archivos_agregados += 1
                    except s3_client.exceptions.ClientError as e:
                        if e.response['Error']['Code'] == '404':
                            print(f"Advertencia: El archivo {s3_key} no fue encontrado en S3 para la entrada {entrada.c_id_remito}.")
                        else:
                            messages.error(request, f"No se pudo descargar el archivo '{s3_key}' de S3.")

        if archivos_agregados == 0:
            messages.warning(request, "No se encontraron archivos en S3 para descargar en este registro.")
            return redirect('detalle_entrada', pk=entrada.pk)

        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="ENTRADA-{entrada.c_id_remito}.zip"'
        return response


# --- VISTAS DE REMISIONES ---

from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin

# Busca la clase RemisionListView y modifícala así:
class RemisionListView(LoginRequiredMixin, PermissionRequiredMixin, ListView):
    permission_required = 'ternium.acceso_remisiones' 
    model = Remision
    template_name = 'ternium/remision_lista.html'
    context_object_name = 'remisiones'
    paginate_by = 15
    
    def has_permission(self):
        user = self.request.user
        return user.has_perm('ternium.acceso_remisiones') or user.has_perm('ternium.view_remision')
    
    def get_queryset(self):
        # 1. Queryset base optimizado
        queryset = Remision.objects.select_related(
            'empresa', 'origen', 'destino'
        ).prefetch_related(
            'detalles__material',
            'facturas' # Agregado por si usas facturas en la tabla
        ).order_by('-pk') 
        
        # 2. FILTRADO POR PERMISOS DE EMPRESA
        if not self.request.user.is_superuser:
            perfil = getattr(self.request.user, 'ternium_profile', None)
            if perfil:
                # Solo mostrar remisiones cuya empresa esté en 'empresas_autorizadas'
                mis_empresas = perfil.empresas_autorizadas.all()
                queryset = queryset.filter(empresa__in=mis_empresas)
            else:
                # Si no tiene perfil o empresas asignadas, no ve nada
                queryset = queryset.none()

        # 3. LÓGICA DE FILTROS
        self.search_params = self.request.GET.copy()
        
        q_remision = self.request.GET.get('q_remision')
        q_folio_ld_raw = self.request.GET.get('q_folio_ld', '').strip()
        q_folio_dlv_raw = self.request.GET.get('q_folio_dlv', '').strip()

        # --- CASO A: SI BUSCAN POR FOLIO PRINCIPAL (BÚSQUEDA GLOBAL) ---
        if q_remision:
            # Buscamos en TODA la base de datos, ignorando las fechas
            queryset = queryset.filter(remision__icontains=q_remision)

        # --- CASO A2: FOLIO LD / DLV — búsqueda global, sin filtros de fecha ---
        elif q_folio_ld_raw or q_folio_dlv_raw:
            if q_folio_ld_raw:
                queryset = queryset.filter(folio_ld__icontains=q_folio_ld_raw)
            if q_folio_dlv_raw:
                queryset = queryset.filter(folio_dlv__icontains=q_folio_dlv_raw)
            return queryset.distinct()

        # --- CASO B: SI NO HAY FOLIO PRINCIPAL (FILTROS NORMALES + FECHAS) ---
        else:
            # Detectamos si hay algún filtro activo (incluyendo los nuevos de carga/descarga)
            filtros_activos = any(
                k.startswith('q_') and v for k, v in self.request.GET.items()
            )
            
            # Si no hay filtros, poner fechas default (últimos 30 días) para no cargar toda la BD
            if not filtros_activos:
                today = timezone.now().date()
                month_ago = today - datetime.timedelta(days=30)
                self.search_params['q_fecha_desde'] = month_ago.strftime('%Y-%m-%d')
                self.search_params['q_fecha_hasta'] = today.strftime('%Y-%m-%d')
            
            # Aplicar filtros de fecha
            q_fecha_desde = self.search_params.get('q_fecha_desde')
            q_fecha_hasta = self.search_params.get('q_fecha_hasta')
            
            if q_fecha_desde:
                queryset = queryset.filter(fecha__gte=q_fecha_desde)
            if q_fecha_hasta:
                queryset = queryset.filter(fecha__lte=q_fecha_hasta)

        # 4. APLICAR RESTO DE FILTROS (Siempre se aplican si están seleccionados)
        filters = {
            'empresa__prefijo__icontains': self.search_params.get('q_prefijo'),
            'detalles__material_id': self.search_params.get('q_material'),
            'origen_id': self.search_params.get('q_origen'),
            'destino_id': self.search_params.get('q_destino'),
            'status': self.search_params.get('q_status'),
        }
        
        for key, value in filters.items():
            if value:
                queryset = queryset.filter(**{key: value})

        # --- NUEVO FILTRO DE OPERADOR LIBRE ---
        q_operador = self.search_params.get('q_operador')
        if q_operador:
            # Busca coincidencias tanto en el nombre del operador del catálogo como en el ingresado manualmente
            queryset = queryset.filter(
                Q(operador__nombre__icontains=q_operador) | 
                Q(operador_manual__icontains=q_operador)
            )

        # --- NUEVOS FILTROS DE FOLIOS OPERATIVOS ---
        q_folio_ld = self.search_params.get('q_folio_ld')
        q_folio_dlv = self.search_params.get('q_folio_dlv')

        if q_folio_ld:
            queryset = queryset.filter(folio_ld__icontains=q_folio_ld)
        
        if q_folio_dlv:
            queryset = queryset.filter(folio_dlv__icontains=q_folio_dlv)
            
        # ==============================================================
        # --- NUEVO: FILTRO DE DESTRUCCIÓN FISCAL (COMPLETO / PENDIENTE)
        # ==============================================================
        q_destruccion = self.search_params.get('q_destruccion')
        if q_destruccion:
            from .models import ConfiguracionManifiesto
            
            # Condición estricta de que tenga los 5 datos mínimos
            completo_q = Q(
                fecha_destruccion__isnull=False,
                destruccion_material_1__isnull=False,
                foto_ingreso__isnull=False,
                foto_vertido__isnull=False,
                foto_destruccion__isnull=False
            ) & ~Q(destruccion_material_1='') & ~Q(foto_ingreso='') & ~Q(foto_vertido='') & ~Q(foto_destruccion='')
            
            if q_destruccion == 'COMPLETO':
                queryset = queryset.filter(completo_q)
            
            elif q_destruccion == 'PENDIENTE':
                # Buscamos cuáles remisiones REQUIEREN destrucción fiscal según las configuraciones
                configuraciones = ConfiguracionManifiesto.objects.all()
                q_requiere = Q()
                for conf in configuraciones:
                    q_requiere |= Q(origen=conf.origen, detalles__material=conf.material)
                
                if configuraciones.exists():
                    # Requiere destrucción PERO le faltan datos (se excluyen los completos)
                    queryset = queryset.filter(q_requiere).exclude(completo_q)
                else:
                    # Si no hay reglas configuradas, no hay pendientes obligatorios
                    queryset = queryset.none()
        # ==============================================================
                
        # Evitar duplicados si se filtra por relaciones ManyToMany (como materiales)
        if filters.get('detalles__material_id') or q_destruccion:
            queryset = queryset.distinct()
            
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Pasamos los parámetros de búsqueda para mantenerlos en los inputs del HTML
        context['search_params'] = self.search_params
        
        # Datos para llenar los selectores del filtro
        context['prefijos'] = Empresa.objects.values_list('prefijo', flat=True).distinct().order_by('prefijo')
        context['materiales'] = Material.objects.all().order_by('nombre')
        context['origenes'] = Lugar.objects.filter(tipo__in=['ORIGEN', 'AMBOS']).order_by('nombre')
        context['destinos'] = Lugar.objects.filter(tipo__in=['DESTINO', 'AMBOS']).order_by('nombre')
        context['estatus_choices'] = Remision.STATUS_CHOICES
        context['total_pendientes'] = self.object_list.filter(status='PENDIENTE').count()
        context['total_cancelados'] = self.object_list.filter(status='CANCELADO').count()
        
        from .models import ControlManifiestoTrane
        context['trane_pendientes'] = ControlManifiestoTrane.objects.filter(remision_vinculada__isnull=True).order_by('-id')
        context['empresas_disponibles'] = Empresa.objects.all().order_by('nombre')
        
        return context
    
def calcular_siguiente_folio(prefijo):
    """
    Calcula el siguiente folio basado en números enteros para evitar
    que 'MTY-999' sea mayor que 'MTY-1000' alfabéticamente.
    """
    prefix_with_dash = f"{prefijo.strip().upper()}-"
    
    # Obtenemos solo los textos de las remisiones que coinciden con el prefijo
    remisiones_existentes = Remision.objects.filter(
        remision__startswith=prefix_with_dash
    ).values_list('remision', flat=True)

    max_num = 0
    
    for rem_str in remisiones_existentes:
        try:
            # Separamos el texto por guiones y tomamos la última parte
            # Ejemplo: "MTY-1005" -> "1005" -> 1005 (int)
            parts = rem_str.split('-')
            if len(parts) > 1:
                # Intentamos convertir a entero para comparar numéricamente
                num = int(parts[-1])
                if num > max_num:
                    max_num = num
        except ValueError:
            continue

    next_num = max_num + 1
    # Rellenamos con ceros a la izquierda (mínimo 3 dígitos)
    return f"{prefix_with_dash}{str(next_num).zfill(3)}"


@login_required
def get_next_remision_number(request, empresa_id):
    """
    Obtiene el siguiente folio para una empresa, verificando permisos.
    """
    try:
        # 1. VERIFICACIÓN DE SEGURIDAD
        if not request.user.is_superuser:
            perfil = getattr(request.user, 'ternium_profile', None)
            # Verificamos si la empresa solicitada está en las autorizadas del usuario
            if not perfil or not perfil.empresas_autorizadas.filter(pk=empresa_id).exists():
                return JsonResponse({'error': 'No tienes permiso para generar folios de esta empresa.'}, status=403)

        empresa = Empresa.objects.get(pk=empresa_id)
        
        if empresa.prefijo:
            # Usamos la función auxiliar que ya tenías
            next_remision = calcular_siguiente_folio(empresa.prefijo)
            return JsonResponse({'next_remision': next_remision, 'is_manual': False})
        else:
            return JsonResponse({'is_manual': True})

    except Empresa.DoesNotExist:
        return JsonResponse({'error': 'Empresa no encontrada'}, status=404)

@login_required
@permission_required('ternium.add_remision', raise_exception=True)
def crear_remision(request):
    DetalleFormSet = inlineformset_factory(
        Remision, DetalleRemision, form=DetalleRemisionForm, extra=1, can_delete=True
    )
    empresa_seleccionada = None
    
    # --- NUEVO: Se inicializan los 5 campos manuales ---
    valores_manuales = {
        'operador': '', 
        'unidad': '', 
        'placas_unidad': '',
        'contenedor': '', 
        'placas_contenedor': ''
    }

    if request.method == 'POST':
        # 1. Capturar valores manuales
        valores_manuales['operador'] = request.POST.get('operador_texto', '').strip().upper()
        valores_manuales['unidad'] = request.POST.get('unidad_texto', '').strip().upper()
        valores_manuales['placas_unidad'] = request.POST.get('placas_unidad_texto', '').strip().upper() # <--- NUEVO
        valores_manuales['contenedor'] = request.POST.get('contenedor_texto', '').strip().upper()
        valores_manuales['placas_contenedor'] = request.POST.get('placas_contenedor_texto', '').strip().upper() # <--- NUEVO
        
        # 2. Buscar empresa seleccionada
        empresa_id = request.POST.get('empresa')
        if empresa_id:
            try:
                empresa_seleccionada = Empresa.objects.get(pk=empresa_id)
            except (Empresa.DoesNotExist, ValueError):
                pass
        
        # 3. Inicializar Formulario Principal
        form = RemisionForm(request.POST, request.FILES, empresa=empresa_seleccionada, user=request.user)
        
        # 4. Filtrar querysets para el Formset
        material_qs = Material.objects.filter(empresas=empresa_seleccionada) if empresa_seleccionada else Material.objects.none()
        lugar_qs = Lugar.objects.filter(Q(empresas=empresa_seleccionada) & (Q(tipo__in=['DESTINO', 'AMBOS']) | Q(es_patio=True))) if empresa_seleccionada else Lugar.objects.none()
        
        # 5. Inicializar Formset
        formset = DetalleFormSet(
            request.POST, 
            prefix='detalles', 
            form_kwargs={'material_queryset': material_qs, 'lugar_queryset': lugar_qs}
        )

        # 6. VALIDACIÓN Y GUARDADO
        if form.is_valid() and formset.is_valid():
            try:
                with transaction.atomic(): 
                    remision = form.save(commit=False)

                    # --- NUEVO: CAPTURAR DATOS DE MATERIALES DEL MODAL ---
                    remision.destruccion_material_1 = request.POST.get('destruccion_material_1')
                    peso1 = request.POST.get('destruccion_peso_1')
                    remision.destruccion_peso_1 = float(peso1) if peso1 else None
                    
                    remision.destruccion_material_2 = request.POST.get('destruccion_material_2')
                    peso2 = request.POST.get('destruccion_peso_2')
                    remision.destruccion_peso_2 = float(peso2) if peso2 else None
                    # -----------------------------------------------------

                    # --- Lógica Manuales ---
                    if valores_manuales['operador']:
                        remision.operador_manual = valores_manuales['operador']; remision.operador = None 
                    else: remision.operador_manual = None

                    # --- NUEVO: Guardar Unidad y Placas ---
                    if valores_manuales['unidad']:
                        remision.unidad_manual = valores_manuales['unidad']
                        remision.placas_unidad_manual = valores_manuales['placas_unidad'] 
                        remision.unidad = None
                    else: 
                        remision.unidad_manual = None
                        remision.placas_unidad_manual = None

                    # --- NUEVO: Guardar Contenedor y Placas ---
                    if valores_manuales['contenedor']:
                        remision.contenedor_manual = valores_manuales['contenedor']
                        remision.placas_contenedor_manual = valores_manuales['placas_contenedor']
                        remision.contenedor = None
                    else: 
                        remision.contenedor_manual = None
                        remision.placas_contenedor_manual = None
                    
                    # --- Seguridad ---
                    if not request.user.is_superuser and remision.empresa:
                        perfil = getattr(request.user, 'ternium_profile', None)
                        if not perfil or not perfil.empresas_autorizadas.filter(pk=remision.empresa.pk).exists():
                            raise PermissionDenied("No tienes permiso para esta empresa.")

                    # --- Generación de Folio ---
                    if empresa_seleccionada and empresa_seleccionada.prefijo:
                        remision.remision = calcular_siguiente_folio(empresa_seleccionada.prefijo)
                    
                    # GUARDAMOS PRIMERO LA REMISIÓN PARA TENER EL ID/FOLIO
                    remision.save() 
                    
                    # --- FOTOS DESTRUCCIÓN FISCAL (NUEVO) ---
                    # --- FOTOS DESTRUCCIÓN FISCAL (NUEVO) ---
                    fotos_destruccion = ['foto_ingreso', 'foto_ingreso_2', 'foto_vertido', 'foto_vertido_2', 'foto_destruccion', 'foto_destruccion_2']
                    hubo_fotos = False
                    for campo_foto in fotos_destruccion:
                        if campo_foto in request.FILES:
                            archivo = request.FILES[campo_foto]
                            nombre_limpio = archivo.name.replace(" ", "_")
                            s3_path = f"remisiones/{remision.remision}/{campo_foto}_{nombre_limpio}"
                            
                            ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                            if ruta_s3:
                                setattr(remision, campo_foto, ruta_s3)
                                hubo_fotos = True
                                
                    if hubo_fotos:
                        remision.save() # Guardar de nuevo si se inyectaron fotos
                    # -------------------------------------------

                    # --- LÓGICA: SUBIR MANIFIESTO A S3 ---
                    if 'manifiesto' in request.FILES:
                        archivo = request.FILES['manifiesto']
                        nombre_limpio = archivo.name.replace(" ", "_")
                        # Ruta: remisiones/[FOLIO]/manifiesto_[NOMBRE_ARCHIVO]
                        s3_path = f"remisiones/{remision.remision}/manifiesto_{nombre_limpio}"
                        
                        ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                        
                        if ruta_s3:
                            remision.manifiesto = ruta_s3
                            remision.save() # Guardamos de nuevo para actualizar el campo
                    # -------------------------------------------

                    # --- LÓGICA: SUBIR BOLETA MEDLINE A S3 (NUEVO) ---
                    if 'boleta_salida_medline' in request.FILES:
                        archivo = request.FILES['boleta_salida_medline']
                        import re
                        nombre_limpio = re.sub(r'[^a-zA-Z0-9_\-\.]', '', archivo.name.replace(" ", "_"))
                        s3_path = f"remisiones/{remision.remision}/medline_{nombre_limpio}"
                        
                        ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                        
                        if ruta_s3:
                            remision.boleta_salida_medline = ruta_s3
                            remision.save() # Guardamos de nuevo para actualizar el campo
                    # -------------------------------------------

                    # --- Archivos Múltiples de Evidencia (S3 Manual) ---
                    files = request.FILES.getlist('evidencia_documento')
                    if files:
                        for i, archivo in enumerate(files):
                            nombre_limpio = archivo.name.replace(" ", "_")
                            s3_path = f"remisiones/{remision.remision}/evidencia_{i}_{nombre_limpio}"
                            
                            ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                            
                            if ruta_s3:
                                EvidenciaRemision.objects.create(remision=remision, archivo=ruta_s3)
                            else:
                                print(f"❌ Error al subir archivo {nombre_limpio} a S3")
                    
                    # --- Historial ---
                    HistorialRemision.objects.create(
                        remision=remision,
                        usuario=request.user,
                        cambio="Creación de la remisión"
                    )
                    
                    # --- Formset (Detalles) ---
                    detalles = formset.save(commit=False)
                    for detalle in detalles:
                        detalle.remision = remision
                        
                        # FORZAR CLIENTE = DESTINO
                        if remision.destino:
                            detalle.cliente = remision.destino
                        
                        # Guardamos directamente en KG
                        detalle.save()
                    
                    for obj in formset.deleted_objects: obj.delete()

                    # ==========================================================
                    # NUEVO: ASIGNAR FOLIO MEDLINE YA QUE SE GUARDÓ EL MATERIAL
                    # ==========================================================
                    asignar_folio_medline(remision)

                    # --- Inventarios ---
                    _update_inventory_from_remision(remision, revert=False)
                    enviar_alerta_merma(remision)
                    messages.success(request, f'Remisión {remision.remision} creada exitosamente.')
                    return redirect('remision_lista')
                
            except Exception as e:
                print(f"❌ Error Crítico al guardar: {e}")
                messages.error(request, f'Error crítico al guardar: {e}')
        
        else:
            print("❌ El formulario NO es válido.")
            if form.errors:
                messages.error(request, f"Error en datos generales: {form.errors}")
            if formset.errors:
                errores_reales = [e for e in formset.errors if e]
                if errores_reales:
                    messages.error(request, f"Error en los materiales: {errores_reales}")

    else:
        # Método GET
        form = RemisionForm(user=request.user)
        formset = DetalleFormSet(prefix='detalles', form_kwargs={'material_queryset': Material.objects.none(), 'lugar_queryset': Lugar.objects.none()})

    # --- NUEVO: Obtener configuración para el JS del Frontend ---
    configs_destruccion = list(ConfiguracionManifiesto.objects.values('origen_id', 'material_id'))
    
    context = {
        'form': form, 
        'formset': formset, 
        'titulo': 'Nueva Remisión', 
        'is_editing': False, 
        'valores_manuales': valores_manuales, 
        'remision': None,
        'configs_destruccion': json.dumps(configs_destruccion) # <--- Variable para HTML
    }
    return render(request, 'ternium/remision_formulario.html', context)


@login_required
@permission_required('ternium.change_remision', raise_exception=True)
def editar_remision(request, pk):
    remision_original = get_object_or_404(Remision, pk=pk)
    
    # 1. Seguridad
    if not request.user.is_superuser:
        perfil = getattr(request.user, 'ternium_profile', None)
        if not perfil or not perfil.empresas_autorizadas.filter(pk=remision_original.empresa.pk).exists():
             messages.error(request, "Acceso denegado.")
             return redirect('remision_lista')

    if remision_original.status == 'AUDITADO':
        messages.error(request, 'No se puede editar una remisión auditada.')
        return redirect('detalle_remision', pk=remision_original.pk)

    DetalleFormSet = inlineformset_factory(
        Remision, DetalleRemision, form=DetalleRemisionForm, extra=0, can_delete=True, min_num=1
    )
    
    empresa_para_form = remision_original.empresa
    
    # --- NUEVO: Precargar TODOS los valores manuales ---
    valores_manuales = {
        'operador': remision_original.operador_manual or '',
        'unidad': remision_original.unidad_manual or '',
        'placas_unidad': remision_original.placas_unidad_manual or '',
        'contenedor': remision_original.contenedor_manual or '',
        'placas_contenedor': remision_original.placas_contenedor_manual or ''
    }

    if request.method == 'POST':
        nuevo_op_manual = request.POST.get('operador_texto', '').strip().upper()
        nuevo_uni_manual = request.POST.get('unidad_texto', '').strip().upper()
        nueva_placa_uni_manual = request.POST.get('placas_unidad_texto', '').strip().upper() # <--- NUEVO
        nuevo_cont_manual = request.POST.get('contenedor_texto', '').strip().upper()
        nueva_placa_cont_manual = request.POST.get('placas_contenedor_texto', '').strip().upper() # <--- NUEVO

        # Actualizar el diccionario para que el form lo retenga si hay error de validación
        valores_manuales['operador'] = nuevo_op_manual
        valores_manuales['unidad'] = nuevo_uni_manual
        valores_manuales['placas_unidad'] = nueva_placa_uni_manual
        valores_manuales['contenedor'] = nuevo_cont_manual
        valores_manuales['placas_contenedor'] = nueva_placa_cont_manual

        form = RemisionForm(request.POST, request.FILES, instance=remision_original, empresa=empresa_para_form, user=request.user)
        
        material_qs = Material.objects.filter(empresas=empresa_para_form)
        lugar_qs = Lugar.objects.filter(Q(empresas=empresa_para_form) & (Q(tipo__in=['DESTINO', 'AMBOS']) | Q(es_patio=True))) if empresa_para_form else Lugar.objects.none()

        formset = DetalleFormSet(
            request.POST, instance=remision_original, prefix='detalles',
            form_kwargs={'material_queryset': material_qs, 'lugar_queryset': lugar_qs}
        )
        
        if form.is_valid() and formset.is_valid():
            try:
                with transaction.atomic():
                    remision_db = Remision.objects.get(pk=pk)
                    cambios_log = []

                    def get_val(manual, catalogo):
                        if manual and str(manual).strip(): return f"{manual} (Manual)"
                        if catalogo: return f"{catalogo}" 
                        return "Vacío"

                    # --- Comparaciones de Campos Manuales Principales ---
                    val_ant = get_val(remision_db.operador_manual, remision_db.operador)
                    val_nue = get_val(nuevo_op_manual, form.cleaned_data.get('operador'))
                    if val_ant != val_nue: cambios_log.append(f"Operador: '{val_ant}' ➝ '{val_nue}'")

                    val_ant = get_val(remision_db.unidad_manual, remision_db.unidad)
                    val_nue = get_val(nuevo_uni_manual, form.cleaned_data.get('unidad'))
                    if val_ant != val_nue: cambios_log.append(f"Unidad: '{val_ant}' ➝ '{val_nue}'")

                    val_ant = get_val(remision_db.contenedor_manual, remision_db.contenedor)
                    val_nue = get_val(nuevo_cont_manual, form.cleaned_data.get('contenedor'))
                    if val_ant != val_nue: cambios_log.append(f"Contenedor: '{val_ant}' ➝ '{val_nue}'")

                    # --- NUEVO: Comparar Placas ---
                    if (remision_db.placas_unidad_manual or '') != nueva_placa_uni_manual:
                        cambios_log.append(f"Placas Unidad: '{remision_db.placas_unidad_manual or 'Vacío'}' ➝ '{nueva_placa_uni_manual or 'Vacío'}'")

                    if (remision_db.placas_contenedor_manual or '') != nueva_placa_cont_manual:
                        cambios_log.append(f"Placas Contenedor: '{remision_db.placas_contenedor_manual or 'Vacío'}' ➝ '{nueva_placa_cont_manual or 'Vacío'}'")

                    # --- Comparaciones Campos Simples ---
                    # Agregamos los campos de placas para que no salgan doble en el log
                    campos_excluidos = [
                        'operador', 'unidad', 'contenedor', 'evidencia_documento', 
                        'operador_manual', 'unidad_manual', 'contenedor_manual', 
                        'empresa', 'manifiesto', 'placas_unidad_manual', 'placas_contenedor_manual',
                        'foto_ingreso', 'foto_ingreso_2', 'foto_vertido', 'foto_vertido_2', 'foto_destruccion', 'foto_destruccion_2',
                        'boleta_salida_medline' # <--- Añadido aquí para evitar error en log de string vs file
                    ]
                    if form.has_changed():
                        for field_name in form.changed_data:
                            if field_name not in campos_excluidos:
                                label = form.fields[field_name].label or field_name
                                val_nue = form.cleaned_data.get(field_name)
                                val_ant = getattr(remision_db, field_name)
                                s_ant = str(val_ant) if val_ant not in [None, ''] else "Vacío"
                                s_nue = str(val_nue) if val_nue not in [None, ''] else "Vacío"
                                if s_ant != s_nue: cambios_log.append(f"{label}: '{s_ant}' ➝ '{s_nue}'")

                    # --- Comparaciones Formset (Materiales) ---
                    if formset.has_changed():
                        for f in formset:
                            if not f.instance.pk and f.has_changed() and not f.cleaned_data.get('DELETE'):
                                mat = f.cleaned_data.get('material')
                                peso = f.cleaned_data.get('peso_ld')
                                unidad = f.cleaned_data.get('unidad_medida') 
                                cambios_log.append(f"Material AGREGADO: {mat} ({peso} {unidad})")
                            
                            elif f.cleaned_data.get('DELETE') and f.instance.pk:
                                mat = f.instance.material
                                cambios_log.append(f"Material ELIMINADO: {mat}")
                            
                            elif f.instance.pk and f.has_changed():
                                for campo in f.changed_data:
                                    if campo in ['peso_ld', 'peso_dlv', 'material', 'cliente', 'unidad_medida']:
                                        cambios_log.append(f"Detalle modificado: {campo}")

                    # --- Guardado ---
                    _update_inventory_from_remision(remision_db, revert=True)
                    remision = form.save(commit=False)

                    # --- NUEVO: CAPTURAR DATOS DE MATERIALES DEL MODAL ---
                    remision.destruccion_material_1 = request.POST.get('destruccion_material_1')
                    peso1 = request.POST.get('destruccion_peso_1')
                    remision.destruccion_peso_1 = float(peso1) if peso1 else None
                    
                    remision.destruccion_material_2 = request.POST.get('destruccion_material_2')
                    peso2 = request.POST.get('destruccion_peso_2')
                    remision.destruccion_peso_2 = float(peso2) if peso2 else None
                    # -----------------------------------------------------

                    if nuevo_op_manual: remision.operador_manual = nuevo_op_manual; remision.operador = None 
                    else: remision.operador_manual = None
                    
                    # --- NUEVO: Guardado de Unidades y Placas Manuales ---
                    if nuevo_uni_manual: 
                        remision.unidad_manual = nuevo_uni_manual
                        remision.placas_unidad_manual = nueva_placa_uni_manual
                        remision.unidad = None
                    else: 
                        remision.unidad_manual = None
                        remision.placas_unidad_manual = None
                        
                    if nuevo_cont_manual: 
                        remision.contenedor_manual = nuevo_cont_manual
                        remision.placas_contenedor_manual = nueva_placa_cont_manual
                        remision.contenedor = None
                    else: 
                        remision.contenedor_manual = None
                        remision.placas_contenedor_manual = None

                    # --- ACTUALIZAR FOTOS DESTRUCCIÓN FISCAL (NUEVO) ---
                    fotos_destruccion = ['foto_ingreso', 'foto_ingreso_2', 'foto_vertido', 'foto_vertido_2', 'foto_destruccion', 'foto_destruccion_2']
                    for campo_foto in fotos_destruccion:
                        if campo_foto in request.FILES:
                            # 1. Eliminar foto anterior de S3 si existe
                            foto_actual = getattr(remision_db, campo_foto)
                            if foto_actual and hasattr(foto_actual, 'name') and foto_actual.name:
                                _eliminar_archivo_de_s3(foto_actual.name)

                            # 2. Subir nueva foto
                            archivo = request.FILES[campo_foto]
                            nombre_limpio = archivo.name.replace(" ", "_")
                            s3_path = f"remisiones/{remision.remision}/{campo_foto}_{nombre_limpio}"
                            
                            ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                            if ruta_s3:
                                setattr(remision, campo_foto, ruta_s3)
                                cambios_log.append(f"Se actualizó la fotografía: {campo_foto}")
                    # -------------------------------------------

                    # --- ACTUALIZAR MANIFIESTO ---
                    if 'manifiesto' in request.FILES:
                        # 1. Eliminar archivo anterior de S3 si existe
                        if remision_db.manifiesto and hasattr(remision_db.manifiesto, 'name') and remision_db.manifiesto.name:
                            _eliminar_archivo_de_s3(remision_db.manifiesto.name)
                        
                        # 2. Subir nuevo archivo
                        archivo = request.FILES['manifiesto']
                        nombre_limpio = archivo.name.replace(" ", "_")
                        s3_path = f"remisiones/{remision.remision}/manifiesto_{nombre_limpio}"
                        
                        ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                        
                        if ruta_s3:
                            remision.manifiesto = ruta_s3
                            cambios_log.append("Se actualizó el documento Manifiesto")
                    # -------------------------------------------

                    # --- ACTUALIZAR BOLETA MEDLINE (NUEVO) ---
                    if 'boleta_salida_medline' in request.FILES:
                        # 1. Eliminar archivo anterior de S3 si existe
                        if remision_db.boleta_salida_medline and hasattr(remision_db.boleta_salida_medline, 'name') and remision_db.boleta_salida_medline.name:
                            _eliminar_archivo_de_s3(remision_db.boleta_salida_medline.name)
                        
                        # 2. Subir nuevo archivo
                        archivo = request.FILES['boleta_salida_medline']
                        nombre_limpio = archivo.name.replace(" ", "_")
                        s3_path = f"remisiones/{remision.remision}/medline_{nombre_limpio}"
                        
                        ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                        
                        if ruta_s3:
                            remision.boleta_salida_medline = ruta_s3
                            cambios_log.append("Se actualizó la Boleta MEDLINE")
                    # -------------------------------------------

                    remision.save()

                    # Archivos de Evidencia (Múltiples)
                    if request.FILES.getlist('evidencia_documento'):
                        archivos = request.FILES.getlist('evidencia_documento')
                        conteo_existente = remision.evidencias.count()
                        for i, archivo in enumerate(archivos):
                            nombre_limpio = archivo.name.replace(" ", "_")
                            idx = conteo_existente + i + 1 
                            s3_path = f"remisiones/{remision.remision}/evidencia_{idx}_{nombre_limpio}"
                            ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                            if ruta_s3:
                                EvidenciaRemision.objects.create(remision=remision, archivo=ruta_s3)
                        cambios_log.append(f"Se agregaron {len(archivos)} archivos de evidencia")
                    
                    # Historial
                    if cambios_log:
                        cambios_unicos = list(dict.fromkeys(cambios_log))
                        texto_historial = " | ".join(cambios_unicos)
                        HistorialRemision.objects.create(remision=remision, usuario=request.user, cambio=texto_historial)

                    # --- Guardado Formset ---
                    detalles = formset.save(commit=False)
                    for detalle in detalles:
                        detalle.remision = remision
                        # Se guarda directo en KG
                        detalle.save()
                    
                    for obj in formset.deleted_objects:
                        obj.delete()

                    # ==========================================================
                    # NUEVO: ASIGNAR FOLIO MEDLINE YA QUE SE GUARDÓ EL MATERIAL
                    # ==========================================================
                    asignar_folio_medline(remision)
                    
                    _update_inventory_from_remision(remision, revert=False)
                    enviar_alerta_merma(remision) 
                    messages.success(request, 'Remisión actualizada correctamente.')
                    return redirect('detalle_remision', pk=remision.pk)
            except Exception as e:
                messages.error(request, f'Error: {e}')
    else:
        form = RemisionForm(instance=remision_original, empresa=empresa_para_form, user=request.user)
        material_qs = Material.objects.filter(empresas=empresa_para_form)
        lugar_qs = Lugar.objects.filter(empresas=empresa_para_form, tipo__in=['DESTINO', 'AMBOS'])
        formset = DetalleFormSet(
            instance=remision_original, prefix='detalles',
            form_kwargs={'material_queryset': material_qs, 'lugar_queryset': lugar_qs}
        )

    # --- NUEVO: Obtener configuración para el JS del Frontend ---
    configs_destruccion = list(ConfiguracionManifiesto.objects.values('origen_id', 'material_id'))

    context = {
        'form': form, 'formset': formset, 'remision': remision_original, 
        'is_editing': True, 'valores_manuales': valores_manuales,
        'configs_destruccion': json.dumps(configs_destruccion) # <--- Variable para HTML
    }
    return render(request, 'ternium/remision_formulario.html', context)
@login_required
def export_reportes_especificos(request):
    tipo_reporte = request.GET.get('tipo_reporte', '')
    
    # Nuevos parámetros de fechas
    tipo_filtro = request.GET.get('tipo_filtro', 'mes') # 'mes' o 'rango'
    mes_req = request.GET.get('mes', '')
    fecha_inicio = request.GET.get('fecha_inicio', '')
    fecha_fin = request.GET.get('fecha_fin', '')
    
    precio_hudson_req = request.GET.get('precio_hudson', '0')
    try:
        precio_hudson_val = float(precio_hudson_req)
    except ValueError:
        precio_hudson_val = 0.0

    # 1. Base Queryset (Solo Terminados y Auditados)
    queryset = Remision.objects.filter(status__in=['TERMINADO', 'AUDITADO']).select_related(
        'empresa', 'origen', 'destino', 'linea_transporte', 'operador', 'unidad'
    ).prefetch_related('detalles', 'detalles__material', 'evidencias')

    # 2. Filtrar por tipo de reporte (Origen)
    if tipo_reporte in ['huaxing', 'bachoco', 'trane']:
        queryset = queryset.filter(origen__nombre__icontains=tipo_reporte.upper())
    elif tipo_reporte == 'hudson':
        queryset = queryset.filter(origen__nombre__icontains='HUDSON')
    else:
        return HttpResponse("Tipo de reporte no válido.", status=400)

    # 3. Aplicar Filtro Dinámico de Fechas
    if tipo_filtro == 'mes' and mes_req:
        try:
            year, month = mes_req.split('-')
            queryset = queryset.filter(fecha__year=year, fecha__month=month)
        except ValueError:
            pass
    elif tipo_filtro == 'rango':
        if fecha_inicio:
            queryset = queryset.filter(fecha__gte=fecha_inicio)
        if fecha_fin:
            queryset = queryset.filter(fecha__lte=fecha_fin)

    # 4. Preparar el Excel
    wb = Workbook()
    ws = wb.active
    center_style = Alignment(horizontal='center', vertical='center')

    # ==========================================
    # LÓGICA PARA INSERCIÓN DE IMAGEN DESDE S3
    # ==========================================
    # Dejamos 5 filas vacías para que la tabla comience en la fila 6
    start_row_table = 6 
    for _ in range(start_row_table - 1):
        ws.append([])

    # Mapa de imágenes en S3 (Exactamente con las extensiones que proporcionaste)
    mapa_imagenes = {
        'bachoco': 'BACHOCO.png',
        'huaxing': 'HUAXING.jpg',
        'hudson': 'HUDSON.png',
        'trane': 'TRANE.png'
    }

    nombre_img = mapa_imagenes.get(tipo_reporte)
    if nombre_img:
        try:
            # Reemplaza '3rrecycling' por el nombre exacto de tu bucket si es distinto en la URL
            s3_url = f"https://3rrecycling.s3.amazonaws.com/static/Fotos_Reportes/{nombre_img}"
            req = urllib.request.Request(s3_url, headers={'User-Agent': 'Mozilla/5.0'})
            
            with urllib.request.urlopen(req) as response:
                img_data = BytesIO(response.read())
                
            logo = OpenXLImage(img_data)
            
            # Ajuste de tamaño manteniendo la proporción (Alto fijo en 75px)
            aspect_ratio = logo.width / logo.height
            logo.height = 75
            logo.width = 75 * aspect_ratio
            
            ws.add_image(logo, 'A1') # Se coloca en la celda A1
        except Exception as e:
            print(f"No se pudo cargar la imagen de S3: {e}") # Falla en silencio para que el Excel sí se descargue

    # =========================================================================
    # ESTRUCTURA PARA HUAXING, BACHOCO O TRANE
    # =========================================================================
    if tipo_reporte in ['huaxing', 'bachoco', 'trane']:
        ws.title = f"Reporte_{tipo_reporte.capitalize()}"
        headers = [
            "Remisión", "Origen", "Destino", "Fecha", "Folio Carga", 
            "Folio Descarga", "Operador", "Material", "Peso Carga (Kg)", 
            "Peso Descarga (Kg)", "Links Evidencias"
        ]
        ws.append(headers)
        
        # Estilo encabezados (que ahora están en ws[start_row_table])
        header_font = Font(bold=True)
        yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        for cell in ws[start_row_table]: 
            cell.alignment = center_style
            cell.font = header_font
            cell.fill = yellow_fill

        for remision in queryset:
            evidencias_urls = [request.build_absolute_uri(ev.archivo.url) for ev in remision.evidencias.all() if ev.archivo]
            texto_evidencias = "Descargar Evidencia" if len(evidencias_urls) == 1 else ("Múltiples Archivos" if len(evidencias_urls) > 1 else "Sin adjuntos")
            
            for d in remision.detalles.all():
                row_data = [
                    remision.remision,
                    remision.origen.nombre if remision.origen else '',
                    remision.destino.nombre if remision.destino else '',
                    remision.fecha.strftime("%d/%m/%Y") if remision.fecha else "",
                    remision.folio_ld or '',
                    remision.folio_dlv or '',
                    remision.operador.nombre if remision.operador else (remision.operador_manual or ''),
                    d.material.nombre if d.material else "S/M",
                    float(d.peso_ld or 0),
                    float(d.peso_dlv or 0),
                    texto_evidencias
                ]
                ws.append(row_data)
                
                current_row = ws.max_row
                for col_idx in range(1, 12):
                    ws.cell(row=current_row, column=col_idx).alignment = center_style
                    
                ws.cell(row=current_row, column=9).number_format = '#,##0.000'
                ws.cell(row=current_row, column=10).number_format = '#,##0.000'
                
                # Link de evidencia
                if len(evidencias_urls) == 1:
                    cell_link = ws.cell(row=current_row, column=11)
                    cell_link.hyperlink = evidencias_urls[0]
                    cell_link.font = Font(color="0000FF", underline="single")

    # =========================================================================
    # ESTRUCTURA PARA HUDSON
    # =========================================================================
    elif tipo_reporte == 'hudson':
        ws.title = "Reporte_Hudson"
        headers = [
            "Remisión", "Empresa", "Operador", "Unidad", "Material", 
            "Peso Carga (Kg)", "Folio Carga", "Inicia Carga", "PRECIO", "VENTA"
        ]
        ws.append(headers)
        
        pink_fill = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid")
        header_pink = PatternFill(start_color="FF66B2", end_color="FF66B2", fill_type="solid")
        green_header = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
        
        for col_idx, cell in enumerate(ws[start_row_table], start=1): 
            cell.alignment = center_style
            cell.font = Font(bold=True, color="FFFFFF" if col_idx <= 8 else "000000")
            if col_idx <= 8: cell.fill = header_pink
            elif col_idx in [9, 10]: cell.fill = green_header

        total_peso = 0
        total_venta = 0

        for remision in queryset:
            val_inicia = timezone.localtime(remision.inicia_ld).date() if remision.inicia_ld else remision.fecha
            fecha_str = val_inicia.strftime("%d/%m/%Y") if val_inicia else ""

            for d in remision.detalles.all():
                mat_nom = d.material.nombre.upper() if d.material else "S/M"
                peso_ld = float(d.peso_ld or 0)
                
                venta_calculada = peso_ld * precio_hudson_val
                
                total_peso += peso_ld
                total_venta += venta_calculada

                row_data = [
                    remision.remision,
                    remision.empresa.nombre if remision.empresa else 'HUDSON',
                    remision.operador.nombre if remision.operador else (remision.operador_manual or ''),
                    remision.unidad.internal_id if remision.unidad else (remision.unidad_manual or ''),
                    mat_nom,
                    peso_ld,
                    remision.folio_ld or '',
                    fecha_str,
                    precio_hudson_val,
                    venta_calculada
                ]
                ws.append(row_data)
                
                current_row = ws.max_row
                for col_idx in range(1, 11):
                    cell = ws.cell(row=current_row, column=col_idx)
                    cell.alignment = center_style
                    if col_idx <= 8: cell.fill = pink_fill
                    if col_idx == 6: cell.number_format = '#,##0.000'
                    if col_idx == 9: cell.number_format = '"$"#,##0.00'
                    if col_idx == 10: cell.number_format = '"$"#,##0.00'

        # Fila de Totales de Hudson
        ws.append(["", "", "", "", "TOTALES:", total_peso, "", "", "", total_venta])
        last_row = ws.max_row
        ws.cell(row=last_row, column=6).font = Font(bold=True)
        ws.cell(row=last_row, column=6).number_format = '#,##0.000'
        ws.cell(row=last_row, column=10).font = Font(bold=True)
        ws.cell(row=last_row, column=10).number_format = '"$"#,##0.00'

    # Ajustar ancho de columnas automáticamente
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                # Evitar revisar toda la imagen si por alguna razón iteramos sobre ella
                if len(str(cell.value)) > max_length and cell.value:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[col_letter].width = (max_length + 2)

    # 5. Retornar el archivo Excel
    nombre_archivo = f"Reporte_{tipo_reporte.capitalize()}_{timezone.now().strftime('%Y%m%d')}.xlsx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
    wb.save(response)
    
    return response
@login_required
@require_POST
@permission_required('ternium.can_audit_remision', raise_exception=True) 
def auditar_remision(request, pk):
    remision = get_object_or_404(Remision, pk=pk)
    
    if remision.status == 'TERMINADO':
        try:
            with transaction.atomic():
                remision.status = 'AUDITADO'
                remision.auditado_por = request.user
                remision.auditado_en = timezone.now()
                remision.save()

                HistorialRemision.objects.create(
                    remision=remision,
                    usuario=request.user,
                    cambio="ESTATUS CAMBIADO A: AUDITADO"
                )

            messages.success(request, f'La remisión {remision.remision} ha sido auditada.')
        except Exception as e:
            messages.error(request, f"Error al auditar: {e}")
    else:
        messages.error(request, 'Esta remisión no puede ser auditada (debe estar TERMINADA).')
        
    return redirect('detalle_remision', pk=pk)


@method_decorator(login_required, name='dispatch')
class RemisionDeleteView(DeleteView):
    model = Remision
    template_name = 'ternium/remision_confirm_delete.html'
    success_url = reverse_lazy('remision_lista')

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        
        if self.object.status == 'AUDITADO':
            messages.error(self.request, 'No se puede eliminar una remisión auditada.')
            return redirect('remision_lista')
            
        try:
            with transaction.atomic():
                # Revertir el inventario primero
                _update_inventory_from_remision(self.object, revert=True)

                # --- LÓGICA MANUAL DE BORRADO EN S3 ---
                if self.object.evidencia_carga and hasattr(self.object.evidencia_carga, 'name'):
                    _eliminar_archivo_de_s3(self.object.evidencia_carga.name)
                if self.object.evidencia_descarga and hasattr(self.object.evidencia_descarga, 'name'):
                    _eliminar_archivo_de_s3(self.object.evidencia_descarga.name)

                # Llamar al método original de borrado de la base de datos
                response = super().delete(request, *args, **kwargs)

            messages.success(self.request, f'Remisión {self.object.remision} eliminada y el inventario ha sido ajustado.')
            return response
        except Exception as e:
            messages.error(self.request, f"Ocurrió un error al eliminar la remisión: {e}")
            return redirect('remision_lista')


# --- VISTAS DE REGISTRO LOGISTICO ---

@method_decorator(login_required, name='dispatch')
class RegistroLogisticoListView(ListView):
    model = RegistroLogistico
    template_name = 'ternium/lista_logistica_ternium.html'
    context_object_name = 'registros'
    paginate_by = 20

    def get_queryset(self):
        # ... (Tu código get_queryset existente se queda igual) ...
        queryset = super().get_queryset().select_related('transportista', 'material').order_by('-id')
        
        q = self.request.GET.get('q')
        transportista = self.request.GET.get('transportista')
        material_id = self.request.GET.get('material')
        status = self.request.GET.get('status')
        fecha_inicio = self.request.GET.get('fecha_inicio')
        fecha_fin = self.request.GET.get('fecha_fin')
        merma = self.request.GET.get('merma') 

        if q:
            queryset = queryset.filter(Q(remision__icontains=q) | Q(boleta_bascula__icontains=q))
        if transportista:
            queryset = queryset.filter(transportista__nombre__icontains=transportista)
        if material_id:
            queryset = queryset.filter(material__id=material_id)
        if status:
            queryset = queryset.filter(status=status)
        if fecha_inicio:
            queryset = queryset.filter(fecha_carga__gte=fecha_inicio)
        if fecha_fin:
            queryset = queryset.filter(fecha_carga__lte=fecha_fin)

        if merma:
            queryset = queryset.annotate(
                abs_diff=Abs(F('toneladas_remisionadas') - F('toneladas_recibidas')),
                pct_calc=Case(
                    When(toneladas_remisionadas__gt=0, then=F('abs_diff') / F('toneladas_remisionadas') * 100),
                    default=Value(0.0),
                    output_field=FloatField()
                )
            )
            if merma == 'SI':
                queryset = queryset.filter(pct_calc__gt=1.0)
            elif merma == 'NO':
                queryset = queryset.filter(pct_calc__lte=1.0)
            
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # 1. Mantener los filtros en la paginación y formulario
        context['filtros'] = self.request.GET
        
        # 2. Cargar lista de materiales
        try:
            ternium = Empresa.objects.get(nombre__iexact="TERNIUM")
            context['materiales_list'] = Material.objects.filter(empresas=ternium).order_by('nombre')
        except Empresa.DoesNotExist:
            context['materiales_list'] = Material.objects.all().order_by('nombre')
            
        # =========================================================
        # 3. LÓGICA KPI: Totales del Mes Actual
        # =========================================================
        now = timezone.now()
        
        # Filtramos por mes/año actual Y SOLO TERMINADOS O AUDITADOS
        qs_mes_actual = RegistroLogistico.objects.filter(
            fecha_carga__year=now.year,
            fecha_carga__month=now.month,
            status__in=['TERMINADO', 'AUDITADO']  # <--- FILTRO AÑADIDO AQUÍ
        )
        
        # Calculamos la suma
        totales = qs_mes_actual.aggregate(
            total_remisionadas=Sum('toneladas_remisionadas'),
            total_recibidas=Sum('toneladas_recibidas')
        )
        
        # Asignamos al contexto (0 si no hay datos)
        context['kpi_remisionadas_mes'] = totales['total_remisionadas'] or 0
        context['kpi_recibidas_mes'] = totales['total_recibidas'] or 0
        
        # Nombre del mes en español
        meses_es = {
            1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
            5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
            9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
        }
        context['mes_actual_nombre'] = meses_es[now.month]
        
        return context

@method_decorator(login_required, name='dispatch')
class RegistroLogisticoDetailView(DetailView):
    model = RegistroLogistico
    template_name = 'ternium/detalle_logistica_ternium.html' 
    context_object_name = 'registro'


@method_decorator(login_required, name='dispatch')
class RegistroLogisticoCreateView(CreateView):
    model = RegistroLogistico
    form_class = RegistroLogisticoForm
    template_name = 'ternium/formulario_logistica_ternium.html'

    def get_success_url(self):
        return reverse_lazy('lista_registros_logistica')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nuevo Registro de Logística'
        return context
    
    def form_valid(self, form):
        self.object = form.save(commit=False)
        remision_num = form.cleaned_data.get('remision', 'sin_remision').strip()
        request = self.request

        # Mapa de campos a sufijos
        archivos_map = {
            'pdf_registro_camion_remision': '4.pdf', 
            'pdf_hoja_circulacion': '5.pdf',
            'foto_superior_vacia': '0', 
            'foto_frontal': '1',
            'foto_superior_llena': '2', 
            'foto_trasera': '3',
            'pdf_factura': 'factura.pdf', # <--- NUEVO
            'xml_factura': 'factura.xml', # <--- NUEVO
            'acuse_pdf': 'acuse.pdf'      # <--- NUEVO
        }

        # Timestamp único para este guardado (evita caché desde el inicio)
        timestamp = int(time.time())

        for campo, sufijo in archivos_map.items():
            if campo in request.FILES:
                archivo = request.FILES[campo]
                
                # Construcción del nombre con timestamp
                if sufijo.endswith('.pdf'):
                    base_sufijo = sufijo.replace('.pdf', '')
                    # Ej: MTY-100-5_1709823.pdf o MTY-100-factura_1709823.pdf
                    s3_path = f"logistica_ternium/{remision_num}/{remision_num}-{base_sufijo}_{timestamp}.pdf"
                
                # ---> ¡AQUÍ ESTÁ LA CORRECCIÓN QUE TE FALTABA! <---
                elif sufijo.endswith('.xml'):
                    base_sufijo = sufijo.replace('.xml', '')
                    # Ej: MTY-100-factura_1709823.xml
                    s3_path = f"logistica_ternium/{remision_num}/{remision_num}-{base_sufijo}_{timestamp}.xml"
                
                else:
                    _nombre_base, extension = os.path.splitext(archivo.name)
                    # Ej: MTY-100-1_1709823.jpg
                    s3_path = f"logistica_ternium/{remision_num}/{remision_num}-{sufijo}_{timestamp}{extension}"
                
                ruta_guardada = _subir_archivo_a_s3(archivo, s3_path)
                if ruta_guardada:
                    setattr(self.object, campo, ruta_guardada)

        self.object.save()
        form.save_m2m() 
        enviar_alerta_merma_logistica(self.object)
        messages.success(self.request, f"Registro {self.object.remision} creado correctamente.")
        return redirect(self.get_success_url())


@method_decorator(login_required, name='dispatch')
class RegistroLogisticoUpdateView(UpdateView):
    model = RegistroLogistico
    form_class = RegistroLogisticoForm
    template_name = 'ternium/formulario_logistica_ternium.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f"Editando Registro: {self.object.remision}"
        context['query_string'] = self.request.GET.urlencode()
        return context

    def get_success_url(self):
        url = reverse_lazy('lista_registros_logistica')
        query_string = self.request.GET.urlencode()
        if query_string:
            return f"{url}?{query_string}"
        return url

    def form_valid(self, form):
        registro_original = self.get_object()
        self.object = form.save(commit=False)
        remision_num = form.cleaned_data.get('remision', 'sin_remision').strip()
        request = self.request

        archivos_map = {
            'pdf_registro_camion_remision': '4.pdf', 
            'pdf_hoja_circulacion': '5.pdf',
            'foto_superior_vacia': '0', 
            'foto_frontal': '1',
            'foto_superior_llena': '2', 
            'foto_trasera': '3',
            'pdf_factura': 'factura.pdf', # <--- NUEVO
            'xml_factura': 'factura.xml', # <--- NUEVO
            'acuse_pdf': 'acuse.pdf'      # <--- NUEVO
        }

        # Timestamp para evitar caché en edición
        timestamp = int(time.time())

        for campo, sufijo in archivos_map.items():
            # 1. ¿Se solicitó eliminar? (Input hidden desde JS)
            clear_flag = request.POST.get(f'clear_{campo}') == 'true'
            
            # 2. ¿Se subió archivo nuevo?
            new_file_uploaded = campo in request.FILES

            ruta_antigua = getattr(registro_original, campo)
            
            # CASO A: LIMPIEZA O REEMPLAZO -> Borrar archivo viejo de S3
            if (clear_flag or new_file_uploaded) and ruta_antigua:
                if hasattr(ruta_antigua, 'name'):
                    _eliminar_archivo_de_s3(ruta_antigua.name)
                
                # Si es solo borrar (y no hay nuevo), limpiar campo en BD
                if clear_flag and not new_file_uploaded:
                    setattr(self.object, campo, None)

            # CASO B: SUBIR NUEVO ARCHIVO (Con nombre nuevo anti-caché)
            if new_file_uploaded:
                archivo = request.FILES[campo]
                
                if sufijo.endswith('.pdf'):
                    base_sufijo = sufijo.replace('.pdf', '')
                    s3_path = f"logistica_ternium/{remision_num}/{remision_num}-{base_sufijo}_{timestamp}.pdf"
                
                # ---> CORRECCIÓN APLICADA AQUÍ PARA EL XML <---
                elif sufijo.endswith('.xml'):
                    base_sufijo = sufijo.replace('.xml', '')
                    s3_path = f"logistica_ternium/{remision_num}/{remision_num}-{base_sufijo}_{timestamp}.xml"
                
                else:
                    _nombre_base, extension = os.path.splitext(archivo.name)
                    s3_path = f"logistica_ternium/{remision_num}/{remision_num}-{sufijo}_{timestamp}{extension}"

                ruta_guardada = _subir_archivo_a_s3(archivo, s3_path)
                if ruta_guardada:
                    setattr(self.object, campo, ruta_guardada)

        self.object.save()
        form.save_m2m()
        enviar_alerta_merma_logistica(self.object)
        messages.success(self.request, f"Registro {self.object.remision} actualizado correctamente.")
        return redirect(self.get_success_url())

@login_required
@require_POST
def auditar_registro_logistico(request, pk):
    registro = get_object_or_404(RegistroLogistico, pk=pk)
    if registro.status == 'TERMINADO':
        registro.status = 'AUDITADO'
        registro.auditado_por = request.user
        registro.auditado_en = timezone.now()
        registro.save(update_fields=['status', 'auditado_por', 'auditado_en'])
        messages.success(request, f'El registro {registro.remision} ha sido auditado.')
    else:
        messages.error(request, 'Este registro no cumple los requisitos para ser auditado.')
    return redirect('detalle_registro_logistica', pk=pk)


@method_decorator(login_required, name='dispatch')
class DescargarPaqueteZipView(View):
    """
    Vista CORREGIDA para descargar un ZIP con todos los archivos de un RegistroLogistico.
    Usa boto3 para descargar explícitamente cada archivo desde S3.
    """
    def get(self, request, *args, **kwargs):
        registro = get_object_or_404(RegistroLogistico, pk=self.kwargs.get('pk'))

        try:
            s3_client = boto3.client(
                's3',
                aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
                region_name=settings.AWS_S3_REGION_NAME
            )
        except (BotoCoreError, NoCredentialsError) as e:
            messages.error(request, f"Error de configuración con S3: {e}")
            return redirect('detalle_registro_logistica', pk=registro.pk)

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            file_fields = [
                'pdf_registro_camion_remision',  # Será -4.pdf
                'pdf_hoja_circulacion',          # Será -5.pdf (EL QUE FALTABA)
                'foto_superior_vacia',           # Será -0
                'foto_frontal',                  # Será -1
                'foto_superior_llena',           # Será -2
                'foto_trasera'                   # Será -3
            ]
            
            archivos_agregados = 0
            for field_name in file_fields:
                file_field = getattr(registro, field_name)
                if file_field and file_field.name:
                    s3_key = f"{settings.AWS_MEDIA_LOCATION}/{file_field.name}"
                    file_content_buffer = io.BytesIO()
                    
                    try:
                        s3_client.download_fileobj(settings.AWS_STORAGE_BUCKET_NAME, s3_key, file_content_buffer)
                        file_content_buffer.seek(0)
                        
                        filename_in_zip = os.path.basename(file_field.name)
                        zip_file.writestr(filename_in_zip, file_content_buffer.read())
                        archivos_agregados += 1
                    except s3_client.exceptions.ClientError as e:
                        if e.response['Error']['Code'] == '404':
                            print(f"Advertencia: El archivo {s3_key} no fue encontrado en S3 para el registro {registro.remision}.")
                        else:
                            messages.error(request, f"No se pudo descargar el archivo '{s3_key}' de S3.")
        
        if archivos_agregados == 0:
            messages.warning(request, "No se encontraron archivos en S3 para descargar en este registro.")
            return redirect('detalle_registro_logistica', pk=registro.pk)
            
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{registro.remision}.zip"'
        return response

# --- VISTAS DE EXPORTACIÓN Y API ---

@login_required
def export_logistica_to_excel(request):
    """
    Opción 3: Base de Datos completa.
    - Columnas reordenadas según solicitud.
    - Formato Tabla Excel (con filtros automáticos).
    - Lógica 3R para colores en porcentaje.
    """
    # 1. Configuración del Libro
    wb = Workbook()
    ws = wb.active
    ws.title = "Base de Datos Logística"
    
    # Estilos para Alerta 3R (Solo para la celda de porcentaje)
    red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid") 
    green_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid") 
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    link_font = Font(color="0563C1", underline="single") # <--- AÑADIDO: Estilo para links Azules
    
    # 2. Definir Encabezados (Orden Solicitado)
    headers = [
        'FECHA CARGA',                # 1
        'ESTATUS',                    # 2
        '# BOLETA BASCULA',           # 3
        'REMISION',                   # 4
        'TRANSPORTISTA',              # 5
        'PLACA TRACTOR',              # 6
        'PLACA TOLVA',                # 7
        'No. TRACTOR',                # 8
        'No. TOLVA',                  # 9
        'CHOFER',                     # 10
        'DESCRIPCION DEL MATERIAL',   # 11
        'CODIGO MATERIAL',            # 12
        'TONELADAS REMISIONADAS',     # 13
        'TONELADAS RECIBIDAS TERNIUM',# 14
        'MERMA (TON)',                # 15
        'PORCENTAJE (%)',             # 16
        'PERMISO CIRCULACION',        # 17
        'PAPELES',                    # 18
        'COMENTARIOS 3R',             # 19
        'FECHA ENTREGA A TERNIUM',    # 20
        'FECHA DE CORREO',            # 21  <--- NUEVO
        'PDF FACTURA',                # 22  <--- NUEVO
        'XML FACTURA',                # 23  <--- NUEVO
        'ACUSE (PDF)'                 # 24  <--- NUEVO
    ]
    
    ws.append(headers)
    
    # 3. Obtener Queryset (Respetando Filtros)
    queryset = RegistroLogistico.objects.select_related('transportista', 'material').order_by('-fecha_carga')

    # --- Aplicación de Filtros ---
    q = request.GET.get('q')
    transportista = request.GET.get('transportista')
    material_id = request.GET.get('material')
    status = request.GET.get('status')
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')
    merma = request.GET.get('merma') 

    if q:
        queryset = queryset.filter(Q(remision__icontains=q) | Q(boleta_bascula__icontains=q))
    if transportista:
        queryset = queryset.filter(transportista__nombre__icontains=transportista)
    if material_id:
        queryset = queryset.filter(material__id=material_id)
    if status:
        queryset = queryset.filter(status=status)
    if fecha_inicio:
        queryset = queryset.filter(fecha_carga__gte=fecha_inicio)
    if fecha_fin:
        queryset = queryset.filter(fecha_carga__lte=fecha_fin)
    
    if merma:
        queryset = queryset.annotate(
            abs_diff=Abs(F('toneladas_remisionadas') - F('toneladas_recibidas')),
            pct_calc=Case(
                When(toneladas_remisionadas__gt=0, then=F('abs_diff') / F('toneladas_remisionadas') * 100),
                default=Value(0.0),
                output_field=FloatField()
            )
        )
        if merma == 'SI':
            queryset = queryset.filter(pct_calc__gt=1.0)
        elif merma == 'NO':
            queryset = queryset.filter(pct_calc__lte=1.0)

    # 4. Escribir Datos
    for registro in queryset:
        
        # Lógica Papeles
        faltantes = []
        if not registro.pdf_registro_camion_remision: faltantes.append("Reg. Camión")
        if not registro.pdf_hoja_circulacion: faltantes.append("Hoja Circulación")
        papeles_status = "Ok" if not faltantes else "Falta: " + ", ".join(faltantes)

        # Cálculo Decimal Porcentaje
        rem = registro.toneladas_remisionadas or 0
        rec = registro.toneladas_recibidas or 0
        merma_decimal = 0
        
        # Calcular merma absoluta si no existe en el modelo (fallback)
        merma_abs = registro.merma_absoluta
        if merma_abs is None:
            merma_abs = abs(rem - rec)

        if rem > 0 and rec > 0:
            # Fórmula: (Recibido - Remisionado) / Remisionado
            # Si es negativo es pérdida, positivo es ganancia
            merma_decimal = float((rec - rem) / rem)

        # Mapeo de datos según el nuevo orden de columnas
        row = [
            registro.fecha_carga,                          # 1
            registro.get_status_display(),                 # 2
            registro.boleta_bascula,                       # 3
            registro.remision,                             # 4
            registro.transportista.nombre if registro.transportista else '', # 5
            registro.placas_tractor or '',                 # 6
            registro.placas_tolva or '',                   # 7
            registro.tractor or '',                        # 8
            registro.tolva or '',                          # 9
            registro.chofer or '',                         # 10
            registro.material.nombre if registro.material else '', # 11
            registro.material.id if registro.material else '',     # 12
            registro.toneladas_remisionadas,               # 13
            registro.toneladas_recibidas,                  # 14
            merma_abs,                                     # 15
            merma_decimal,                                 # 16
            registro.numero_permiso_sct or '',             # 17
            papeles_status,                                # 18
            registro.comentario or '',                     # 19
            registro.fecha_envio,                          # 20
            
            # --- NUEVOS CAMPOS ---
            registro.fecha_correo or '',                   # 21
            "Ver Archivo" if registro.pdf_factura else "Sin evidencia", # 22
            "Ver Archivo" if registro.xml_factura else "Sin evidencia", # 23
            "Ver Archivo" if registro.acuse_pdf else "Sin evidencia"    # 24
        ]
        
        ws.append(row)
        
        # Aplicar alineación y Formato Condicional a la fila actual
        current_row = ws.max_row
        is_terminado = (registro.status == 'TERMINADO') 

        for col_idx, cell in enumerate(ws[current_row], start=1):
            cell.alignment = center_align
            
            # Columna 16: PORCENTAJE (%)
            if col_idx == 16:
                cell.number_format = '0.00%' 
                try:
                    val = float(cell.value)
                    if is_terminado:
                        if val < 0:
                            cell.fill = red_fill
                        elif val > 0:
                            cell.fill = green_fill
                except:
                    pass
            
            # Columna 21: FECHA DE CORREO
            if col_idx == 21 and cell.value:
                cell.number_format = 'dd/mm/yyyy'
                
            # --- LÓGICA DE HYPERLINKS PARA LOS ARCHIVOS ---
            if col_idx == 22 and registro.pdf_factura:
                cell.hyperlink = request.build_absolute_uri(registro.pdf_factura.url)
                cell.font = link_font
                
            if col_idx == 23 and registro.xml_factura:
                cell.hyperlink = request.build_absolute_uri(registro.xml_factura.url)
                cell.font = link_font
                
            if col_idx == 24 and registro.acuse_pdf:
                cell.hyperlink = request.build_absolute_uri(registro.acuse_pdf.url)
                cell.font = link_font

    # 5. CREAR TABLA DE EXCEL (Esto agrega filtros y diseño azul)
    last_col_letter = get_column_letter(len(headers))
    last_row = ws.max_row
    
    if last_row > 1: 
        ref = f"A1:{last_col_letter}{last_row}"
        tabla = Table(displayName="TablaLogistica", ref=ref)
        
        style = TableStyleInfo(
            name="TableStyleMedium9", 
            showFirstColumn=False, 
            showLastColumn=False, 
            showRowStripes=True, 
            showColumnStripes=False
        )
        tabla.tableStyleInfo = style
        ws.add_table(tabla)

    # 6. Ajustar Ancho de Columnas
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    length = len(str(cell.value))
                    if length > max_length: max_length = length
            except: pass
        adjusted_width = min(max_length + 4, 50) # Tope máximo de 50
        ws.column_dimensions[column].width = adjusted_width

    # 7. Retornar Archivo
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f'BaseDatos_Logistica_{datetime.date.today()}.xlsx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    wb.save(response)
    return response

def get_catalogos_por_empresa(request, empresa_id):
    """
    Devuelve los catálogos filtrados por una empresa específica.
    ESTA ES LA VERSIÓN CORREGIDA.
    """
    try:
        # La línea 'unidades' fallaba porque 'models.F' no estaba definido.
        # Al importar 'F' directamente, ahora funciona correctamente.
        data = {
            'operadores': list(Operador.objects.filter(empresas__id=empresa_id).distinct().values('id', 'nombre')),
            'lineas_transporte': list(LineaTransporte.objects.filter(empresas__id=empresa_id).values('id', 'nombre')),
            'materiales': list(Material.objects.filter(empresas__id=empresa_id).values('id', 'nombre')),
            'unidades': list(Unidad.objects.filter(empresas__id=empresa_id).values('id', nombre=F('internal_id'), placas=F('license_plate'))),
            'contenedores': list(Contenedor.objects.filter(empresas__id=empresa_id).values('id', 'nombre', 'placas')),
            'lugares_origen': list(Lugar.objects.filter(empresas__id=empresa_id, tipo__in=['ORIGEN', 'AMBOS']).values('id', 'nombre')),
            'lugares_destino': list(Lugar.objects.filter(empresas__id=empresa_id, tipo__in=['DESTINO', 'AMBOS']).values('id', 'nombre')),
            'patios': list(Lugar.objects.filter(empresas__id=empresa_id, es_patio=True).values('id', 'nombre')),
        }
        return JsonResponse(data)
    except Exception as e:
        # En caso de otro error, devolvemos una respuesta vacía con un error 500
        # para que sea más fácil de depurar en el futuro.
        print(f"Error en get_catalogos_por_empresa: {e}")
        return JsonResponse({'error': 'Ocurrió un error en el servidor'}, status=500)


class CatalogoListView(ListView):
    """
    Clase base Premium. Incluye búsqueda avanzada, filtro dinámico por 'Empresa'
    y paginación.
    """
    template_name = 'ternium/catalogo_lista.html'
    paginate_by = 15

    def get_queryset(self):
        # Esta versión ya está corregida y no debería dar errores.
        queryset = super().get_queryset()
        query = self.request.GET.get('q')
        
        empresa_id = self.request.GET.get('empresa')

        if empresa_id and hasattr(self.model, 'empresas'):
            queryset = queryset.filter(empresas__id=empresa_id)

        if query:
            if hasattr(self.model, 'search_fields'):
                q_objects = Q()
                for field in self.model.search_fields:
                    q_objects |= Q(**{f'{field}__icontains': query})
                queryset = queryset.filter(q_objects).distinct()
            # Fallback seguro por si un modelo no define `search_fields`
            elif hasattr(self.model, 'internal_id'):
                queryset = queryset.filter(internal_id__icontains=query)
            elif hasattr(self.model, 'nombre'):
                 queryset = queryset.filter(nombre__icontains=query)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        model_name = self.model._meta.model_name
        
        context['verbose_name_plural'] = self.model._meta.verbose_name_plural
        context['verbose_name'] = self.model._meta.verbose_name
        context['model_name'] = model_name
        context['search_query'] = self.request.GET.get('q', '')
        
        # Check if the model has 'empresas' m2m field
        context['has_empresa_filter'] = hasattr(self.model, 'empresas') and model_name != 'empresa'
        if context['has_empresa_filter']:
            context['empresas'] = Empresa.objects.all().order_by('nombre')
            try:
                context['selected_empresa'] = int(self.request.GET.get('empresa', ''))
            except (ValueError, TypeError):
                context['selected_empresa'] = ''

        return context
    
# --- VISTAS DE CATÁLOGOS ---
class EmpresaListView(CatalogoListView): 
    model = Empresa
    template_name = 'ternium/empresa_list.html' # Template nuevo

class EmpresaCreateView(CreateView): 
    model = Empresa
    form_class = EmpresaForm
    template_name = 'ternium/empresa_form.html' # Template modificado
    success_url = reverse_lazy('lista_empresas')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Crear Nueva Empresa'
        return context

class EmpresaDetailView(DetailView): 
    model = Empresa
    template_name = 'ternium/empresa_detail.html' # Template nuevo

class EmpresaUpdateView(UpdateView): 
    model = Empresa
    form_class = EmpresaForm
    template_name = 'ternium/empresa_form.html' # Template modificado
    success_url = reverse_lazy('lista_empresas')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Editar Empresa: {self.object.nombre}'
        return context

class LineaTransporteListView(CatalogoListView): model = LineaTransporte
class LineaTransporteCreateView(CreateView): model = LineaTransporte; form_class = LineaTransporteForm; success_url = reverse_lazy('lista_lineas_transporte')
class LineaTransporteDetailView(DetailView): model = LineaTransporte
class LineaTransporteUpdateView(UpdateView): model = LineaTransporte; form_class = LineaTransporteForm; success_url = reverse_lazy('lista_lineas_transporte')

class OperadorListView(CatalogoListView): model = Operador
class OperadorCreateView(CreateView): model = Operador; form_class = OperadorForm; success_url = reverse_lazy('lista_operadores')
class OperadorDetailView(DetailView): model = Operador
class OperadorUpdateView(UpdateView): model = Operador; form_class = OperadorForm; success_url = reverse_lazy('lista_operadores')

class MaterialListView(CatalogoListView): model = Material
class MaterialCreateView(CreateView): model = Material; form_class = MaterialForm; success_url = reverse_lazy('lista_materiales')
class MaterialDetailView(DetailView): model = Material
class MaterialUpdateView(UpdateView): model = Material; form_class = MaterialForm; success_url = reverse_lazy('lista_materiales')


@method_decorator(login_required, name='dispatch')
class UnidadListView(CatalogoListView): 
    model = Unidad
    template_name = 'ternium/unidad_list.html' # Asegúrate que este template exista
class UnidadCreateView(CreateView): model = Unidad; form_class = UnidadForm; success_url = reverse_lazy('lista_unidades')
class UnidadDetailView(DetailView): model = Unidad
class UnidadUpdateView(UpdateView): model = Unidad; form_class = UnidadForm; success_url = reverse_lazy('lista_unidades')

class ContenedorListView(CatalogoListView): model = Contenedor
class ContenedorCreateView(CreateView): model = Contenedor; form_class = ContenedorForm; success_url = reverse_lazy('lista_contenedores')
class ContenedorDetailView(DetailView): model = Contenedor
class ContenedorUpdateView(UpdateView): model = Contenedor; form_class = ContenedorForm; success_url = reverse_lazy('lista_contenedores')

class LugarListView(CatalogoListView): 
    model = Lugar
    template_name = 'ternium/lugar_lista.html'
    context_object_name = 'lugares'
    paginate_by = 20

    def get_queryset(self):
        # 1. Queryset base
        queryset = Lugar.objects.prefetch_related('empresas').all()
        
        # 2. Filtro por Búsqueda (q)
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(nombre__icontains=query) | 
                Q(rfc__icontains=query) |
                Q(razon_social__icontains=query)
            )

        # 3. Filtro por Empresa
        empresa_id = self.request.GET.get('empresa')
        if empresa_id:
            queryset = queryset.filter(empresas__id=empresa_id)

        # 4. Filtro por Tipo
        tipo = self.request.GET.get('tipo')
        if tipo:
            queryset = queryset.filter(tipo=tipo)
        
        # --- AQUÍ ESTABA EL ERROR: FALTABA ESTE RETURN ---
        return queryset.order_by('nombre')

    # --- CORRECCIÓN DE INDENTACIÓN: Este método va al mismo nivel que get_queryset ---
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Cargar lista para el dropdown
        context['empresas_list'] = Empresa.objects.all().order_by('nombre')
        
        # Mantener seleccionado el filtro
        selected_empresa = self.request.GET.get('empresa')
        if selected_empresa:
            try:
                context['selected_empresa_id'] = int(selected_empresa)
            except ValueError:
                pass
            
        return context
            
        
class LugarCreateView(CreateView): model = Lugar; form_class = LugarForm; success_url = reverse_lazy('lista_lugares')
class LugarDetailView(DetailView): model = Lugar
class LugarUpdateView(UpdateView): model = Lugar; form_class = LugarForm; success_url = reverse_lazy('lista_lugares')

class DescargaListView(ListView):
    model = Descarga
    paginate_by = 15
    def get_queryset(self):
        return Descarga.objects.select_related('origen', 'destino', 'material', 'registrado_por').order_by('-fecha_descarga')

class DescargaCreateView(CreateView):
    model = Descarga
    form_class = DescargaForm
    success_url = reverse_lazy('descarga_lista')
    def form_valid(self, form):
        descarga = form.save(commit=False)
        descarga.registrado_por = self.request.user
        try:
            descarga.save()
            messages.success(self.request, "Descarga registrada y el inventario ha sido actualizado.")
            return redirect(self.success_url)
        except ValidationError as e:
            form.add_error(None, e)
            return self.form_invalid(form)
        
@login_required
def obtener_precio_medline(request):
    mes = request.GET.get('mes')
    if mes:
        precio_obj = PrecioMedline.objects.filter(mes=mes).first()
        if precio_obj:
            return JsonResponse({
                'success': True, 
                'precio_carton': float(precio_obj.precio_carton),
                'precio_archivo': float(precio_obj.precio_archivo)
            })
    return JsonResponse({'success': False})

@login_required
def export_remisiones_to_excel(request):
    # 1. Recuperar filtros del GET
    tipo_reporte = request.GET.get('tipo_reporte', 'normal') 
    
    # --- NUEVO: Capturamos el material seleccionado para Medline ---
    material_medline = request.GET.get('material_medline', '')
    # ---------------------------------------------------------------
    
    # Nuevos parámetros para Medline (Cartón y Archivo Muerto)
    precio_carton_req = request.GET.get('precio_carton', '0')
    precio_archivo_req = request.GET.get('precio_archivo', '0')
    mes_medline = request.GET.get('mes', '')
    
    q_remision = request.GET.get('q_remision', '')
    q_prefijo = request.GET.get('q_prefijo', '')
    q_empresa = request.GET.get('q_empresa', '') 
    q_material = request.GET.get('q_material', '')
    q_status = request.GET.get('q_status', '')
    q_origen = request.GET.get('q_origen', '')
    q_destino = request.GET.get('q_destino', '')
    q_operador = request.GET.get('q_operador', '') 
    q_fecha_desde = request.GET.get('q_fecha_desde', '')
    q_fecha_hasta = request.GET.get('q_fecha_hasta', '')
    q_folio_ld = request.GET.get('q_folio_ld', '')
    q_folio_dlv = request.GET.get('q_folio_dlv', '')

    # Convertir precios a decimal de forma segura
    try:
        precio_carton_val = float(precio_carton_req)
        precio_archivo_val = float(precio_archivo_req)
    except ValueError:
        precio_carton_val = 0.0
        precio_archivo_val = 0.0
    
    # 2. Construir QuerySet Base
    queryset = Remision.objects.all().select_related(
        'empresa', 'origen', 'destino', 'linea_transporte', 'operador', 'unidad', 'contenedor'
    ).prefetch_related('detalles', 'detalles__material', 'evidencias')
    
    # =========================================================================
    # SI ES MEDLINE: solo filtrar por mes y origen Medline, SIN filtros generales
    # =========================================================================
    if tipo_reporte == 'medline':
        queryset = queryset.filter(origen__nombre__icontains='MEDLINE')

        if mes_medline:
            try:
                year, month = mes_medline.split('-')
                queryset = queryset.filter(fecha__year=year, fecha__month=month)
            except ValueError:
                pass

            PrecioMedline.objects.update_or_create(
                mes=mes_medline,
                defaults={
                    'precio_carton': precio_carton_val,
                    'precio_archivo': precio_archivo_val
                }
            )

    # =========================================================================
    # TODOS LOS DEMÁS REPORTES: aplican filtros generales normalmente
    # =========================================================================
    else:
        if q_remision: queryset = queryset.filter(remision__icontains=q_remision)
        if q_prefijo: queryset = queryset.filter(remision__istartswith=q_prefijo)
        if q_empresa: queryset = queryset.filter(empresa_id=q_empresa)
        if q_material: queryset = queryset.filter(detalles__material__id=q_material).distinct()
        if q_status: queryset = queryset.filter(status=q_status)
        if q_origen: queryset = queryset.filter(origen__id=q_origen)
        if q_destino: queryset = queryset.filter(destino__id=q_destino)
        if q_operador: queryset = queryset.filter(operador__id=q_operador) 
        if q_fecha_desde: queryset = queryset.filter(fecha__gte=q_fecha_desde)
        if q_fecha_hasta: queryset = queryset.filter(fecha__lte=q_fecha_hasta)
        if q_folio_ld: queryset = queryset.filter(folio_ld__icontains=q_folio_ld)
        if q_folio_dlv: queryset = queryset.filter(folio_dlv__icontains=q_folio_dlv)

    queryset = queryset.order_by('fecha', 'pk')

    # 3. Crear Libro y Hoja
    wb = Workbook()
    ws = wb.active

    # --- ESTILOS COMUNES ---
    center_style = Alignment(horizontal="center", vertical="center", wrap_text=False)
    red_font = Font(color="FF0000", bold=True)      
    green_font = Font(color="009900", bold=True)    
    link_font = Font(color="0563C1", underline="single") 
    header_font = Font(bold=True)
    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    # =========================================================================
    # LÓGICA REPORTE MEDLINE — dos pestañas: Cartón y Archivo Muerto
    # =========================================================================
    if tipo_reporte == 'medline':
        precio_unitario = precio_carton_val or precio_archivo_val  # mismo precio para ambos

        pink_fill = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid")
        header_pink = PatternFill(start_color="FF66B2", end_color="FF66B2", fill_type="solid")
        yellow_header = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        green_header = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")

        def _fill_medline_sheet(sheet, keyword_include):
            """Rellena una hoja con las filas del material indicado (CARTON o ARCHIVO)."""
            headers = [
                "Remisión", "Origen", "Material", "Peso Carga (Kg)",
                "Folio Carga", "Inicia Carga", "PRECIO", "VENTA", "BULTOS", "FOLIO"
            ]
            sheet.append(headers)
            for col_idx, cell in enumerate(sheet[1], start=1):
                cell.alignment = center_style
                cell.font = Font(bold=True, color="FFFFFF" if col_idx <= 6 else "000000")
                if col_idx <= 6: cell.fill = header_pink
                elif col_idx in [7, 8]: cell.fill = yellow_header
                elif col_idx in [9, 10]: cell.fill = green_header

            total_peso = 0
            total_venta = 0

            for remision in queryset:
                val_inicia = timezone.localtime(remision.inicia_ld).date() if remision.inicia_ld else remision.fecha
                fecha_str = val_inicia.strftime("%d/%m/%Y") if val_inicia else ""

                for d in remision.detalles.all():
                    mat_nom = d.material.nombre.upper() if d.material else "CARTON"
                    if keyword_include not in mat_nom:
                        continue

                    bultos = int(d.bultos or 0)
                    peso_ld = float(d.peso_ld or 0)
                    venta_calculada = peso_ld * precio_unitario
                    total_peso += peso_ld
                    total_venta += venta_calculada
                    folio_generado = remision.folio_medline or "N/A"

                    sheet.append([
                        remision.remision,
                        remision.origen.nombre if remision.origen else 'MEDLINE',
                        mat_nom,
                        peso_ld,
                        remision.folio_ld or '',
                        fecha_str,
                        precio_unitario,
                        venta_calculada,
                        bultos,
                        folio_generado,
                    ])
                    current_row = sheet.max_row
                    for col_idx in range(1, 11):
                        cell = sheet.cell(row=current_row, column=col_idx)
                        cell.alignment = center_style
                        if col_idx <= 6: cell.fill = pink_fill
                        if col_idx == 4: cell.number_format = '#,##0.000'
                        if col_idx in [7, 8]: cell.number_format = '"$"#,##0.00'

            sheet.append(["", "", "", total_peso, "", "", "", total_venta, "", ""])
            last_row = sheet.max_row
            c = sheet.cell(row=last_row, column=4)
            c.font = Font(bold=True); c.number_format = '#,##0.000'
            c = sheet.cell(row=last_row, column=8)
            c.font = Font(bold=True); c.number_format = '"$"#,##0.00'

        # Pestaña 1: Cartón
        ws.title = "Cartón"
        _fill_medline_sheet(ws, "CARTON")

        # Pestaña 2: Archivo Muerto
        ws2 = wb.create_sheet(title="Archivo Muerto")
        _fill_medline_sheet(ws2, "ARCHIVO")

    # =========================================================================
    # LÓGICA REPORTE CONCENTRADO
    # =========================================================================
    elif tipo_reporte == 'concentrado':
        ws.title = "Resumen Concentrado"
        headers = [
            "Remisión", "Fecha", "Estatus", "Empresa", "Origen", "Destino",
            "Transporte", "Operador", "Unidad", "Contenedor", "Materiales",
            "Desglose Carga (Kg)", "Total Carga (Kg)", "Folio Carga", "Inicia Carga",
            "Desglose Descarga (Kg)", "Total Descarga (Kg)", "Folio Descarga", "Termina Descarga",
            "Total Rechazo (Kg)", "Dif. Neta (Kg)", "% Merma Global",
            "Comentarios", "Links Evidencias"
        ]
        ws.append(headers)
        for cell in ws[1]: cell.alignment = center_style; cell.font = header_font

        for remision in queryset:
            is_pendiente = (remision.status == 'PENDIENTE')
            evidencias_urls = [request.build_absolute_uri(ev.archivo.url) for ev in remision.evidencias.all() if ev.archivo]
            texto_evidencias = "Descargar Evidencia" if len(evidencias_urls) == 1 else ("Múltiples Archivos" if len(evidencias_urls) > 1 else "Sin adjuntos")
            
            val_inicia = timezone.localtime(remision.inicia_ld).date() if remision.inicia_ld else ''
            val_termina = timezone.localtime(remision.termina_dlv).date() if remision.termina_dlv else ''

            detalles = remision.detalles.all()
            mat_nombres = " | ".join([d.material.nombre for d in detalles if d.material]) or "Sin Material"
            
            if detalles.count() > 1:
                desglose_ld = " | ".join([f"{float(d.peso_ld or 0):.2f}" for d in detalles])
                desglose_dlv = " | ".join([f"{float(d.peso_dlv or 0):.2f}" for d in detalles])
            else:
                desglose_ld = "N/A"
                desglose_dlv = "N/A"

            total_ld = sum([float(d.peso_ld or 0) for d in detalles])
            total_dlv = sum([float(d.peso_dlv or 0) for d in detalles])
            total_rechazo = sum([float(d.peso_rechazado or 0) for d in detalles])

            if is_pendiente:
                diff = 0; porcentaje_merma = 0
            else:
                diff = (total_dlv + total_rechazo) - total_ld
                porcentaje_merma = (diff / total_ld) if total_ld > 0 else 0

            rechazo_val = total_rechazo if total_rechazo > 0 else "N/A"

            row_data = [
                remision.remision, remision.fecha, remision.get_status_display(),
                remision.empresa.nombre if remision.empresa else '',
                remision.origen.nombre if remision.origen else '',
                remision.destino.nombre if remision.destino else '',
                remision.linea_transporte.nombre if remision.linea_transporte else '',
                remision.operador.nombre if remision.operador else (remision.operador_manual or ''),
                remision.unidad.internal_id if remision.unidad else (remision.unidad_manual or ''),
                remision.contenedor.nombre if remision.contenedor else (remision.contenedor_manual or ''),
                mat_nombres,
                desglose_ld, total_ld, remision.folio_ld or '', val_inicia,
                desglose_dlv, total_dlv, remision.folio_dlv or '', val_termina,
                rechazo_val, diff, porcentaje_merma,
                remision.comentario or '', texto_evidencias
            ]
            ws.append(row_data)
            current_row = ws.max_row

            for cell in ws[current_row]:
                cell.alignment = center_style
                if is_pendiente: cell.fill = yellow_fill

            for col_idx in [2, 15, 19]: ws.cell(row=current_row, column=col_idx).number_format = 'dd/mm/yyyy'
            for col_idx in [13, 17, 21]: ws.cell(row=current_row, column=col_idx).number_format = '#,##0.000'
            if rechazo_val != "N/A": ws.cell(row=current_row, column=20).number_format = '#,##0.000'

            if not is_pendiente:
                cell_diff = ws.cell(row=current_row, column=21)
                if diff < 0: cell_diff.font = red_font
                elif diff > 0: cell_diff.font = green_font
            
            ws.cell(row=current_row, column=22).number_format = '0.00%'

            if len(evidencias_urls) == 1:
                cell_link = ws.cell(row=current_row, column=24) 
                cell_link.hyperlink = evidencias_urls[0]; cell_link.font = link_font

    # =========================================================================
    # LÓGICA REPORTE A DETALLE
    # =========================================================================
    elif tipo_reporte == 'detallado':
        ws.title = "Remisiones a Detalle"
        headers = [
            "Remisión", "Fecha", "Estatus", "Empresa", "Origen", "Destino",
            "Transporte", "Operador", "Unidad", "Contenedor", "Material",
            "Peso Carga (Kg)", "Folio Carga", "Inicia Carga",
            "Peso Descarga (Kg)", "Folio Descarga", "Termina Descarga",
            "Rechazo", "Dif. (Kg)", "% Merma", 
            "Comentarios", "Links Evidencias"
        ]
        ws.append(headers)
        for cell in ws[1]: cell.alignment = center_style; cell.font = header_font

        for remision in queryset:
            is_pendiente = (remision.status == 'PENDIENTE')
            evidencias_urls = [request.build_absolute_uri(ev.archivo.url) for ev in remision.evidencias.all() if ev.archivo]
            texto_evidencias = "Descargar Evidencia" if len(evidencias_urls) == 1 else ("Múltiples Archivos" if len(evidencias_urls) > 1 else "Sin adjuntos")

            val_inicia = timezone.localtime(remision.inicia_ld).date() if remision.inicia_ld else ''
            val_termina = timezone.localtime(remision.termina_dlv).date() if remision.termina_dlv else ''

            detalles = remision.detalles.all()
            if not detalles:
                detalles_list = [{'material': 'Sin Material', 'peso_ld': 0.0, 'peso_dlv': 0.0, 'peso_rechazado': 0.0}]
            else:
                detalles_list = [{
                    'material': d.material.nombre if d.material else 'Sin Material',
                    'peso_ld': float(d.peso_ld or 0),
                    'peso_dlv': float(d.peso_dlv or 0),
                    'peso_rechazado': float(d.peso_rechazado or 0)
                } for d in detalles]

            for d in detalles_list:
                peso_ld = d['peso_ld']
                peso_dlv = d['peso_dlv']
                peso_rechazado = d['peso_rechazado']
                
                if is_pendiente:
                    diff = 0; porcentaje_merma = 0
                else:
                    diff = (peso_dlv + peso_rechazado) - peso_ld
                    porcentaje_merma = (diff / peso_ld) if peso_ld > 0 else 0
                
                rechazo_val = peso_rechazado if peso_rechazado > 0 else "N/A"

                row_data = [
                    remision.remision, remision.fecha, remision.get_status_display(),
                    remision.empresa.nombre if remision.empresa else '',
                    remision.origen.nombre if remision.origen else '',
                    remision.destino.nombre if remision.destino else '',
                    remision.linea_transporte.nombre if remision.linea_transporte else '',
                    remision.operador.nombre if remision.operador else (remision.operador_manual or ''),
                    remision.unidad.internal_id if remision.unidad else (remision.unidad_manual or ''),
                    remision.contenedor.nombre if remision.contenedor else (remision.contenedor_manual or ''),
                    d['material'], 
                    peso_ld, remision.folio_ld or '', val_inicia,
                    peso_dlv, remision.folio_dlv or '', val_termina,
                    rechazo_val, diff, porcentaje_merma,
                    remision.comentario or '', texto_evidencias
                ]
                
                ws.append(row_data)
                current_row = ws.max_row

                for cell in ws[current_row]:
                    cell.alignment = center_style
                    if is_pendiente: cell.fill = yellow_fill

                for col_idx in [2, 14, 17]: ws.cell(row=current_row, column=col_idx).number_format = 'dd/mm/yyyy'
                for col_idx in [12, 15, 19]: ws.cell(row=current_row, column=col_idx).number_format = '#,##0.000'
                if rechazo_val != "N/A": ws.cell(row=current_row, column=18).number_format = '#,##0.000'

                if not is_pendiente:
                    cell_diff = ws.cell(row=current_row, column=19)
                    if diff < 0: cell_diff.font = red_font
                    elif diff > 0: cell_diff.font = green_font
                
                ws.cell(row=current_row, column=20).number_format = '0.00%'

                if len(evidencias_urls) == 1:
                    cell_link = ws.cell(row=current_row, column=22) 
                    cell_link.hyperlink = evidencias_urls[0]; cell_link.font = link_font

    # =========================================================================
    # LÓGICA REPORTE NORMAL
    # =========================================================================
    else:
        ws.title = "Remisiones (Normal)"
        headers = [
            "Remisión", "Fecha", "Estatus", "Empresa", "Origen", "Destino",
            "Transporte", "Operador", "Unidad", "Contenedor", "Materiales",
            "Peso Carga (Kg)", "Folio Carga", "Inicia Carga",
            "Peso Descarga (Kg)", "Folio Descarga", "Termina Descarga",
            "Rechazo", "Dif. (Kg)", "% Merma", 
            "Comentarios", "Links Evidencias"
        ]
        ws.append(headers)
        for cell in ws[1]: cell.alignment = center_style; cell.font = header_font

        for remision in queryset:
            is_pendiente = (remision.status == 'PENDIENTE')
            evidencias_urls = [request.build_absolute_uri(ev.archivo.url) for ev in remision.evidencias.all() if ev.archivo]
            texto_evidencias = "Descargar Evidencia" if len(evidencias_urls) == 1 else ("Múltiples Archivos" if len(evidencias_urls) > 1 else "Sin adjuntos")

            val_inicia = timezone.localtime(remision.inicia_ld).date() if remision.inicia_ld else ''
            val_termina = timezone.localtime(remision.termina_dlv).date() if remision.termina_dlv else ''

            detalles = remision.detalles.all()
            materiales = ", ".join([d.material.nombre for d in detalles if d.material]) or "Sin Material"
            
            peso_ld = sum([float(d.peso_ld or 0) for d in detalles])
            peso_dlv = sum([float(d.peso_dlv or 0) for d in detalles])
            peso_rechazado = sum([float(d.peso_rechazado or 0) for d in detalles])

            if is_pendiente:
                diff = 0; porcentaje_merma = 0
            else:
                diff = (peso_dlv + peso_rechazado) - peso_ld
                porcentaje_merma = (diff / peso_ld) if peso_ld > 0 else 0
            
            rechazo_val = peso_rechazado if peso_rechazado > 0 else "N/A"

            row_data = [
                remision.remision, remision.fecha, remision.get_status_display(),
                remision.empresa.nombre if remision.empresa else '',
                remision.origen.nombre if remision.origen else '',
                remision.destino.nombre if remision.destino else '',
                remision.linea_transporte.nombre if remision.linea_transporte else '',
                remision.operador.nombre if remision.operador else (remision.operador_manual or ''),
                remision.unidad.internal_id if remision.unidad else (remision.unidad_manual or ''),
                remision.contenedor.nombre if remision.contenedor else (remision.contenedor_manual or ''),
                materiales, 
                peso_ld, remision.folio_ld or '', val_inicia,
                peso_dlv, remision.folio_dlv or '', val_termina,
                rechazo_val, diff, porcentaje_merma,
                remision.comentario or '', texto_evidencias
            ]
            
            ws.append(row_data)
            current_row = ws.max_row

            for cell in ws[current_row]:
                cell.alignment = center_style
                if is_pendiente: cell.fill = yellow_fill

            for col_idx in [2, 14, 17]: ws.cell(row=current_row, column=col_idx).number_format = 'dd/mm/yyyy'
            for col_idx in [12, 15, 19]: ws.cell(row=current_row, column=col_idx).number_format = '#,##0.000'
            if rechazo_val != "N/A": ws.cell(row=current_row, column=18).number_format = '#,##0.000'

            if not is_pendiente:
                cell_diff = ws.cell(row=current_row, column=19)
                if diff < 0: cell_diff.font = red_font
                elif diff > 0: cell_diff.font = green_font
            
            ws.cell(row=current_row, column=20).number_format = '0.00%'

            if len(evidencias_urls) == 1:
                cell_link = ws.cell(row=current_row, column=22) 
                cell_link.hyperlink = evidencias_urls[0]; cell_link.font = link_font

    if ws.max_row > 1 and tipo_reporte != 'medline':
        last_col_letter = get_column_letter(len(headers))
        full_range = f"A1:{last_col_letter}{ws.max_row}"
        tab = Table(displayName="TablaRemisiones", ref=full_range)
        style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
        tab.tableStyleInfo = style
        ws.add_table(tab)

    for col in ws.columns:
        max_length = 0; column_letter = col[0].column_letter
        for cell in col:
            try:
                if cell.value and len(str(cell.value)) > max_length: max_length = len(str(cell.value))
            except: pass
        ws.column_dimensions[column_letter].width = min(max_length + 3, 60)

    nombre_base = f"Remisiones_{tipo_reporte.capitalize()}"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{nombre_base}_{datetime.date.today()}.xlsx"'
    wb.save(response)
    return response

@login_required
def detalles_genericos(request, model_name, pk):
    model_map = {
        'empresa': Empresa, 'material': Material, 'unidad': Unidad,
        'contenedor': Contenedor, 'operador': Operador, 'lugar': Lugar,
        'lineatransporte': LineaTransporte,
    }
    model = model_map.get(model_name)
    if not model:
        return HttpResponse('Modelo no encontrado', status=404)
    
    objeto = get_object_or_404(model, pk=pk)
    
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return render(request, 'ternium/detalles_genericos.html', {
            'objeto': objeto, 'model_name': model_name,
            'verbose_name': model._meta.verbose_name
        })
    
    return redirect(f'/{model_name}/editar/{pk}/')

@login_required
def busqueda_avanzada(request):
    query = request.GET.get('q', '')
    filter_type = request.GET.get('filter_type', 'all')
    
    if not query:
        return JsonResponse({'results': []})
    
    models_to_search = {
        'empresa': Empresa, 'material': Material, 'unidad': Unidad,
        'contenedor': Contenedor, 'operador': Operador, 'lugar': Lugar,
        'lineatransporte': LineaTransporte,
    }
    
    results = []
    
    for model_name, model in models_to_search.items():
        queryset = model.objects.all()
        
        if filter_type == 'all':
            q_objects = Q(nombre__icontains=query)
            if model_name == 'empresa':
                q_objects |= Q(rfc__icontains=query) | Q(contacto_principal__icontains=query)
            elif model_name == 'operador':
                q_objects |= Q(licencia__icontains=query) | Q(telefono__icontains=query)
            elif model_name in ['unidad', 'contenedor']:
                q_objects |= Q(placas__icontains=query)
            queryset = queryset.filter(q_objects)
        elif filter_type == 'nombre':
            queryset = queryset.filter(nombre__icontains=query)
        elif filter_type == 'rfc' and hasattr(model, 'rfc'):
            queryset = queryset.filter(rfc__icontains=query)
        
        for obj in queryset[:5]:
            edit_url_name = f'editar_{model_name}'
            try:
                url = reverse(edit_url_name, args=[obj.id])
            except:
                url = '#' 

            results.append({
                'model': model_name,
                'model_verbose': model._meta.verbose_name,
                'id': obj.id,
                'nombre': obj.nombre,
                'detalles': getattr(obj, 'rfc', '') or getattr(obj, 'licencia', '') or getattr(obj, 'placas', '') or '',
                'url': url
            })
    
    return JsonResponse({'results': results})


import os
import re
import io
import logging
import pandas as pd
from .models import RegistroLogistico, EntradaMaquila, InventarioPatio # Asegúrate de importar tus modelos
import boto3
from botocore.exceptions import NoCredentialsError, BotoCoreError
import json
from django.http import JsonResponse
from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils import timezone

from langchain_deepseek import ChatDeepSeek
from langchain_community.utilities import SQLDatabase
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import PromptTemplate
from sqlalchemy import create_engine
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# Configura un logger para un mejor seguimiento de errores
logger = logging.getLogger(__name__)

# Función auxiliar para extraer SQL de la respuesta del LLM
def _extraer_sql(texto_respuesta_ia: str) -> str:
    """
    Extrae el código SQL de la respuesta de un LLM de forma más robusta.
    """
    match = re.search(r"```sql\n(.*?)\n```", texto_respuesta_ia, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    
    select_pos = texto_respuesta_ia.upper().find("SELECT")
    if select_pos != -1:
        # Limpia el texto que a veces el LLM añade después del SQL
        sql_text = texto_respuesta_ia[select_pos:]
        if ';' in sql_text:
            sql_text = sql_text.split(';')[0]
        return sql_text.strip()
        
    return texto_respuesta_ia.strip()


@login_required
@permission_required('ternium.acceso_ia', raise_exception=True)
@csrf_exempt
def asistente_ia(request):
    """
    Gestiona las solicitudes al Asistente de IA con CONTEXTO DE NEGOCIO ENRIQUECIDO
    para identificar Mermas, Cargas, Descargas, Lugares y Operaciones.
    """
    if request.method == 'POST':
        pregunta = request.POST.get('pregunta', '').strip()
        if not pregunta:
            return JsonResponse({'error': 'La pregunta no puede estar vacía.'}, status=400)

        try:
            # 1. Verificación de configuración
            api_key = os.environ.get('DEEPSEEK_API_KEY')
            db_url = os.environ.get('DATABASE_URL_READONLY')
            if not api_key or not db_url:
                error_msg = 'Error de configuración: Faltan claves API o DB URL.'
                logger.error(error_msg)
                raise ValueError(error_msg)

            # 2. Configuración de IA y Base de Datos
            db_engine = create_engine(db_url)
            db = SQLDatabase(engine=db_engine)
            llm = ChatDeepSeek(model="deepseek-chat", api_key=api_key, temperature=0)

            # 3. Router de Intención (Sin cambios)
            router_prompt = PromptTemplate.from_template(
                """Analiza la pregunta y clasifica en: 'saludo', 'excel', o 'consulta'.
                Pregunta: "{pregunta}" 
                Responde SOLO la categoría."""
            )
            router_chain = router_prompt | llm | StrOutputParser()
            intencion = router_chain.invoke({"pregunta": pregunta}).strip().lower()

            if 'saludo' in intencion:
                return JsonResponse({'respuesta': "¡Hola! Puedo analizar cargas, descargas, mermas, lugares y operaciones. ¿Qué necesitas saber?"})

            # ==============================================================================
            # 4. CONTEXTO DE NEGOCIO (EL CEREBRO DE LA OPERACIÓN)
            # ==============================================================================
            business_rules = """
            DICCIONARIO DE DATOS Y REGLAS DE NEGOCIO (ÚSALO OBLIGATORIAMENTE):
            
            1. **REMISIONES (Operación Interna / Clientes / Patios):**
               - Tablas: `ternium_remision` (Cabecera) JOIN `ternium_detalleremision` (Detalles).
               - **Material:** Está en `ternium_detalleremision` unido con `ternium_material`.
               - **Carga (Peso Origen):** Campo `ternium_detalleremision.peso_ld`.
               - **Descarga (Peso Destino):** Campo `ternium_detalleremision.peso_dlv`.
               - **Merma (Faltante):** Fórmula SQL: `(ternium_detalleremision.peso_ld - ternium_detalleremision.peso_dlv)`.
               - **Porcentaje Merma:** `CASE WHEN peso_ld > 0 THEN ((peso_ld - peso_dlv) / peso_ld) * 100 ELSE 0 END`.
               - **Lugares:** `origen_id` y `destino_id` se unen con `ternium_lugar`.

            2. **LOGÍSTICA (Salidas a Ternium / Ventas):**
               - Tabla: `ternium_registrologistico`.
               - **Carga (Remisionado):** Campo `toneladas_remisionadas`.
               - **Descarga (Recibido):** Campo `toneladas_recibidas`.
               - **Merma:** `(toneladas_remisionadas - toneladas_recibidas)`.
               - **Proveedor/Transportista:** Tabla `ternium_lineatransporte`.

            3. **ENTRADAS MAQUILA (Compras / Materia Prima):**
               - Tabla: `ternium_entradamaquila`.
               - **Carga (Peso Remisión):** Campo `peso_remision`.
               - **Descarga (Peso Neto 3R):** Campo `peso_neto`.
               - **Merma:** `(peso_remision - peso_neto)`.
            
            4. **LUGARES Y DESTINOS:**
               - Tabla central: `ternium_lugar`.
               - Para saber el nombre del origen/destino, haz JOIN con esta tabla usando `origen_id` o `destino_id`.

            NOTA: Si el usuario pregunta por "operaciones", busca en las 3 tablas principales (Remision, Logistica, Entradas) o pregunta a cuál se refiere, pero prioriza `ternium_remision`.
            """

            # 5. Generación de SQL con Contexto
            table_info = db.get_table_info()
            
            template_sql = """Eres un experto en PostgreSQL y Logística. Genera una consulta SELECT precisa.
            
            ESQUEMA DE BASE DE DATOS:
            {table_info}
            
            {business_rules}

            REGLAS TÉCNICAS:
            1. Genera SOLO el código SQL (sin markdown, sin explicaciones).
            2. Usa JOINs explícitos (LEFT JOIN o INNER JOIN) para obtener nombres de materiales, lugares y empresas.
            3. Si piden 'Mermas', calcula la diferencia y ordénalo de mayor a menor pérdida.
            4. Usa LIMIT 1000.
            5. Para fechas usa CURRENT_DATE.
            
            PREGUNTA DEL USUARIO: {question}
            SQL QUERY:"""
            
            prompt_sql = PromptTemplate(
                input_variables=["question", "table_info", "business_rules"], 
                template=template_sql
            )
            
            chain_sql = prompt_sql | llm | StrOutputParser()
            
            # Invocamos pasando las reglas de negocio
            respuesta_sql = chain_sql.invoke({
                "question": pregunta, 
                "table_info": table_info,
                "business_rules": business_rules
            })

            if 'INVALIDO' in respuesta_sql.upper():
                return JsonResponse({'respuesta': 'No pude entender la relación de datos. Intenta ser más específico (ej: "mermas en remisiones").'})

            consulta_sql = _extraer_sql(respuesta_sql)
            
            # Validación de seguridad básica
            if not consulta_sql.upper().startswith('SELECT'):
                return JsonResponse({'error': 'Consulta no permitida.'}, status=400)

            # 6. Ejecución (Igual que antes)
            with db_engine.connect() as connection:
                df = pd.read_sql(consulta_sql, connection)

            file_url = None
            tabla_html = ""

            # 7. Generación de Excel (Igual que antes, manteniendo tu lógica S3)
            if 'excel' in intencion and not df.empty:
                nombre_base = re.sub(r'[^a-zA-Z0-9]+', '_', pregunta.lower())[:30] or "data"
                timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
                file_name = f"reportes_ia/{nombre_base}_{timestamp}.xlsx"
                
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df_excel = df.copy()
                    # Eliminar zonas horarias para Excel
                    for col in df_excel.select_dtypes(include=['datetimetz']).columns:
                        df_excel[col] = df_excel[col].dt.tz_localize(None)
                    df_excel.to_excel(writer, index=False, sheet_name='Datos')
                    
                    # Estilos básicos
                    worksheet = writer.sheets['Datos']
                    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                    header_font = Font(color="FFFFFF", bold=True)
                    for col_num, value in enumerate(df_excel.columns.values):
                        cell = worksheet.cell(row=1, column=col_num + 1)
                        cell.fill = header_fill
                        cell.font = header_font
                
                excel_buffer.seek(0)
                
                # Subida a S3 (Tu lógica original)
                file_path_relative = _subir_archivo_a_s3(excel_buffer, file_name)
                
                if file_path_relative:
                    try:
                        s3_client = boto3.client(
                            's3',
                            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
                            region_name=settings.AWS_S3_REGION_NAME,
                            config=boto3.session.Config(signature_version='s3v4')
                        )
                        full_s3_key = f"{settings.AWS_MEDIA_LOCATION}/{file_path_relative}"
                        file_url = s3_client.generate_presigned_url(
                            'get_object',
                            Params={'Bucket': settings.AWS_STORAGE_BUCKET_NAME, 'Key': full_s3_key},
                            ExpiresIn=172800
                        )
                    except Exception as e:
                        logger.error(f"Error S3 URL: {e}")
                        file_url = default_storage.url(file_path_relative)

            # 8. Respuesta de Texto (Interpretación de datos)
            if not df.empty:
                # Prompt para interpretar los datos usando términos de negocio
                template_respuesta = """Actúa como un analista de logística.
                PREGUNTA: "{pregunta}"
                
                DATOS ENCONTRADOS ({total_registros} registros):
                {muestra_datos}
                
                INSTRUCCIONES:
                1. Si hay mermas, destaca la más alta.
                2. Menciona claramente Origen y Destino si aplica.
                3. Usa formato markdown simple (negritas).
                4. Si se generó Excel, avisa al usuario.
                """
                
                prompt_respuesta = PromptTemplate.from_template(template_respuesta)
                chain_respuesta = prompt_respuesta | llm | StrOutputParser()
                
                muestra_datos = df.head(5).to_string() # Damos un poco más de contexto
                respuesta_final = chain_respuesta.invoke({
                    "pregunta": pregunta, 
                    "total_registros": len(df),
                    "muestra_datos": muestra_datos
                })
                
                tabla_html = df.head(10).to_html(classes='table table-sm table-striped mt-3', index=False, border=0, escape=False)
                if len(df) > 10:
                    preview_text = f'Mostrando 10 de {len(df)} registros.'
                    if file_url:
                        preview_text += f' <a href="{file_url}" download class="text-decoration-none fw-bold">📥 Descargar reporte completo</a>'
                    tabla_html += f'<div class="mt-2 text-center text-muted"><small>{preview_text}</small></div>'
            else:
                respuesta_final = "No encontré registros que coincidan con esos criterios de búsqueda."

            # Ensamblaje final con botón de descarga
            if file_url:
                respuesta_final += f"""
                <div class="alert alert-success mt-3 shadow-sm">
                    <div class="d-flex justify-content-between align-items-center">
                        <div>
                            <i class="fas fa-file-excel fa-2x text-success me-2"></i>
                            <strong>Reporte Generado</strong>
                        </div>
                        <a href="{file_url}" download class="btn btn-sm btn-success">
                            <i class="fas fa-download"></i> Descargar Excel
                        </a>
                    </div>
                </div>
                """
            
            return JsonResponse({
                'respuesta': respuesta_final,
                'tabla_html': tabla_html,
                'url_excel': file_url
            })

        except Exception as e:
            logger.error(f"Error IA: {str(e)}", exc_info=True)
            return JsonResponse({'error': f"Error procesando la solicitud: {str(e)}"}, status=500)

    return render(request, 'ternium/asistente_ia.html')


@method_decorator(login_required, name='dispatch')
class UnidadListView(CatalogoListView): # <-- MODIFICADO: Hereda de CatalogoListView
    """
    Vista para listar TODOS los activos (Unidades).
    Incluye filtros avanzados.
    """
    model = Unidad
    template_name = 'ternium/unidad_list.html'
    context_object_name = 'unidades' # Mantenemos 'unidades' para el template
    paginate_by = 20

    def get_queryset(self):
        # 1. Empezamos con el queryset del padre (que ya filtra por 'q' y 'empresa')
        queryset = super().get_queryset().prefetch_related('empresas') # <-- Optimizado con prefetch
        
        # 2. Obtenemos los filtros adicionales
        asset_type = self.request.GET.get('asset_type')
        status = self.request.GET.get('status')

        # 3. Aplicamos los filtros adicionales
        if asset_type:
            queryset = queryset.filter(asset_type=asset_type)
        if status:
            queryset = queryset.filter(operational_status=status)
        
        return queryset.order_by('internal_id') # Ordenamos por ID

    def get_context_data(self, **kwargs):
        # 1. Obtenemos el contexto del padre (que ya incluye 'empresas', 'search_query', etc.)
        context = super().get_context_data(**kwargs)
        
        # 2. Renombramos 'object_list' a 'unidades' para que el template funcione
        context['unidades'] = context.get('object_list')
        
        # 3. Añadimos las opciones para los nuevos filtros
        context['asset_type_choices'] = Unidad.AssetType.choices
        context['status_choices'] = Unidad.OperationalStatus.choices
        
        # 4. Pasamos todos los filtros aplicados para mantener el estado del form
        context['filtros_aplicados'] = self.request.GET
        return context

@login_required
def editar_unidad(request, pk):
    unidad_original = get_object_or_404(Unidad, pk=pk)
    
    if request.method == 'POST':
        form = UnidadForm(request.POST, request.FILES, instance=unidad_original)
        if form.is_valid():
            unidad = form.save(commit=False)
            unidad_id_folder = form.cleaned_data.get('internal_id', 'sin_id').strip()

            # Lógica para actualizar foto
            if 'display_photo' in request.FILES:
                _eliminar_archivo_de_s3(unidad_original.display_photo.name if unidad_original.display_photo else None)
                archivo = request.FILES['display_photo']
                s3_path = f"activos_unidades/{unidad_id_folder}/foto_{archivo.name}"
                unidad.display_photo = _subir_archivo_a_s3(archivo, s3_path)
            
            # Lógica para actualizar documentos
            if 'unit_documents' in request.FILES:
                _eliminar_archivo_de_s3(unidad_original.unit_documents.name if unidad_original.unit_documents else None)
                archivo = request.FILES['unit_documents']
                s3_path = f"activos_unidades/{unidad_id_folder}/doc_{archivo.name}"
                unidad.unit_documents = _subir_archivo_a_s3(archivo, s3_path)

            unidad.save()
            form.save_m2m()
            messages.success(request, f'Activo "{unidad.internal_id}" actualizado.')
            return redirect('lista_unidades')
    else:
        form = UnidadForm(instance=unidad_original)

    context = {
        'form': form, 'object': unidad_original, 'titulo': f'Editar Activo: {unidad_original.internal_id}'
    }
    return render(request, 'ternium/unidad_form.html', context)

@method_decorator(login_required, name='dispatch')
class UnidadDetailView(DetailView): 
    model = Unidad
    template_name = 'ternium/unidad_detail.html' # Asegúrate que este template exista


from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.forms import PasswordChangeForm
from .models import Profile # Asegúrate de que Profile esté importado

@login_required
def vista_perfil(request):
    user = request.user
    
    # Inicializamos el formulario de contraseña (se usará en GET y POST)
    password_form = PasswordChangeForm(user)

    if request.method == 'POST':
        # --- CASO 1: ACTUALIZAR INFORMACIÓN DEL PERFIL ---
        if 'update_profile' in request.POST:
            try:
                # Actualizar datos del modelo User
                user.first_name = request.POST.get('first_name', user.first_name)
                user.last_name = request.POST.get('last_name', user.last_name)
                user.email = request.POST.get('email', user.email)
                user.save()

                # Actualizar datos del modelo Profile (ternium_profile)
                if hasattr(user, 'ternium_profile'):
                    profile = user.ternium_profile
                    profile.telefono = request.POST.get('telefono', profile.telefono)
                    profile.area = request.POST.get('area', profile.area)
                    profile.empresa = request.POST.get('empresa', profile.empresa)
                    
                    # Manejo de archivo (Avatar)
                    if 'avatar' in request.FILES:
                        profile.avatar = request.FILES['avatar']
                    
                    profile.save()
                
                messages.success(request, 'Tu información de perfil ha sido actualizada.')
                return redirect('perfil')
            except Exception as e:
                messages.error(request, f'Ocurrió un error al actualizar el perfil: {e}')

        # --- CASO 2: CAMBIAR CONTRASEÑA ---
        elif 'change_password' in request.POST:
            password_form = PasswordChangeForm(user, request.POST)
            if password_form.is_valid():
                user = password_form.save()
                # Importante: Mantener la sesión activa tras cambiar contraseña
                update_session_auth_hash(request, user) 
                messages.success(request, 'Tu contraseña ha sido actualizada correctamente.')
                return redirect('perfil')
            else:
                messages.error(request, 'Error al cambiar contraseña. Revisa los campos.')

    # =======================================================
    # LÓGICA PARA MOSTRAR PERMISOS (SOLICITADO)
    # =======================================================
    permisos_usuario = user.get_all_permissions()
    lista_permisos_legibles = []
    
    # Diccionario para traducir los códigos técnicos a texto amigable
    nombres_amigables = {
        # Permisos Ternium / Operaciones
        'ternium.acceso_dashboard_patio': 'Dashboard de Patios',
        'ternium.acceso_remisiones': 'Módulo de Remisiones',
        'ternium.acceso_ia': 'Asistente de IA',
        'ternium.acceso_catalogos': 'Catálogos Operativos',
        'ternium.acceso_reportes_kpi': 'Reportes y KPIs',
        'ternium.view_ternium_module': 'Logística General',
        
        # Permisos Compras
        'compras.acceso_compras': 'Gestión de Compras',
        'compras.aprobar_solicitudes': 'Aprobar Solicitudes',
        
        # Permisos CXP
        'cuentas_por_pagar.acceso_cxp': 'Cuentas por Pagar',
        'cuentas_por_pagar.autorizar_pagos': 'Autorizar Pagos',
    }

    for perm_code in permisos_usuario:
        if perm_code in nombres_amigables:
            lista_permisos_legibles.append(nombres_amigables[perm_code])
        # Opcional: Si quieres mostrar otros permisos estándar de Django, descomenta esto:
        # else:
        #     lista_permisos_legibles.append(perm_code.split('.')[1].replace('_', ' ').capitalize())

    lista_permisos_legibles.sort()

    context = {
        'password_form': password_form,
        'user_groups': user.groups.all(),
        'permisos_detallados': lista_permisos_legibles, # <--- Enviamos la lista al HTML
    }
    return render(request, 'ternium/perfil.html', context)


class EmpresaVincularOrigenesView(LoginRequiredMixin, UpdateView):
    model = Empresa
    form_class = EmpresaOrigenesForm
    template_name = 'ternium/empresa_vincular_origenes.html' # <-- Un template nuevo

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f"Vinculando Orígenes para: {self.object.nombre}"
        return context

    def get_success_url(self):
        # Redirige a donde tengas tu lista de empresas o lugares
        # (Ajusta 'lista_lugares' si tienes una lista de empresas)
        return reverse_lazy('lista_lugares')
    
# --- COLOCAR AL FINAL DE ternium/views.py ---

# Asegúrate de que estas importaciones estén presentes

@login_required
@permission_required('ternium.acceso_reportes_kpi', raise_exception=True)
def dashboard_analisis_view(request):
    
    # ==========================================
    # 1. FILTROS DE TIEMPO Y CONSULTAS BASE
    # ==========================================
    now = timezone.now()
    
    # Obtener fechas del GET
    fecha_inicio_str = request.GET.get('fecha_inicio')
    fecha_fin_str = request.GET.get('fecha_fin')

    # --- QUERYSETS BASE ---
    entradas_all = EntradaMaquila.objects.all()
    logistica_all = RegistroLogistico.objects.all()

    # --- QUERYSETS DE ESTADÍSTICA (Filtrados por Status Válido) ---
    entradas_stats = entradas_all.exclude(status='CANCELADO')
    logistica_stats = logistica_all.filter(status__in=['TERMINADO', 'AUDITADO'])

    # --- APLICACIÓN DEL FILTRO DE FECHAS ---
    
    if fecha_inicio_str:
        entradas_stats = entradas_stats.filter(fecha_ingreso__gte=fecha_inicio_str)
        logistica_stats = logistica_stats.filter(fecha_carga__gte=fecha_inicio_str)
        entradas_all = entradas_all.filter(fecha_ingreso__gte=fecha_inicio_str) 
        logistica_all = logistica_all.filter(fecha_carga__gte=fecha_inicio_str)

    if fecha_fin_str:
        # CORRECCIÓN AQUÍ: Al ser DateField, pasamos solo la fecha string (YYYY-MM-DD).
        # El filtro __lte incluye todo el día indicado.
        entradas_stats = entradas_stats.filter(fecha_ingreso__lte=fecha_fin_str)
        logistica_stats = logistica_stats.filter(fecha_carga__lte=fecha_fin_str)
        entradas_all = entradas_all.filter(fecha_ingreso__lte=fecha_fin_str)
        logistica_all = logistica_all.filter(fecha_carga__lte=fecha_fin_str)

    # Si NO hay filtros, por defecto mostramos estadísticas del AÑO actual (o lo que prefieras)
    if not fecha_inicio_str and not fecha_fin_str:
        start_of_year = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
        entradas_stats = entradas_stats.filter(fecha_ingreso__gte=start_of_year)
        logistica_stats = logistica_stats.filter(fecha_carga__gte=start_of_year)


    # ==========================================
    # 2. CÁLCULO DE KPIS DE VOLUMEN
    # ==========================================
    
    # Totales (Sumas)
    toneladas_entradas = float(entradas_stats.aggregate(Sum('peso_neto'))['peso_neto__sum'] or 0)
    toneladas_salidas = float(logistica_stats.aggregate(Sum('toneladas_remisionadas'))['toneladas_remisionadas__sum'] or 0)
    
    # Inventario Actual
    inventario_agg = InventarioPatio.objects.aggregate(total_kg=Coalesce(Sum('cantidad'), 0.0, output_field=FloatField()))
    inventario_actual_tons = float(inventario_agg['total_kg']) / 1000 if inventario_agg['total_kg'] else 0

    # --- CÁLCULO DE MERMAS GLOBALES EN EL PERIODO ---
    merma_maq_tons = float(entradas_stats.aggregate(
        merma_tons=Sum(F('peso_neto') - F('peso_remision'), output_field=FloatField())
    )['merma_tons'] or 0)
    
    merma_log_tons = float(logistica_stats.aggregate(
        merma_tons=Sum(F('toneladas_recibidas') - F('toneladas_remisionadas'), output_field=FloatField())
    )['merma_tons'] or 0)
    
    merma_total_global = merma_maq_tons + merma_log_tons
    total_manejado = toneladas_entradas + toneladas_salidas
    
    porcentaje_merma_global = (merma_total_global / total_manejado) * 100 if total_manejado > 0 else 0
    
    movimientos_totales = entradas_stats.count() + logistica_stats.count()

    # ==========================================
    # 3. DESGLOSE DE ESTATUS (RESUMEN)
    # ==========================================
    
    status_summary = {
        'PENDIENTE': {'ent': 0, 'sal': 0, 'icon': 'fa-clock', 'color': 'warning'},
        'TERMINADO': {'ent': 0, 'sal': 0, 'icon': 'fa-check', 'color': 'primary'},
        'AUDITADO':  {'ent': 0, 'sal': 0, 'icon': 'fa-lock', 'color': 'success'},
        'CANCELADO': {'ent': 0, 'sal': 0, 'icon': 'fa-ban', 'color': 'danger'},
        'RECHAZADO': {'ent': 0, 'sal': 0, 'icon': 'fa-times-circle', 'color': 'dark'},
    }

    # Contar Entradas
    for item in entradas_all.values('status').annotate(count=Count('id')):
        s = item['status']
        if s in status_summary:
            status_summary[s]['ent'] = item['count']

    # Contar Salidas
    for item in logistica_all.values('status').annotate(count=Count('id')):
        s = item['status']
        if s in status_summary:
            status_summary[s]['sal'] = item['count']

    # ==========================================
    # 4. GRÁFICAS PRINCIPALES
    # ==========================================

    # Gráfica 1: Volumen Mensual
    entradas_mensuales = entradas_stats.annotate(mes=TruncMonth('fecha_ingreso')).values('mes').annotate(toneladas=Sum('peso_neto')).order_by('mes')
    salidas_mensuales = logistica_stats.annotate(mes=TruncMonth('fecha_carga')).values('mes').annotate(toneladas=Sum('toneladas_remisionadas')).order_by('mes')
    
    timeline_data = {}
    for entry in entradas_mensuales:
        if entry['mes']:
            mes_str = entry['mes'].strftime("%Y-%m")
            timeline_data[mes_str] = {'entrada': float(entry['toneladas'] or 0), 'salida': 0.0}

    for entry in salidas_mensuales:
        if entry['mes']:
            mes_str = entry['mes'].strftime("%Y-%m")
            val_salida = float(entry['toneladas'] or 0)
            if mes_str in timeline_data:
                timeline_data[mes_str]['salida'] = val_salida
            else:
                timeline_data[mes_str] = {'entrada': 0.0, 'salida': val_salida}
    
    sorted_timeline_keys = sorted(timeline_data.keys())
    chart_timeline_labels = sorted_timeline_keys
    chart_timeline_entrada = [timeline_data[k]['entrada'] for k in sorted_timeline_keys]
    chart_timeline_salida = [timeline_data[k]['salida'] for k in sorted_timeline_keys]

    # Gráfica 2: Mermas Mensuales
    merma_mensual = entradas_stats.annotate(mes=TruncMonth('fecha_ingreso')).values('mes').annotate(avg_merma=Avg('porcentaje_faltante')).order_by('mes')
    chart_merma_labels = [item['mes'].strftime("%Y-%m") for item in merma_mensual if item['mes']]
    chart_merma_data = [float(item['avg_merma'] or 0) for item in merma_mensual if item['mes']]

    # ==========================================
    # 5. TABLAS DE ANÁLISIS
    # ==========================================
    
    # A) Análisis por Material
    materiales_entradas = entradas_stats.values('calidad').annotate(
        entrada=Sum('peso_neto'), 
        merma_count=Count('id', filter=Q(alerta=True))
    ).order_by('-entrada')
    
    materiales_salidas = logistica_stats.values('material__nombre').annotate(
        salida=Sum('toneladas_remisionadas')
    ).order_by('-salida')
    
    inventario_por_material = InventarioPatio.objects.values('material__nombre').annotate(stock=Sum('cantidad'))

    material_analysis = {}
    
    for item in materiales_entradas:
        calidad = item['calidad'] or "Sin Especificar"
        material_analysis[calidad] = {
            'entrada': float(item['entrada'] or 0), 
            'salida': 0.0, 
            'merma_incidencias': item['merma_count'], 
            'stock': 0.0
        }
    
    for item in materiales_salidas:
        name = item['material__nombre'] or "Sin Especificar"
        val_salida = float(item['salida'] or 0)
        if name in material_analysis: 
            material_analysis[name]['salida'] = val_salida
        else: 
            material_analysis[name] = {
                'entrada': 0.0, 
                'salida': val_salida, 
                'merma_incidencias': 0, 
                'stock': 0.0
            }
    
    for item in inventario_por_material:
         name = item['material__nombre'] or "Sin Especificar"
         stock_ton = float(item['stock']) / 1000 if item['stock'] else 0
         if name in material_analysis: 
             material_analysis[name]['stock'] = stock_ton
    
    sorted_material_analysis = sorted(
        material_analysis.items(), 
        key=lambda item: item[1]['entrada'] + item[1]['salida'], 
        reverse=True
    )[:10]

    # B) Rankings Transportes (Entradas)
    top_transportes_entradas = entradas_stats.values('transporte').annotate(
        toneladas=Sum('peso_neto'), 
        viajes=Count('id'),
        avg_merma_perc=Avg('porcentaje_faltante')
    ).order_by('-toneladas')[:5]

    # C) Rankings Transportes (Salidas)
    top_transportes_salidas = logistica_stats.values('transportista__nombre').annotate(
        toneladas=Sum('toneladas_remisionadas'),
        viajes=Count('id'),
        avg_merma_perc=Avg(
            ExpressionWrapper(
                (F('toneladas_recibidas') - F('toneladas_remisionadas')) / F('toneladas_remisionadas') * 100,
                output_field=FloatField()
            )
        )
    ).order_by('-toneladas')[:5]

    # ==========================================
    # 6. BITÁCORA RECIENTE (Respetando filtro de fecha)
    # ==========================================
    
    # Entradas
    raw_entradas = entradas_stats.order_by('-fecha_ingreso', '-creado_en')[:10]
    ultimas_entradas = []
    
    for e in raw_entradas:
        pct = None
        if e.peso_remision and e.peso_remision > 0:
            pct = ((e.peso_neto - e.peso_remision) / e.peso_remision) * 100
            
        ultimas_entradas.append({
            'id': e.id,
            'status': e.status,
            'folio': e.c_id_remito,
            'calidad_mat': e.calidad,
            'peso': e.peso_neto,
            'trans': e.transporte,
            'ganancia_pct': pct
        })
    
    # Salidas
    raw_salidas = logistica_stats.order_by('-fecha_carga', '-creado_en')[:10]
    ultimas_salidas = []
    
    for s in raw_salidas:
        pct = None
        if s.toneladas_remisionadas and s.toneladas_remisionadas > 0:
            recibido = s.toneladas_recibidas or 0
            pct = ((recibido - s.toneladas_remisionadas) / s.toneladas_remisionadas) * 100
            
        ultimas_salidas.append({
            'id': s.id,
            'status': s.status,
            'remision': s.remision,
            'material__nombre': s.material.nombre if s.material else "S/M",
            'toneladas_remisionadas': s.toneladas_remisionadas,
            'transportista__nombre': s.transportista.nombre if s.transportista else "S/T",
            'ganancia_pct': pct
        })
    
    context = {
        'kpi_ent_year': round(toneladas_entradas, 2),
        'kpi_sal_year': round(toneladas_salidas, 2),
        'kpi_inv_act': round(inventario_actual_tons, 2),
        'kpi_merma_perc': round(porcentaje_merma_global, 2),
        'kpi_merma_tons': round(merma_total_global, 2),
        'kpi_movs_total': movimientos_totales,
        
        'chart_tl_labels': json.dumps(chart_timeline_labels),
        'chart_tl_entrada': json.dumps(chart_timeline_entrada),
        'chart_tl_salida': json.dumps(chart_timeline_salida),
        'chart_merma_labels': json.dumps(chart_merma_labels),
        'chart_merma_data': json.dumps(chart_merma_data),
        
        'material_analysis': sorted_material_analysis,
        'top_transportes_entradas': top_transportes_entradas,
        'top_transportes_salidas': top_transportes_salidas,
        
        'ultimas_entradas': ultimas_entradas,
        'ultimas_salidas': ultimas_salidas,
        
        'status_summary': status_summary
    }

    return render(request, 'ternium/dashboard_analisis.html', context)



import pandas as pd
from datetime import timedelta
from django.shortcuts import render, redirect
from django.contrib import messages
from django.db import transaction
from django.db.models import Q
from django.utils import timezone
from django.contrib.auth.decorators import login_required

from .models import (
    Remision, Empresa, Operador, LineaTransporte, 
    Unidad, Contenedor, Lugar, Cliente, Material, DetalleRemision
)
from .forms import ImportarRemisionesForm

@login_required
def importar_remisiones_excel(request):
    lista_errores = [] 
    
    if request.method == 'POST':
        form = ImportarRemisionesForm(request.POST, request.FILES)
        
        if form.is_valid():
            archivo = request.FILES['archivo_excel']
            
            try:
                # 1. Leer el Excel
                df = pd.read_excel(archivo)
                df.columns = df.columns.str.strip().str.upper()
                
                conteo_creadas = 0
                conteo_actualizadas = 0

                # --- BUSCAR EMPRESA ---
                empresa_mty = Empresa.objects.filter(
                    Q(nombre__icontains="Monterrey") | Q(prefijo__icontains="MTY")
                ).first()
                if not empresa_mty:
                    empresa_mty = Empresa.objects.first()

                # --- FUNCIONES DE LIMPIEZA ---
                def limpiar_texto(valor, default="N/A"):
                    if pd.isna(valor) or str(valor).lower() in ['nan', 'nat', 'none', '']:
                        return default
                    return str(valor).strip()

                def limpiar_folio(valor):
                    # Si viene vacío regresa cadena vacía "" para poder validar después
                    val_str = str(valor).strip()
                    if val_str.lower() in ['nan', 'nat', 'none', '', '0', '0.0']: 
                        return '' 
                    if val_str.endswith('.0'): return val_str[:-2]
                    return val_str
                
                def obtener_datetime_ficticio(valor):
                    if pd.isna(valor) or str(valor).strip() == '': return None
                    try:
                        dt = pd.to_datetime(valor)
                        if dt.hour == 0 and dt.minute == 0:
                            dt = dt.replace(hour=12, minute=0)
                        return dt
                    except:
                        return None
                
                def obtener_fecha_simple(valor):
                    dt = obtener_datetime_ficticio(valor)
                    return dt.date() if dt else timezone.now().date()

                def obtener_float(valor):
                    try:
                        return float(valor) if pd.notnull(valor) else 0.0
                    except:
                        return 0.0

                # --- ITERAR FILAS ---
                for index, row in df.iterrows():
                    fila_excel = index + 2
                    try:
                        remision_num = limpiar_texto(row.get('REMISION'), default='')
                        if not remision_num:
                            continue 

                        # =======================================================
                        # 1. CATALOGOS
                        # =======================================================
                        nom_operador = limpiar_texto(row.get('OPERADOR'), 'Sin Operador')
                        operador_obj, _ = Operador.objects.get_or_create(nombre__iexact=nom_operador, defaults={'nombre': nom_operador})
                        if hasattr(operador_obj, 'empresas'): operador_obj.empresas.add(empresa_mty)

                        nom_linea = limpiar_texto(row.get('LINEA DE TRANSPORTE'), 'Terceros')
                        linea_obj, _ = LineaTransporte.objects.get_or_create(nombre__iexact=nom_linea, defaults={'nombre': nom_linea})
                        if hasattr(linea_obj, 'empresas'): linea_obj.empresas.add(empresa_mty)

                        eco_unidad = limpiar_texto(row.get('UNIDAD'), 'N/A')
                        unidad_obj, _ = Unidad.objects.get_or_create(internal_id__iexact=eco_unidad, defaults={'internal_id': eco_unidad})
                        if hasattr(unidad_obj, 'empresas'): unidad_obj.empresas.add(empresa_mty)

                        num_cont = limpiar_texto(row.get('CONT'), 'N/A')
                        placas_cont = limpiar_texto(row.get('PLACAS CONT'), '')
                        contenedor_obj, _ = Contenedor.objects.get_or_create(
                            nombre__iexact=num_cont, 
                            defaults={'nombre': num_cont, 'placas': placas_cont}
                        )

                        nom_origen = limpiar_texto(row.get('ORIGEN'), 'Origen Desconocido')
                        origen_obj, _ = Lugar.objects.get_or_create(nombre__iexact=nom_origen, defaults={'nombre': nom_origen, 'tipo': 'ORIGEN'})

                        nom_destino = limpiar_texto(row.get('DESTINO'), 'Destino Desconocido')
                        destino_obj, _ = Lugar.objects.get_or_create(nombre__iexact=nom_destino, defaults={'nombre': nom_destino, 'tipo': 'DESTINO'})

                        cliente_obj, _ = Cliente.objects.get_or_create(nombre__iexact=nom_destino, defaults={'nombre': nom_destino})
                        
                        nom_material = limpiar_texto(row.get('MATERIAL'), 'Generico')
                        material_obj, _ = Material.objects.get_or_create(nombre__iexact=nom_material, defaults={'nombre': nom_material})

                        # =======================================================
                        # 2. FECHAS Y FOLIOS
                        # =======================================================
                        
                        fecha_carga_base = obtener_datetime_ficticio(row.get('INICIA CARGA'))
                        if fecha_carga_base:
                            dt_inicia_carga = fecha_carga_base
                            dt_termina_carga = fecha_carga_base + timedelta(hours=1)
                        else:
                            dt_inicia_carga = None
                            dt_termina_carga = None
                        
                        folio_carga = limpiar_folio(row.get('FOLIO CARGA')) 

                        fecha_descarga_base = obtener_datetime_ficticio(row.get('INICIA DESCARGA'))
                        if fecha_descarga_base:
                            dt_inicia_descarga = fecha_descarga_base
                            dt_termina_descarga = fecha_descarga_base + timedelta(hours=1)
                        else:
                            dt_inicia_descarga = None
                            dt_termina_descarga = None

                        folio_descarga = limpiar_folio(row.get('FOLIO DESCARGA')) 

                        # =======================================================
                        # 3. LÓGICA DE ESTATUS (CORREGIDA)
                        # =======================================================
                        
                        # Verificamos si tiene los datos esenciales (texto del excel)
                        # Ignoramos fotos (evidencia), solo nos importan folios y fechas
                        tiene_datos_completos = (
                            folio_carga != '' and 
                            folio_descarga != '' and 
                            dt_inicia_carga is not None and 
                            dt_inicia_descarga is not None
                        )

                        if 'PTE' in nom_destino.upper():
                            # REGLA 1: Si es destino PTE -> PENDIENTE (sin importar nada más)
                            status_calculado = 'PENDIENTE'
                        elif tiene_datos_completos:
                            # REGLA 2: Si NO es PTE y tiene folios/fechas -> TERMINADO
                            # (Se marca terminado aunque no tenga fotos aún)
                            status_calculado = 'TERMINADO'
                        else:
                            # REGLA 3: Si le faltan folios o fechas -> PENDIENTE
                            status_calculado = 'PENDIENTE'

                        # =======================================================
                        # 4. GUARDAR
                        # =======================================================
                        with transaction.atomic():
                            remision_existente = Remision.objects.filter(remision=remision_num).first()
                            
                            datos_remision = {
                                'empresa': empresa_mty,
                                'fecha': obtener_fecha_simple(row.get('FECHA')),
                                'status': status_calculado,
                                'operador': operador_obj,
                                'linea_transporte': linea_obj,
                                'unidad': unidad_obj,
                                'contenedor': contenedor_obj,
                                'origen': origen_obj,
                                'destino': destino_obj,
                                'cliente': cliente_obj,
                                
                                'inicia_ld': dt_inicia_carga,
                                'termina_ld': dt_termina_carga,
                                'folio_ld': folio_carga,
                                
                                'inicia_dlv': dt_inicia_descarga,
                                'termina_dlv': dt_termina_descarga,
                                'folio_dlv': folio_descarga,
                                
                                'descripcion': f"Carga Excel {nom_material}",
                            }

                            if remision_existente:
                                for key, value in datos_remision.items():
                                    setattr(remision_existente, key, value)
                                remision_existente.save()
                                remision_obj = remision_existente
                                DetalleRemision.objects.filter(remision=remision_obj).delete()
                                conteo_actualizadas += 1
                            else:
                                datos_remision['remision'] = remision_num
                                remision_obj = Remision.objects.create(**datos_remision)
                                conteo_creadas += 1

                            DetalleRemision.objects.create(
                                remision=remision_obj,
                                material=material_obj,
                                cliente=destino_obj,
                                peso_ld=obtener_float(row.get('PESO CARGA')),
                                peso_dlv=obtener_float(row.get('PESO'))
                            )

                    except Exception as e:
                        print(f"Error Fila {fila_excel}: {e}")
                        lista_errores.append(f"Fila {fila_excel} ({remision_num}): {str(e)}")

                if conteo_creadas > 0 or conteo_actualizadas > 0:
                    messages.success(request, f"✅ Éxito: {conteo_creadas} nuevas, {conteo_actualizadas} actualizadas.")
                
                if lista_errores:
                    err_msg = " | ".join(lista_errores[:3])
                    messages.warning(request, f"⚠ Errores en {len(lista_errores)} filas. Ej: {err_msg}")
                    return render(request, 'ternium/importar_remisiones.html', {'form': form, 'lista_errores': lista_errores})
                
                return redirect('remision_lista')

            except Exception as e:
                messages.error(request, f"❌ Error Crítico: {e}")
                return render(request, 'ternium/importar_remisiones.html', {'form': form})
        else:
            messages.error(request, "Formulario inválido.")
    else:
        form = ImportarRemisionesForm()
        
    if 'form_zip' not in locals():
        form_zip = ImportarEvidenciasZipForm()
        
    context = {
        'form': form, # El form de Excel que ya tenías
        'form_zip': form_zip, # El nuevo form de ZIP
        # ... otros datos de contexto si tienes ...
    }
    # Asegúrate de pasar 'lista_errores' si existe en tu lógica original
    if 'lista_errores' in locals():
        context['lista_errores'] = lista_errores

    return render(request, 'ternium/importar_remisiones.html', context)


# --- AGREGAR AL FINAL DE ternium/views.py ---


from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import Sum, Count, FloatField, F, Q
from django.db.models.functions import TruncMonth, Coalesce
from django.utils import timezone
import json

from .models import Remision, Lugar

@login_required
def dashboard_remisiones_view(request):
    now = timezone.now()
    start_of_year = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)

    # 0. FILTROS DE URL
    filtro_origen_id = request.GET.get('origen')
    filtro_destino_id = request.GET.get('destino')

    # 1. BASE COMPLETA (INCLUYE PATIOS)
    # Esta base tiene TODO lo que está Terminado/Auditado y con Peso > 0
    qs_completa = Remision.objects.filter(
        status__in=['TERMINADO', 'AUDITADO'],
        detalles__peso_ld__gt=0,
        detalles__peso_dlv__gt=0
    )

    # Aplicamos filtros de usuario a la base completa
    if filtro_origen_id:
        qs_completa = qs_completa.filter(origen_id=filtro_origen_id)
    if filtro_destino_id:
        qs_completa = qs_completa.filter(destino_id=filtro_destino_id)

    # =======================================================
    # LÓGICA DE NO DUPLICIDAD (CLEAN DATA)
    # =======================================================
    # Regla: 
    # - Si va a PATIO (Destino='Patio') -> NO CONTAR (Es entrada/stock)
    # - Si sale de PATIO o va directo -> CONTAR (Es salida/venta)
    #
    # Esto se usará para: Gráficas, Materiales y KPIs para que todo cuadre.
    qs_sin_duplicados = qs_completa.exclude(destino__nombre__icontains='patio')

    # =======================================================
    # 3. GRÁFICA EVOLUCIÓN
    # =======================================================
    timeline_qs = qs_sin_duplicados.annotate(mes=TruncMonth('fecha')).values('mes').annotate(
        carga=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        descarga=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField())
    ).order_by('mes')

    chart_labels = []
    chart_carga = []
    chart_descarga = []
    for entry in timeline_qs:
        if entry['mes']:
            chart_labels.append(entry['mes'].strftime("%Y-%m"))
            chart_carga.append(entry['carga'])
            chart_descarga.append(entry['descarga'])

    # =======================================================
    # 4. ANÁLISIS MATERIALES (MODIFICADO: Usa qs_sin_duplicados)
    # =======================================================
    # Antes usaba qs_analisis, ahora usa la lógica estricta de no duplicar patios
    materiales_data = qs_sin_duplicados.values('detalles__material__nombre').annotate(
        total_carga=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        total_descarga=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField())
    ).order_by('-total_carga')

    mat_labels = []
    mat_carga_data = []
    mat_descarga_data = []
    
    lista_materiales = []
    # Evitamos división por cero si no hay datos
    total_volumen = materiales_data[0]['total_carga'] if (materiales_data and materiales_data[0]['total_carga'] > 0) else 1

    for mat in materiales_data:
        c = mat['total_carga']
        d = mat['total_descarga']
        n = mat['detalles__material__nombre']
        diff = c - d  # <--- CÁLCULO DE DIFERENCIA
        
        mat_labels.append(n)
        mat_carga_data.append(c)
        mat_descarga_data.append(d)
        
        lista_materiales.append({
            'nombre': n,
            'carga': c,
            'descarga': d,
            'diff': diff,  # <--- AGREGADO AL DICCIONARIO
            'porcentaje_relativo': (c / total_volumen * 100)
        })

    # =======================================================
    # 5. DATA PARA EL DESLIZADOR (OFFCANVAS) (MODIFICADO)
    # =======================================================
    # También debe usar qs_sin_duplicados para que coincida con la tabla
    raw_details = qs_sin_duplicados.values(
        'fecha', 'remision', 'origen__nombre', 'destino__nombre', 
        'detalles__material__nombre', 'detalles__peso_ld', 'detalles__peso_dlv'
    ).order_by('-fecha')

    materiales_detalle_map = {}
    for item in raw_details:
        mat_nombre = item['detalles__material__nombre'] or "Sin Material"
        if mat_nombre not in materiales_detalle_map:
            materiales_detalle_map[mat_nombre] = []
        
        # Conversiones y cálculos
        p_carga = float(item['detalles__peso_ld'] or 0)
        p_descarga = float(item['detalles__peso_dlv'] or 0)
        faltante = p_carga - p_descarga
        
        # Calcular porcentaje de faltante (evitando división por cero)
        porcentaje = (faltante / p_carga * 100) if p_carga > 0 else 0.0
        
        materiales_detalle_map[mat_nombre].append({
            'fecha': item['fecha'].strftime("%d/%m/%Y"),
            'remision': item['remision'],
            'material': mat_nombre,
            'origen': item['origen__nombre'],
            'destino': item['destino__nombre'],
            'carga': p_carga,
            'descarga': p_descarga,
            'faltante': faltante,
            'porcentaje': porcentaje
        })

    # =======================================================
    # 6. RANKING OPERADORES (Mantiene qs_completa)
    # =======================================================
    # Aquí SIEMPRE usamos la completa porque al operador se le paga todo el movimiento
    operadores_data = qs_completa.values('operador__nombre').annotate(
        total_viajes=Count('id', distinct=True),
        total_cargado=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        
        # CÁLCULO DE MERMA ACUMULADA (Solo suma si Carga > Descarga)
        merma_acumulada=Sum(
            Case(
                When(
                    detalles__peso_ld__gt=F('detalles__peso_dlv'),
                    then=F('detalles__peso_ld') - F('detalles__peso_dlv')
                ),
                default=0.0,
                output_field=FloatField()
            )
        )
    ).order_by('-merma_acumulada')[:50] # Ordenamos por quién ha perdido más material

    ranking_operadores = []
    
    for op in operadores_data:
        cargado = op['total_cargado'] or 0
        merma_real = op['merma_acumulada'] or 0 # Esta es la suma pura de faltantes
        
        # Porcentaje de Riesgo: Cuánto material pierde del total que mueve
        riesgo_perc = ((merma_real / cargado) * 100) if cargado > 0 else 0
        
        # Obtener detalle de viajes CON FALTANTE para el deslizador
        # Filtramos merma > 0.02 (20kg) para limpiar tolerancias mínimas
        viajes_con_faltante = qs_completa.filter(
            operador__nombre=op['operador__nombre'],
            detalles__peso_ld__gt=F('detalles__peso_dlv') + 0.02 
        ).values(
            'fecha', 'remision', 'detalles__material__nombre',
            'origen__nombre', 'destino__nombre',
            'detalles__peso_ld', 'detalles__peso_dlv'
        ).order_by('-fecha')

        detalles_list = []
        for v in viajes_con_faltante:
            p_carga = float(v['detalles__peso_ld'] or 0)
            p_descarga = float(v['detalles__peso_dlv'] or 0)
            faltante = p_carga - p_descarga
            
            detalles_list.append({
                'fecha': v['fecha'].strftime("%d/%m/%Y"),
                'remision': v['remision'],
                'material': v['detalles__material__nombre'] or 'S/M',
                'origen': v['origen__nombre'],
                'destino': v['destino__nombre'],
                'carga': p_carga,
                'descarga': p_descarga,
                'faltante': round(faltante, 3)
            })

        ranking_operadores.append({
            'nombre': op['operador__nombre'] or "S/N", 
            'viajes': op['total_viajes'],
            'cargado': cargado,
            'diff': merma_real, # Ahora 'diff' representa solo pérdida
            'riesgo_perc': riesgo_perc,
            'json_detalles': json.dumps(detalles_list)
        })

    # =======================================================
    # 7. PENDIENTES
    # =======================================================
    pendientes_qs = Remision.objects.filter(status='PENDIENTE')
    if filtro_origen_id: pendientes_qs = pendientes_qs.filter(origen_id=filtro_origen_id)
    if filtro_destino_id: pendientes_qs = pendientes_qs.filter(destino_id=filtro_destino_id)

    resumen_pend = pendientes_qs.aggregate(
        total_carga=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        total_descarga=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField())
    )
    
    lista_pendientes = []
    hoy_date = now.date()
    for rem in pendientes_qs:
        try: det=rem.detalles.first(); p_ld=float(det.peso_ld); p_dlv=float(det.peso_dlv)
        except: p_ld=0; p_dlv=0
        dias = (hoy_date - rem.fecha).days
        if p_dlv == 0: e="En Tránsito"; c="bg-secondary"
        elif abs(p_ld-p_dlv)>0.1: e="Diferencia"; c="bg-warning text-dark"
        else: e="Por Cerrar"; c="bg-info"
        
        lista_pendientes.append({
            'remision': rem.remision, 'fecha': rem.fecha, 'dias': dias,
            'empresa': rem.empresa.nombre if rem.empresa else '', 
            'cliente': rem.cliente.nombre if rem.cliente else '',
            'unidad': rem.unidad.internal_id if rem.unidad else '', 
            'peso_origen': p_ld, 'peso_destino': p_dlv,
            'estado_txt': e, 'badge_cls': c
        })
    lista_pendientes.sort(key=lambda x: x['dias'], reverse=True)

    # =======================================================
    # 8. KPIS & CONTEXTO FINAL (Usamos qs_sin_duplicados)
    # =======================================================
    # Actualizamos los KPIs generales para que coincidan con la gráfica y tabla
    kpis_year = qs_sin_duplicados.filter(fecha__gte=start_of_year).aggregate(
        total_carga=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        total_descarga=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField())
    )

    origenes_list = Lugar.objects.filter(tipo__in=['ORIGEN', 'AMBOS']).order_by('nombre')
    destinos_list = Lugar.objects.filter(tipo__in=['DESTINO', 'AMBOS']).order_by('nombre')

    context = {
        'kpi_carga': round(kpis_year['total_carga'], 2), 
        'kpi_descarga': round(kpis_year['total_descarga'], 2),
        'kpi_merma_ton': round(kpis_year['total_carga'] - kpis_year['total_descarga'], 2),
        
        # Gráficas
        'chart_labels': json.dumps(chart_labels), 
        'chart_carga': json.dumps(chart_carga), 
        'chart_descarga': json.dumps(chart_descarga),
        
        # Materiales
        'mat_labels': json.dumps(mat_labels),
        'mat_carga_data': json.dumps(mat_carga_data),
        'mat_descarga_data': json.dumps(mat_descarga_data),
        'lista_materiales': lista_materiales,
        'materiales_detalle_json': json.dumps(materiales_detalle_map),

        # Filtros y Tablas
        'origenes_list': origenes_list, 'destinos_list': destinos_list,
        'filtro_origen_sel': int(filtro_origen_id) if filtro_origen_id else None,
        'filtro_destino_sel': int(filtro_destino_id) if filtro_destino_id else None,
        'ranking_operadores': ranking_operadores,
        'pendientes_carga_total': resumen_pend['total_carga'], 
        'pendientes_descarga_total': resumen_pend['total_descarga'],
        'lista_pendientes': lista_pendientes,
    }
    
    if context['kpi_carga'] > 0:
        context['kpi_merma_perc'] = round((context['kpi_merma_ton'] / context['kpi_carga']) * 100, 2)
    else:
        context['kpi_merma_perc'] = 0

    return render(request, 'ternium/dashboard_remisiones.html', context)


from django.contrib.auth import login
from django.contrib.auth.views import LoginView
from .forms import CustomLoginForm  # <--- CORREGIDO

class CustomLoginView(LoginView):
    form_class = CustomLoginForm    # <--- CORREGIDO
    template_name = 'registration/login.html'

    def form_valid(self, form):
        # Lógica de "Recordar sesión"
        remember_me = form.cleaned_data.get('remember_me')
        if not remember_me:
            self.request.session.set_expiry(0)
        else:
            self.request.session.modified = True
        
        return super().form_valid(form)
    
    
@login_required
@require_POST
@permission_required('ternium.change_remision', raise_exception=True)
def cancelar_remision(request, pk):
    """
    Cancela una remisión, revierte el movimiento de inventario 
    y evita que aparezca en el Dashboard.
    """
    remision = get_object_or_404(Remision, pk=pk)
    
    if remision.status == 'AUDITADO':
        messages.error(request, 'No se puede cancelar una remisión que ya fue auditada.')
        return redirect('remision_lista')
        
    if remision.status == 'CANCELADO':
        messages.warning(request, 'Esta remisión ya estaba cancelada.')
        return redirect('remision_lista')

    try:
        with transaction.atomic():
            # 1. Revertir inventario
            _update_inventory_from_remision(remision, revert=True)

            # 2. Guardar folio Medline antes de liberarlo
            folio_liberado = remision.folio_medline
            remision.folio_medline = None

            # 3. Actualizar estatus
            remision.status = 'CANCELADO'

            usuario_nombre = request.user.username
            fecha_str = timezone.now().strftime("%d/%m/%Y %H:%M")
            remision.descripcion += f" [CANCELADA por {usuario_nombre} el {fecha_str}]"

            remision.save()

            # 4. Renumerar folios Medline posteriores para cerrar el hueco
            api_views._renumerar_folios_medline(folio_liberado)

            # --- HISTORIAL: REGISTRAR CANCELACIÓN ---
            HistorialRemision.objects.create(
                remision=remision,
                usuario=request.user,
                cambio="ESTATUS CAMBIADO A: CANCELADO (Inventario revertido)"
            )

        messages.success(request, f'La remisión {remision.remision} ha sido CANCELADA.')
    except Exception as e:
        messages.error(request, f'Error al cancelar: {e}')

    return redirect('remision_lista')

# --- ESTA ES LA PARTE QUE CAUSA EL ERROR ---
@xframe_options_exempt  
@login_required
def detalle_remision(request, pk):
    remision = get_object_or_404(Remision, pk=pk)
    return render(request, 'ternium/detalle_remision.html', {'remision': remision})


from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.styles import Alignment, PatternFill, Font
from openpyxl.utils import get_column_letter
import datetime

@login_required
def export_catalogo_excel(request, model_name):
    """
    Exporta a Excel cualquier catálogo simple basándose en el nombre del modelo.
    Incluye la columna 'Empresas Asociadas' donde es aplicable.
    """
    # 1. Configuración de Modelos y Columnas
    config = {
        'empresa': {
            'model': Empresa,
            'headers': ['ID', 'Nombre', 'Prefijo', 'Contacto', 'Teléfono', 'Email'],
            'fields': ['id', 'nombre', 'prefijo', 'contacto_principal', 'telefono', 'email']
        },
        'lugar': {
            'model': Lugar,
            # MODIFICADO para incluir Empresas
            'headers': ['ID', 'Nombre', 'Tipo', 'Es Patio', 'RFC', 'Razón Social', 'Dirección Completa', 'Empresas Asociadas'],
            'fields': ['id', 'nombre', 'tipo', 'es_patio', 'rfc', 'razon_social', 'direccion_completa', 'empresas_str']
        },
        'lineatransporte': {
            'model': LineaTransporte,
            # Ya incluida, pero aseguramos el orden
            'headers': ['ID', 'Nombre', 'Empresas Asociadas'],
            'fields': ['id', 'nombre', 'empresas_str']
        },
        'operador': {
            'model': Operador,
            # MODIFICADO para incluir Empresas
            'headers': ['ID', 'Nombre', 'Licencia', 'Teléfono', 'Empresas Asociadas'],
            'fields': ['id', 'nombre', 'licencia', 'telefono', 'empresas_str']
        },
        'material': {
            'model': Material,
            # MODIFICADO para incluir Empresas
            'headers': ['ID', 'Nombre', 'Clave SAT', 'Unidad SAT', 'Empresas Asociadas'],
            'fields': ['id', 'nombre', 'clave_sat', 'clave_unidad_sat', 'empresas_str']
        },
        'unidad': {
            'model': Unidad,
            # MODIFICADO para incluir Empresas y campos específicos de unidad
            'headers': ['ID Interno', 'Placas', 'Marca/Modelo', 'Tipo', 'Estatus', 'Dueño', 'Empresas Asociadas'],
            'fields': ['internal_id', 'license_plate', 'make_model', 'asset_type', 'operational_status', 'ownership', 'empresas_str']
        },
        'contenedor': {
            'model': Contenedor,
            # MODIFICADO para incluir Empresas
            'headers': ['ID/Nombre', 'Placas', 'Empresas Asociadas'],
            'fields': ['nombre', 'placas', 'empresas_str']
        }
    }

    conf = config.get(model_name.lower())
    if not conf:
        messages.error(request, "Modelo no válido para exportación.")
        return redirect('home')

    # 2. Obtener Queryset Base
    model = conf['model']
    # Optimizamos la consulta si el modelo tiene la relación 'empresas'
    if hasattr(model, 'empresas'):
        queryset = model.objects.all().prefetch_related('empresas')
    else:
        queryset = model.objects.all()
    
    # --- FILTROS COMUNES (SIN CAMBIOS) ---
    empresa_id = request.GET.get('empresa')
    if empresa_id and hasattr(model, 'empresas'):
        queryset = queryset.filter(empresas__id=empresa_id)
    
    query = request.GET.get('q')
    if query:
        if hasattr(model, 'search_fields'):
            q_objects = Q()
            for field in model.search_fields:
                q_objects |= Q(**{f'{field}__icontains': query})
            queryset = queryset.filter(q_objects).distinct()
        elif hasattr(model, 'nombre'):
             queryset = queryset.filter(nombre__icontains=query)

    # --- FILTROS ESPECÍFICOS PARA UNIDAD (SIN CAMBIOS) ---
    if model_name.lower() == 'unidad':
        asset_type = request.GET.get('asset_type')
        status = request.GET.get('status')
        
        if asset_type:
            queryset = queryset.filter(asset_type=asset_type)
        if status:
            queryset = queryset.filter(operational_status=status)

    # 3. Crear Excel
    wb = Workbook()
    ws = wb.active
    ws.title = f"Catálogo {model._meta.verbose_name_plural}"
    
    # Escribir encabezados
    ws.append(conf['headers'])

    # Escribir datos
    for obj in queryset:
        row = []
        for field in conf['fields']:
            # LÓGICA DE EXTRACCIÓN MODIFICADA/AÑADIDA
            if field == 'empresas_str':
                # Si el objeto tiene relación ManyToMany a 'empresas', la mostramos
                if hasattr(obj, 'empresas'):
                    val = ", ".join([e.nombre for e in obj.empresas.all()])
                else:
                    val = ""
            elif field == 'direccion_completa' and hasattr(obj, 'direccion_completa'):
                val = obj.direccion_completa()
            else:
                # Obtener valor, manejar campos con choices (get_FOO_display)
                if hasattr(obj, f'get_{field}_display'):
                    val = getattr(obj, f'get_{field}_display')()
                else:
                    val = getattr(obj, field, '')
                
                if val is None: val = ""
            row.append(str(val))
        ws.append(row)

    # 4. Formato de Tabla (SIN CAMBIOS)
    last_col = get_column_letter(len(conf['headers']))
    last_row = ws.max_row
    if last_row > 1:
        tab = Table(displayName=f"Tabla{model_name}", ref=f"A1:{last_col}{last_row}")
        style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
        tab.tableStyleInfo = style
        ws.add_table(tab)
        
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except: pass
            # Aseguramos un ancho máximo razonable (ej. 70)
            adjusted_width = min(max_length + 2, 70) 
            ws.column_dimensions[column].width = adjusted_width

    # 5. Respuesta
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Catalogo_{model_name}_{datetime.date.today()}.xlsx"'
    wb.save(response)
    return response

@login_required
def importar_evidencias_masivas(request):
    """
    Procesa un archivo ZIP con evidencias (imágenes/PDFs).
    Asocia cada archivo a una Remisión basándose en el nombre del archivo.
    Ejemplo: 'MTY-500.jpg' se guarda en la remisión con folio 'MTY-500'.
    """
    if request.method == 'POST':
        form = ImportarEvidenciasZipForm(request.POST, request.FILES)
        if form.is_valid():
            archivo_zip = request.FILES['archivo_zip']
            
            conteo_exitos = 0
            conteo_errores = 0
            errores_detalle = []

            try:
                # Abrimos el ZIP en memoria
                with zipfile.ZipFile(archivo_zip, 'r') as zf:
                    # Iteramos sobre cada archivo dentro del ZIP
                    for filename in zf.namelist():
                        # Ignorar carpetas o archivos ocultos de Mac (__MACOSX)
                        if filename.endswith('/') or '__MACOSX' in filename or filename.startswith('.'):
                            continue
                        
                        # Obtener nombre base (folio) y extensión
                        # Ej: filename="MTY-1611.jpg" -> nombre_base="MTY-1611", ext=".jpg"
                        nombre_base, extension = os.path.splitext(os.path.basename(filename))
                        nombre_base = nombre_base.strip() # Limpiamos espacios
                        
                        # Buscar la remisión en la BD
                        remision = Remision.objects.filter(remision__iexact=nombre_base).first()
                        
                        if remision:
                            try:
                                # Leer el archivo del ZIP a un buffer de memoria
                                file_content = zf.read(filename)
                                file_io = io.BytesIO(file_content)
                                
                                # Definir ruta S3 (mismo estándar que al crear individualmente)
                                # remisiones/MTY-1611/evidencia_MTY-1611.jpg
                                s3_path = f"remisiones/{remision.remision}/evidencia_{nombre_base}{extension}"
                                
                                # Usamos tu función auxiliar existente para subir a S3
                                ruta_guardada = _subir_archivo_a_s3(file_io, s3_path)
                                
                                if ruta_guardada:
                                    remision.evidencia_documento = ruta_guardada
                                    remision.save()
                                    conteo_exitos += 1
                                else:
                                    conteo_errores += 1
                                    errores_detalle.append(f"{filename}: Error al subir a S3")
                                    
                            except Exception as e:
                                conteo_errores += 1
                                errores_detalle.append(f"{filename}: Error procesando archivo ({str(e)})")
                        else:
                            conteo_errores += 1
                            errores_detalle.append(f"{filename}: No se encontró la remisión '{nombre_base}'")

                # Mensajes finales
                if conteo_exitos > 0:
                    messages.success(request, f"✅ Se vincularon {conteo_exitos} evidencias exitosamente.")
                
                if conteo_errores > 0:
                    messages.warning(request, f"⚠ Hubo {conteo_errores} archivos no procesados.")
                    # Opcional: Mostrar los primeros 5 errores
                    for err in errores_detalle[:5]:
                        messages.warning(request, err)

            except zipfile.BadZipFile:
                messages.error(request, "El archivo subido no es un ZIP válido.")
            except Exception as e:
                messages.error(request, f"Error crítico procesando el ZIP: {e}")
                
            return redirect('importar_remisiones_excel') # Redirigimos a la misma página
    
    # Si intentan entrar por GET a esta URL específica, los mandamos al importador general
    return redirect('importar_remisiones_excel')

@login_required
@require_POST
def cancelar_entrada(request, pk):
    entrada = get_object_or_404(EntradaMaquila, pk=pk)
    
    if entrada.status == 'AUDITADO':
        messages.error(request, 'No se puede cancelar una entrada auditada.')
    elif entrada.status == 'CANCELADO':
        messages.warning(request, 'Esta entrada ya estaba cancelada.')
    else:
        entrada.status = 'CANCELADO'
        entrada.save(update_fields=['status']) 
        messages.success(request, f'La entrada {entrada.c_id_remito} ha sido CANCELADA.')
        
    return redirect('lista_entradas')

@login_required
@require_POST
def cancelar_registro_logistica(request, pk):
    registro = get_object_or_404(RegistroLogistico, pk=pk)
    
    if registro.status == 'AUDITADO':
        messages.error(request, 'No se puede cancelar un registro auditado.')
    elif registro.status == 'CANCELADO':
        messages.warning(request, 'Este registro ya estaba cancelado.')
    else:
        registro.status = 'CANCELADO'
        registro.save(update_fields=['status'])
        messages.success(request, f'El registro {registro.remision} ha sido CANCELADO.')
        
    return redirect('lista_registros_logistica')

def registrar_historial(remision, usuario, accion, descripcion=""):
    try:
        HistorialRemision.objects.create(
            remision=remision,
            usuario=usuario,
            accion=accion,
            descripcion=descripcion
        )
    except Exception as e:
        print(f"Error guardando historial: {e}")
        
@login_required
@require_POST
def eliminar_evidencia_individual(request, pk):
    evidencia = get_object_or_404(EvidenciaRemision, pk=pk)
    remision_id = evidencia.remision.pk
    
    # Seguridad
    if evidencia.remision.status == 'AUDITADO':
        messages.error(request, "No se puede eliminar evidencia de una remisión auditada.")
    else:
        # Borrar de S3
        if evidencia.archivo and hasattr(evidencia.archivo, 'name'):
            _eliminar_archivo_de_s3(evidencia.archivo.name)
        
        # Borrar de BD
        evidencia.delete()
        messages.success(request, "Archivo eliminado correctamente.")
        
    return redirect('editar_remision', pk=remision_id)


from .models import Plastico, EvidenciaPlastico, HistorialPlastico
from .forms import PlasticoForm

@login_required
def lista_plastico(request):
    # 1. Filtros Básicos
    q = request.GET.get('q', '')
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')
    status = request.GET.get('status')
    
    # --- NUEVO FILTRO: Capturar ID de empresa ---
    empresa_id = request.GET.get('empresa') 

    registros_list = Plastico.objects.all().order_by('-id')
    
    # 2. Aplicación de Filtros
    if q:
        registros_list = registros_list.filter(
            Q(remision__icontains=q) | 
            Q(descripcion__icontains=q) |
            Q(comentario__icontains=q) |
            Q(unidad__icontains=q)
        )
    
    # --- APLICAR FILTRO DE OPERACIÓN ---
    if empresa_id:
        registros_list = registros_list.filter(empresa_id=empresa_id)

    if fecha_inicio:
        registros_list = registros_list.filter(fecha__gte=fecha_inicio)
    if fecha_fin:
        registros_list = registros_list.filter(fecha__lte=fecha_fin)
    if status:
        registros_list = registros_list.filter(status=status)

    # 3. KPIs
    totales = registros_list.aggregate(
        total_peso=Sum('peso'),
        total_venta=Sum('venta'),
        conteo=Count('id')
    )

    # 4. Obtener lista de Empresas para el Select (Solo PLA/SEA o nombres relacionados)
    empresas_filtro = Empresa.objects.filter(
        Q(prefijo__in=['PLA', 'SEA']) | 
        Q(nombre__icontains='PLASTICO') | 
        Q(nombre__icontains='SEALED')
    ).order_by('nombre')

    # 5. Paginación
    paginator = Paginator(registros_list, 10) 
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'registros': page_obj, 
        'titulo': 'Gestión de Plástico',
        'kpi_peso': totales['total_peso'] or 0,
        'kpi_venta': totales['total_venta'] or 0,
        'kpi_conteo': totales['conteo'] or 0,
        # Pasamos los valores actuales para mantener el formulario lleno
        'q': q,
        'fecha_inicio': fecha_inicio,
        'fecha_fin': fecha_fin,
        'status_filter': status,
        'empresa_id_selected': int(empresa_id) if empresa_id else None, # ID seleccionado
        'empresas_list': empresas_filtro # Lista para el dropdown
    }

    return render(request, 'ternium/plastico_lista.html', context)
@login_required
def crear_plastico(request):
    if request.method == 'POST':
        form = PlasticoForm(request.POST, request.FILES)
        if form.is_valid():
            with transaction.atomic():
                plastico = form.save(commit=False)
                
                # NOTA: El folio (plastico.remision) se generará automáticamente 
                # dentro de plastico.save() gracias a la lógica que agregamos al modelo,
                # basándose en la empresa seleccionada en el formulario.
                
                # --- LÓGICA: PDF ADICIONAL A S3 ---
                if 'pdf_adicional' in request.FILES:
                    archivo_pdf = request.FILES['pdf_adicional']
                    # Usamos un nombre temporal si aún no tiene remision asignada (se asigna en save)
                    # Para asegurar el nombre correcto, guardamos primero el objeto para tener la remisión
                    # O usamos un timestamp si prefieres no guardar dos veces.
                    # Estrategia: Guardar primero para obtener folio.
                    pass 

                # 1. Guardar primero para generar el ID y Folio
                plastico.save() 

                # 2. Ahora que tenemos folio, procesamos archivos con el nombre correcto
                archivos_actualizados = False
                
                # Procesar PDF Adicional
                if 'pdf_adicional' in request.FILES:
                    archivo_pdf = request.FILES['pdf_adicional']
                    nombre_limpio = archivo_pdf.name.replace(" ", "_")
                    s3_path = f"plastico/adicionales/{plastico.remision}_{nombre_limpio}"
                    ruta_guardada = _subir_archivo_a_s3(archivo_pdf, s3_path)
                    if ruta_guardada:
                        plastico.pdf_adicional = ruta_guardada
                        archivos_actualizados = True

                if archivos_actualizados:
                    plastico.save() # Guardar referencias de archivos
                
                # 3. Historial
                HistorialPlastico.objects.create(
                    plastico=plastico, 
                    usuario=request.user, 
                    cambio=f"Creado con folio {plastico.remision} (Operación: {plastico.empresa})"
                )
                
                # 4. Evidencia Principal (Tabla separada)
                archivo = request.FILES.get('evidencia_documento')
                if archivo:
                    nombre_limpio = archivo.name.replace(" ", "_")
                    s3_path = f"plastico/{plastico.remision}_{nombre_limpio}"
                    ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                    if ruta_s3:
                        EvidenciaPlastico.objects.create(plastico=plastico, archivo=ruta_s3)

                messages.success(request, f'Registro {plastico.remision} guardado correctamente.')
                return redirect('lista_plastico')
    else:
        # GET: Formulario vacío. 
        # Ya NO pre-calculamos folio porque no sabemos qué empresa elegirá el usuario.
        form = PlasticoForm()

    return render(request, 'ternium/plastico_formulario.html', {
        'form': form, 
        'titulo': 'Nuevo Registro de Plástico / Sealed Air', 
        'is_editing': False
    })

@login_required
def editar_plastico(request, pk):
    plastico = get_object_or_404(Plastico, pk=pk)
    
    # Validación de permisos: Solo superusuario puede editar si está AUDITADO
    if plastico.status == 'AUDITADO' and not request.user.is_superuser:
        messages.error(request, 'No se puede editar un registro AUDITADO.')
        return redirect('detalle_plastico', pk=pk)

    if request.method == 'POST':
        form = PlasticoForm(request.POST, request.FILES, instance=plastico)
        if form.is_valid():
            with transaction.atomic():
                # 1. Instanciamos el objeto pero no guardamos en DB todavía
                plastico_obj = form.save(commit=False)

                # --- NUEVA LÓGICA: ACTUALIZAR PDF ADICIONAL ---
                if 'pdf_adicional' in request.FILES:
                    # A. Borrar archivo anterior de S3 si existe
                    if plastico.pdf_adicional and hasattr(plastico.pdf_adicional, 'name'):
                        _eliminar_archivo_de_s3(plastico.pdf_adicional.name)
                    
                    # B. Subir nuevo archivo
                    archivo_pdf = request.FILES['pdf_adicional']
                    nombre_limpio_pdf = archivo_pdf.name.replace(" ", "_")
                    # Usamos timestamp para evitar colisiones de nombre
                    s3_path_pdf = f"plastico/adicionales/{plastico.remision}_{timezone.now().timestamp()}_{nombre_limpio_pdf}"
                    
                    ruta_guardada_pdf = _subir_archivo_a_s3(archivo_pdf, s3_path_pdf)
                    
                    if ruta_guardada_pdf:
                        plastico_obj.pdf_adicional = ruta_guardada_pdf
                    else:
                         messages.warning(request, "Hubo un problema al subir el PDF adicional, pero se guardaron los demás datos.")
                # ---------------------------------------------
                
                # 2. Guardamos los cambios en el objeto Plastico (incluye 'pagado' y 'pdf_adicional')
                plastico_obj.save()
                
                # 3. Registrar Historial
                HistorialPlastico.objects.create(
                    plastico=plastico_obj, 
                    usuario=request.user, 
                    cambio="Registro Editado"
                )
                
                # 4. Procesar Evidencias Múltiples (Imágenes/PDFs principales)
                files = request.FILES.getlist('evidencia_documento')
                for i, archivo in enumerate(files):
                    nombre_limpio = archivo.name.replace(" ", "_")
                    s3_path = f"plastico/{plastico.id}_{timezone.now().timestamp()}_{nombre_limpio}"
                    ruta_s3 = _subir_archivo_a_s3(archivo, s3_path)
                    if ruta_s3:
                        EvidenciaPlastico.objects.create(plastico=plastico, archivo=ruta_s3)

                messages.success(request, 'Actualizado correctamente.')
                return redirect('lista_plastico')
    else:
        form = PlasticoForm(instance=plastico)

    return render(request, 'ternium/plastico_formulario.html', {
        'form': form, 
        'titulo': f'Editar {plastico.remision}', 
        'is_editing': True, 
        'plastico': plastico
    })

@login_required
def detalle_plastico(request, pk):
    plastico = get_object_or_404(Plastico, pk=pk)
    return render(request, 'ternium/plastico_detalle.html', {'plastico': plastico})

@login_required
def cancelar_plastico(request, pk):
    plastico = get_object_or_404(Plastico, pk=pk)
    if plastico.status != 'AUDITADO':
        plastico.status = 'CANCELADO'
        plastico.save()
        HistorialPlastico.objects.create(plastico=plastico, usuario=request.user, cambio="Registro CANCELADO")
        messages.warning(request, 'Registro Cancelado.')
    else:
        messages.error(request, 'No puedes cancelar un registro Auditado.')
    return redirect('lista_plastico')

@login_required
@permission_required('ternium.can_audit_plastico', raise_exception=True)
def auditar_plastico(request, pk):
    # ELIMINA EL BLOQUE 'if' DE AQUÍ.
    # El decorador ya hizo el trabajo de seguridad arriba.
    
    plastico = get_object_or_404(Plastico, pk=pk)
    
    if plastico.status == 'BORRADOR':
        plastico.status = 'AUDITADO'
        plastico.save()
        HistorialPlastico.objects.create(plastico=plastico, usuario=request.user, cambio="Registro AUDITADO")
        messages.success(request, 'Registro Auditado y Bloqueado.')
    else:
        messages.warning(request, f'No se puede auditar estatus {plastico.status}.')
        
    return redirect('lista_plastico')
@login_required
def exportar_plastico_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Reporte_Plastico.xlsx"'
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Plastico"
    ws.append(['ID', 'FECHA', 'REMISION', 'PESO', 'PRECIO', 'TOTAL', 'COMENTARIO'])
    
    registros = Plastico.objects.all().order_by('-fecha')
    
    # Filtros
    inicio = request.GET.get('fecha_inicio')
    fin = request.GET.get('fecha_fin')
    
    if inicio:
        registros = registros.filter(fecha__gte=inicio)
    if fin:
        registros = registros.filter(fecha__lte=fin)
        
    for p in registros:
        ws.append([p.id, p.fecha, p.remision, p.peso, p.precio, p.venta, p.comentario])
        
    wb.save(response)
    return response

@login_required
def export_remision_pdf(request, pk):
    remision = get_object_or_404(Remision, pk=pk)
    # Renderizamos una plantilla especial limpia para imprimir
    return render(request, 'ternium/remision_print.html', {'remision': remision})


@login_required
@require_POST
def crear_linea_transporte_ajax(request):
    """Crea una línea de transporte y la vincula a una empresa"""
    try:
        # 1. Decodificar datos
        data = json.loads(request.body)
        nombre = data.get('nombre', '').strip()
        empresa_id = data.get('empresa_id')
        
        # 2. Validaciones básicas
        if not nombre:
            return JsonResponse({'success': False, 'error': 'El nombre es obligatorio.'})
        
        if not empresa_id:
            return JsonResponse({'success': False, 'error': 'Debe seleccionar una operación (Empresa).'})

        # 3. Crear o Recuperar la Línea
        # Usamos get_or_create para evitar duplicados exactos
        nueva_linea, created = LineaTransporte.objects.get_or_create(
            nombre__iexact=nombre, 
            defaults={'nombre': nombre}
        )
        
        # 4. Vincular con la empresa seleccionada
        try:
            empresa = Empresa.objects.get(pk=empresa_id)
            nueva_linea.empresas.add(empresa)
        except Empresa.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'La empresa seleccionada no existe.'})
            
        return JsonResponse({
            'success': True, 
            'id': nueva_linea.id, 
            'nombre': nueva_linea.nombre,
            'mensaje': 'Línea creada y vinculada correctamente.'
        })
    except Exception as e:
        # Esto nos dirá si hay un error de código en el servidor
        print(f"Error AJAX: {str(e)}") 
        return JsonResponse({'success': False, 'error': f'Error interno: {str(e)}'})

@login_required
def listar_lineas_transporte_ajax(request):
    """Retorna la lista actualizada de transportes para refrescar el combo"""
    try:
        # Filtramos por empresa TERNIUM para mantener consistencia con tu lógica
        empresa_ternium = Empresa.objects.get(nombre__iexact="TERNIUM")
        lineas = LineaTransporte.objects.filter(empresas=empresa_ternium).order_by('nombre')
    except Empresa.DoesNotExist:
        # Si no existe TERNIUM, traemos todas
        lineas = LineaTransporte.objects.all().order_by('nombre')
        
    # Preparamos los datos para JSON
    data = [{'id': l.id, 'nombre': l.nombre} for l in lineas]
    return JsonResponse({'lineas': data})

from django.contrib.auth.decorators import user_passes_test
from django.db.models import Sum, Count, Avg
from django.db.models.functions import TruncMonth

# Función de comprobación de seguridad
def is_trane_user(user):
    return user.groups.filter(name='Trane_Usuario').exists() or user.is_superuser

from django.contrib.auth.decorators import user_passes_test
from django.db.models import Sum, Count, FloatField
from django.db.models.functions import Coalesce
import json

"""
=============================================================================
REEMPLAZO OPTIMIZADO DE dashboard_trane_view
=============================================================================
Buscar la función 'def dashboard_trane_view(request):' en views.py
y reemplazarla COMPLETA con esta versión.

CAMBIOS:
1. Fecha default: últimos 90 días si no hay filtros (antes: TODO el historial)
2. Cambio de icontains a iexact para usar índices
3. Fase 2: filtro de fecha aplicado también a salidas_patio_raw
4. ControlManifiestoTrane: limitado a 100 registros más recientes
5. select_related y prefetch_related verificados
"""

@login_required
@user_passes_test(is_trane_user)
def dashboard_trane_view(request):
    import datetime as dt
    import json
    from django.db.models import Sum, Count, FloatField
    from django.db.models.functions import Coalesce
    from django.utils import timezone
    from django.shortcuts import render

    # --- 0. PREPARACIÓN DE FILTROS ---
    materiales_disponibles = Material.objects.filter(
        detalleremision__remision__origen__nombre__iexact='TRANE'
    ).distinct().order_by('nombre')

    filtro_material_id = request.GET.get('material')
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')

    # =============================================
    # MODIFICACIÓN: Por defecto cargar solo últimos 30 días (1 mes)
    # Esto evita el colapso del servidor por exceso de datos históricos.
    # =============================================
    if not fecha_inicio and not fecha_fin and not filtro_material_id:
        fecha_fin = timezone.now().date().isoformat()
        fecha_inicio = (timezone.now().date() - dt.timedelta(days=30)).isoformat()

    # --- 1. CONSULTA PRINCIPAL ---
    queryset = Remision.objects.filter(
        origen__nombre__iexact='TRANE',
        status__in=['TERMINADO', 'AUDITADO']
    )

    # Aplicar filtros
    if filtro_material_id:
        queryset = queryset.filter(detalles__material_id=filtro_material_id)
    if fecha_inicio:
        queryset = queryset.filter(fecha__gte=fecha_inicio)
    if fecha_fin:
        queryset = queryset.filter(fecha__lte=fecha_fin)

    # --- 2. TOTALES GENERALES ---
    agregados = queryset.aggregate(
        total_kg=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField())
    )
    total_tons_val = agregados['total_kg'] / 1000
    total_tons_str = "{:,.2f}".format(total_tons_val)
    total_viajes = queryset.count()

    # --- 3. PREPARAR DESGLOSE Y BITÁCORA FASE 1 ---
    # OPTIMIZACIÓN: Se limita a los últimos 500 registros para proteger la memoria del servidor
    raw_movimientos = queryset.select_related(
        'origen', 'destino', 'linea_transporte', 'operador', 'unidad', 'contenedor'
    ).prefetch_related('detalles__material', 'evidencias').order_by('-fecha')[:500]
    
    breakdown_map = {}
    movimientos_formateados = []
    
    for mov in raw_movimientos:
        detalles_cached = mov.detalles.all()
        detalle = detalles_cached[0] if detalles_cached else None
        mat_nom = detalle.material.nombre if (detalle and detalle.material) else "Sin Clasificar"
        dest_nom = mov.destino.nombre if mov.destino else "Sin Destino"
        key = f"{mat_nom}_{dest_nom}"
        
        peso_ton = (mov.total_peso_ld or 0) / 1000
        peso_str = "{:,.3f}".format(peso_ton)
        
        transporte = mov.linea_transporte.nombre if mov.linea_transporte else 'S/T'
        operador = mov.operador.nombre if mov.operador else (mov.operador_manual or 'S/O')
        unidad = mov.unidad.internal_id if mov.unidad else (mov.unidad_manual or 'S/U')
        placas = getattr(mov.unidad, 'license_plate', getattr(mov.unidad, 'placas', None)) if mov.unidad else (mov.placas_unidad_manual or '')
        unidad_texto = f"{unidad} {f'({placas})' if placas else ''}".strip()
        
        folio_cliente = mov.folio_ld.strip() if mov.folio_ld and mov.folio_ld.strip() else mov.remision
        
        evidencias_extras = [{'url': e.archivo.url} for e in mov.evidencias.all() if e.archivo]
        archivos_desglose = []
        if mov.evidencia_documento: archivos_desglose.append({'url': mov.evidencia_documento.url, 'icon': 'fa-camera', 'color': 'text-danger', 'title': 'Evidencia Principal'})
        if mov.manifiesto: archivos_desglose.append({'url': mov.manifiesto.url, 'icon': 'fa-file-contract', 'color': 'text-primary', 'title': 'Manifiesto'})
        for extra in evidencias_extras: archivos_desglose.append({'url': extra['url'], 'icon': 'fa-paperclip', 'color': 'text-info', 'title': 'Evidencia Extra'})
            
        if key not in breakdown_map:
            breakdown_map[key] = []
            
        trazabilidad_str = mov.trazabilidad_notas if mov.trazabilidad_notas else ""
        
        breakdown_map[key].append({
            'remision': folio_cliente, 'fecha': mov.fecha.strftime("%d/%m/%Y"), 'peso': peso_str,
            'archivos': archivos_desglose, 'origen': mov.origen.nombre if mov.origen else 'N/A',
            'destino': dest_nom, 'operador': operador,
            'trazabilidad': trazabilidad_str
        })
        
        movimientos_formateados.append({
            'remision': folio_cliente, 'fecha': mov.fecha, 'status': mov.get_status_display(),
            'origen_nombre': mov.origen.nombre if mov.origen else 'N/A', 'destino_nombre': dest_nom,
            'transporte': transporte, 'operador': operador, 'unidad': unidad_texto,
            'material_nombre': mat_nom, 'peso_str': peso_str, 'evidencia_documento': mov.evidencia_documento,
            'manifiesto': mov.manifiesto, 'evidencias_extras': evidencias_extras,
            'trazabilidad_notas': trazabilidad_str
        })

    # --- 4. TABLA 1: TRAZABILIDAD FASE 1 ---
    trazabilidad = queryset.values(
        'detalles__material__nombre', 'destino__nombre', 'destino__es_patio'
    ).annotate(
        kilos=Sum('detalles__peso_ld'), viajes=Count('id')
    ).order_by('detalles__material__nombre')
    
    lista_trazabilidad = []
    for item in trazabilidad:
        mat_nom = item['detalles__material__nombre'] or "Sin Clasificar"
        destino_nombre = item['destino__nombre'] or "Sin Destino"
        tipo_destino = "PATIO DE ACOPIO" if item['destino__es_patio'] else "RECICLADOR FINAL"
        key = f"{mat_nom}_{destino_nombre}"
        lista_trazabilidad.append({
            'material': mat_nom, 'destino': destino_nombre, 'tipo_destino': tipo_destino,
            'toneladas': (item['kilos'] or 0) / 1000, 'viajes': item['viajes'], 'reciclaje': '100%',
            'json_breakdown': json.dumps(breakdown_map.get(key, []))
        })

    # --- 5. TABLA 2: TRAZABILIDAD FASE 2 ---
    envios_a_patio = queryset.filter(destino__es_patio=True)
    patios_ids = envios_a_patio.values_list('destino_id', flat=True).distinct()
    mat_ids = [filtro_material_id] if filtro_material_id else list(
        envios_a_patio.values_list('detalles__material_id', flat=True).distinct()
    )

    lista_patios_salidas = []
    
    if patios_ids:
        salidas_patio_raw = Remision.objects.filter(
            origen_id__in=patios_ids,
            detalles__material_id__in=mat_ids,
            status__in=['TERMINADO', 'AUDITADO']
        ).exclude(destino__es_patio=True).select_related(
            'origen', 'destino', 'operador'
        ).prefetch_related('detalles__material', 'evidencias')

        if fecha_inicio:
            salidas_patio_raw = salidas_patio_raw.filter(fecha__gte=fecha_inicio)
        if fecha_fin:
            salidas_patio_raw = salidas_patio_raw.filter(fecha__lte=fecha_fin)

        breakdown_fase2_map = {}
        for mov in salidas_patio_raw:
            det_first = mov.detalles.first()
            mat_nom = det_first.material.nombre if det_first and det_first.material else "S/C"
            orig_nom = mov.origen.nombre if mov.origen else "S/O"
            dest_nom = mov.destino.nombre if mov.destino else "S/D"
            key = f"{orig_nom}_{mat_nom}_{dest_nom}"
            
            folio_cliente_f2 = mov.folio_ld.strip() if mov.folio_ld and mov.folio_ld.strip() else mov.remision
            
            if key not in breakdown_fase2_map: breakdown_fase2_map[key] = []
                
            archivos_f2 = []
            if mov.evidencia_documento: archivos_f2.append({'url': mov.evidencia_documento.url, 'icon': 'fa-camera', 'color': 'text-danger', 'title': 'Evidencia'})
            if mov.manifiesto: archivos_f2.append({'url': mov.manifiesto.url, 'icon': 'fa-file-contract', 'color': 'text-primary', 'title': 'Manifiesto'})
            for ex in mov.evidencias.all(): 
                if ex.archivo: archivos_f2.append({'url': ex.archivo.url, 'icon': 'fa-paperclip', 'color': 'text-info', 'title': 'Extra'})
                
            trazabilidad_str_f2 = mov.trazabilidad_notas if mov.trazabilidad_notas else ""
                
            breakdown_fase2_map[key].append({
                'remision': folio_cliente_f2, 'fecha': mov.fecha.strftime("%d/%m/%Y"), 
                'peso': "{:,.3f}".format((mov.total_peso_ld or 0)/1000), 'archivos': archivos_f2,
                'origen': orig_nom, 'destino': dest_nom, 'operador': mov.operador.nombre if mov.operador else 'S/O',
                'trazabilidad': trazabilidad_str_f2
            })

        agrupado_patios = salidas_patio_raw.values(
            'origen__nombre', 'detalles__material__nombre', 'destino__nombre'
        ).annotate(
            kilos=Sum('detalles__peso_ld'), viajes=Count('id')
        ).order_by('detalles__material__nombre')

        for salida in agrupado_patios:
            key = f"{salida['origen__nombre']}_{salida['detalles__material__nombre']}_{salida['destino__nombre']}"
            lista_patios_salidas.append({
                'patio_origen': salida['origen__nombre'], 'material': salida['detalles__material__nombre'],
                'cliente_final': salida['destino__nombre'], 'toneladas': (salida['kilos'] or 0) / 1000,
                'reciclaje': '100%',
                'json_breakdown': json.dumps(breakdown_fase2_map.get(key, []))
            })

    destinos_unicos = sorted(list(set([item['destino'] for item in lista_trazabilidad])))
    materiales_unicos = sorted(list(set([item['material'] for item in lista_trazabilidad])))

    registros_salidas = ControlManifiestoTrane.objects.select_related(
        'linea_transporte', 'operador', 'destino', 'material'
    )
    if fecha_inicio:
        registros_salidas = registros_salidas.filter(fecha_captura__gte=fecha_inicio)
    if fecha_fin:
        registros_salidas = registros_salidas.filter(fecha_captura__lte=fecha_fin)
    registros_salidas = registros_salidas.order_by('-fecha_captura', '-id')[:100]

    context = {
        'total_tons': total_tons_str, 'total_viajes': total_viajes,
        'trazabilidad_list': lista_trazabilidad, 'trazabilidad_patios': lista_patios_salidas,
        'ultimos_movimientos': movimientos_formateados, 'materiales_filtro': materiales_disponibles,
        'filtro_actual': int(filtro_material_id) if filtro_material_id else None,
        'fecha_inicio': fecha_inicio,
        'fecha_fin': fecha_fin,
        'fecha_actual': timezone.now(),
        'destinos_unicos': destinos_unicos, 'materiales_unicos': materiales_unicos,
        'registros_salidas': registros_salidas
    }
    return render(request, 'ternium/dashboard_trane.html', context)

@login_required
@user_passes_test(is_trane_user)
def dashboard_trane_data(request):
    """
    API JSON para gráficas. 
    Convierte los datos sumados (Kg) a Toneladas (/1000) antes de enviar al Chart.js
    """
    queryset = Remision.objects.filter(
        origen__nombre__icontains='TRANE',
        status__in=['TERMINADO', 'AUDITADO']
    )
    
    filtro_material_id = request.GET.get('material')
    if filtro_material_id:
        queryset = queryset.filter(detalles__material_id=filtro_material_id)

    # GRÁFICA 1: Por Material (Volumen por Tipo de Material)
    # MODIFICADO: Sumando 'peso_ld' (Carga/Salida) en la anotación
    materiales = queryset.values('detalles__material__nombre')\
                         .annotate(total_kg=Sum('detalles__peso_ld'))\
                         .order_by('-total_kg')
    
    # GRÁFICA 2: Por Destino
    # MODIFICADO: Sumando 'peso_ld' (Carga/Salida) en la anotación
    destinos = queryset.values('destino__nombre')\
                       .annotate(total_kg=Sum('detalles__peso_ld'))\
                       .order_by('-total_kg')

    # GRÁFICA 3: Evolución Mensual
    # MODIFICADO: Sumando 'peso_ld' (Carga/Salida) en la anotación
    mensual = queryset.annotate(mes=TruncMonth('fecha'))\
                      .values('mes')\
                      .annotate(total_kg=Sum('detalles__peso_ld'))\
                      .order_by('mes')

    # Procesar Datos (Dividiendo entre 1000 y convirtiendo a float para evitar errores JSON)
    data_mat = {'labels': [], 'data': []}
    for item in materiales:
        data_mat['labels'].append(item['detalles__material__nombre'] or "N/A")
        data_mat['data'].append(float(round((item['total_kg'] or 0) / 1000, 2)))

    data_dest = {'labels': [], 'data': []}
    for item in destinos:
        data_dest['labels'].append(item['destino__nombre'] or "N/A")
        data_dest['data'].append(float(round((item['total_kg'] or 0) / 1000, 2)))

    data_timeline = {'labels': [], 'data': []}
    for item in mensual:
        if item['mes']:
            mes_str = item['mes'].strftime("%b %Y") 
            data_timeline['labels'].append(mes_str)
            data_timeline['data'].append(float(round((item['total_kg'] or 0) / 1000, 2)))

    return JsonResponse({
        'materiales': data_mat,
        'destinos': data_dest,
        'mensual': data_timeline
    })
    
def verificar_remito_existente(request):
    remito_id = request.GET.get('id', None)
    pk_actual = request.GET.get('pk_actual', None) # Para excluir el registro actual si estamos editando
    
    if remito_id:
        # Buscamos coincidencias exactas (case-insensitive)
        qs = EntradaMaquila.objects.filter(c_id_remito__iexact=remito_id)
        
        # Si estamos editando, excluimos el registro actual de la búsqueda
        if pk_actual and pk_actual != 'None' and pk_actual != '':
             qs = qs.exclude(pk=pk_actual)
             
        existe = qs.exists()
        return JsonResponse({'existe': existe})
        
    return JsonResponse({'existe': False})
    
def verificar_remito_existente(request):
    remito_id = request.GET.get('id', None)
    pk_actual = request.GET.get('pk_actual', None) # Para excluir el registro actual si estamos editando
    
    if remito_id:
        # Buscamos coincidencias exactas (case-insensitive)
        qs = EntradaMaquila.objects.filter(c_id_remito__iexact=remito_id)
        
        # Si estamos editando, excluimos el registro actual de la búsqueda
        if pk_actual and pk_actual != 'None' and pk_actual != '':
             qs = qs.exclude(pk=pk_actual)
             
        existe = qs.exists()
        return JsonResponse({'existe': existe})
        
    return JsonResponse({'existe': False})

def enviar_alerta_merma(remision):
    """
    Envía un correo solo si la merma es NEGATIVA (Pérdida) y supera el -1%.
    Fórmula: ((Descarga - Carga) / Carga) * 100
    """
    # 1. Validación de Estatus
    if remision.status != 'TERMINADO':
        return 

    # 2. Obtener pesos (convertir a float para cálculo seguro)
    carga = float(remision.total_peso_ld)
    descarga = float(remision.total_peso_dlv)

    # 3. Cálculo de Diferencia Estándar (Destino - Origen)
    # Si Descarga < Carga, el resultado será NEGATIVO (Pérdida)
    if carga > 0:
        porcentaje = ((descarga - carga) / carga) * 100
    else:
        porcentaje = 0

    # 4. CONDICIÓN SOLICITADA: Solo si es NEGATIVA (menor a -1.0%)
    # Ejemplo: Si resultado es -1.5%, se dispara la alerta.
    if porcentaje < -1.0: 
        asunto = f"ALERTA DE MERMA NEGATIVA: Remisión {remision.remision}"
        mensaje = f"""
        Se ha detectado una diferencia negativa (pérdida) superior al 1% en la remisión {remision.remision}.
        
        Detalles del Movimiento:
        ------------------------------------------------
        - Remisión:   {remision.remision}
        - Fecha:      {remision.fecha}
        - Origen:     {remision.origen}
        - Destino:    {remision.destino}
        
        Transporte:
        - Línea:      {remision.linea_transporte}
        - Operador:   {remision.operador}
        - Unidad:     {remision.unidad}
        
        Balance de Pesos:
        - Peso Carga:    {carga:,.3f} Kg
        - Peso Descarga: {descarga:,.3f} Kg
        - Diferencia:    {descarga - carga:,.3f} Kg
        
        >>> PORCENTAJE DE MERMA: {porcentaje:.2f}% <<<
        
        Por favor revisar en el sistema.
        """
        
        try:
            send_mail(
                asunto,
                mensaje,
                settings.DEFAULT_FROM_EMAIL,
                ['3rrecycling@gmail.com', 'aux.contable@3rrecycling.com.mx', 'calidad@3rrecycling.com.mx', 'mmartinez@3rrecycling.com.mx'], 
                fail_silently=True,
            )
            print(f"📧 Correo de alerta enviado para {remision.remision} (Merma: {porcentaje:.2f}%)")
        except Exception as e:
            print(f"❌ Error enviando correo: {e}")
            
@method_decorator(login_required, name='dispatch')
class ControlTarimaListView(ListView):
    model = ControlTarima
    template_name = 'ternium/lista_tarimas.html'
    context_object_name = 'tarimas'
    paginate_by = 20

    def get_queryset(self):
        qs = super().get_queryset().order_by('-fecha')
        
        # --- FILTROS ---
        q = self.request.GET.get('q')
        fecha_inicio = self.request.GET.get('fecha_inicio')
        fecha_fin = self.request.GET.get('fecha_fin')

        if q:
            qs = qs.filter(
                Q(origen__nombre__icontains=q) | 
                Q(destino__nombre__icontains=q) |
                Q(comentarios__icontains=q)
            )
        
        if fecha_inicio:
            qs = qs.filter(fecha__gte=fecha_inicio)
        
        if fecha_fin:
            qs = qs.filter(fecha__lte=fecha_fin)
            
        return qs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Pasamos los filtros al template para mantenerlos en los inputs
        context['filtros'] = self.request.GET
        return context

@login_required
def crear_tarima(request):
    if request.method == 'POST':
        form = ControlTarimaForm(request.POST, request.FILES)
        if form.is_valid():
            tarima = form.save(commit=False)
            
            # --- Lógica S3 Manual (Igual que tus otros módulos) ---
            if 'evidencia' in request.FILES:
                archivo = request.FILES['evidencia']
                nombre_limpio = archivo.name.replace(" ", "_")
                # Ruta: tarimas/FECHA_ORIGEN_archivo.ext
                s3_path = f"tarimas/{tarima.fecha}_{nombre_limpio}"
                ruta_guardada = _subir_archivo_a_s3(archivo, s3_path)
                if ruta_guardada:
                    tarima.evidencia = ruta_guardada
            # ----------------------------------------------------
            
            tarima.save()
            messages.success(request, 'Registro de tarimas creado exitosamente.')
            return redirect('lista_tarimas')
    else:
        form = ControlTarimaForm(initial={'fecha': timezone.now().date()})
    
    return render(request, 'ternium/formulario_tarima.html', {
        'form': form, 'titulo': 'Nuevo Registro de Tarimas'
    })

@login_required
def editar_tarima(request, pk):
    tarima = get_object_or_404(ControlTarima, pk=pk)
    if request.method == 'POST':
        form = ControlTarimaForm(request.POST, request.FILES, instance=tarima)
        if form.is_valid():
            obj = form.save(commit=False)
            
            # --- Actualización S3 ---
            if 'evidencia' in request.FILES:
                # Borrar anterior si existe
                if tarima.evidencia and hasattr(tarima.evidencia, 'name'):
                    _eliminar_archivo_de_s3(tarima.evidencia.name)
                
                archivo = request.FILES['evidencia']
                s3_path = f"tarimas/{obj.fecha}_{archivo.name.replace(' ', '_')}"
                ruta = _subir_archivo_a_s3(archivo, s3_path)
                if ruta: obj.evidencia = ruta
            # ------------------------

            obj.save()
            messages.success(request, 'Registro actualizado.')
            return redirect('lista_tarimas')
    else:
        form = ControlTarimaForm(instance=tarima)
        
    return render(request, 'ternium/formulario_tarima.html', {
        'form': form, 'titulo': f'Editar Tarima #{tarima.pk}'
    })

@login_required
def eliminar_tarima(request, pk):
    tarima = get_object_or_404(ControlTarima, pk=pk)
    if request.method == 'POST':
        # Borrar archivo de S3
        if tarima.evidencia and hasattr(tarima.evidencia, 'name'):
            _eliminar_archivo_de_s3(tarima.evidencia.name)
        tarima.delete()
        messages.success(request, 'Registro eliminado.')
    return redirect('lista_tarimas')

@login_required
def export_tarimas_excel(request):
    """
    Exporta el listado de tarimas aplicando los mismos filtros que la vista.
    """
    # 1. Obtener filtros
    q = request.GET.get('q')
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')
    
    # 2. Filtrar Queryset
    qs = ControlTarima.objects.all().order_by('-fecha')
    
    if q:
        qs = qs.filter(Q(origen__nombre__icontains=q) | Q(destino__nombre__icontains=q))
    if fecha_inicio:
        qs = qs.filter(fecha__gte=fecha_inicio)
    if fecha_fin:
        qs = qs.filter(fecha__lte=fecha_fin)

    # 3. Crear Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Control Tarimas"
    
    # Estilos
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
    center_align = Alignment(horizontal='center')
    money_fmt = '"$"#,##0.00'

    headers = [
    'ID', 'Fecha', 'Origen', 'Destino', 
    'Cant. Chica', 'Precio Chica', 'Total Chica',
    'Cant. Mediana', 'Precio Mediana', 'Total Mediana', # <-- Agregar
    'Cant. Grande', 'Precio Grande', 'Total Grande',
    'GRAN TOTAL', 'Comentarios'
]
    
    ws.append(headers)
    
    # Aplicar estilo a cabecera
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align

    # Llenar datos
    for obj in qs:
        row = [
            obj.id,
            obj.fecha,
            obj.origen.nombre if obj.origen else 'N/A',
            obj.destino.nombre if obj.destino else 'N/A',
            obj.tarima_chica_cant,
            obj.precio_chica,
            obj.total_chica,
            obj.tarima_grande_cant,
            obj.precio_grande,
            obj.total_grande,
            obj.gran_total,
            obj.comentarios
        ]
        ws.append(row)
        
        # Formato de moneda para columnas de dinero
        current_row = ws.max_row
        # Precio Chica (F), Total Chica (G), Precio Grande (I), Total Grande (J), Gran Total (K)
        for col_idx in [6, 7, 9, 10, 11]: 
            ws.cell(row=current_row, column=col_idx).number_format = money_fmt

    # Ajustar anchos
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except: pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Reporte_Tarimas_{timezone.now().date()}.xlsx"'
    wb.save(response)
    return response

@login_required
def export_logistica_reporte_mensual(request):
    """
    OPCIÓN 1: Resumen Ejecutivo Mensual.
    LÓGICA 3R: Real (Recibidas en Ternium) - Documento (Remisionadas por 3R).
    - Resultado Negativo: Pérdida / Faltante (Rojo).
    - Resultado Positivo: Ganancia / Sobrante (Verde).
    """
    try:
        year = int(request.GET.get('year', datetime.date.today().year))
    except ValueError:
        year = datetime.date.today().year

    wb = Workbook()
    ws = wb.active
    ws.title = f"Resumen Mensual {year}"

    # --- ESTILOS ---
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="0d6efd", end_color="0d6efd", fill_type="solid") # Azul Bootstrap
    
    total_font = Font(bold=True)
    total_fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid") # Amarillo Claro
    
    # Fuentes para Resultados
    red_font = Font(color="DC3545", bold=True)   # Rojo para pérdidas
    green_font = Font(color="198754", bold=True) # Verde para ganancias/ok
    
    center = Alignment(horizontal='center', vertical='center')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    number_fmt = '#,##0.000'
    percent_fmt = '0.00%'
    int_fmt = '#,##0'

    meses_nombres = {
        1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL', 5: 'MAYO', 6: 'JUNIO',
        7: 'JULIO', 8: 'AGOSTO', 9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
    }

    # ==========================
    # 1. ENCABEZADOS
    # ==========================
    ws.merge_cells('A1:H1')
    ws['A1'] = f"RESUMEN EJECUTIVO MENSUAL DE LOGÍSTICA - {year}"
    ws['A1'].font = Font(bold=True, size=14, color="0d6efd")
    ws['A1'].alignment = center

    ws.merge_cells('A2:H2')
    ws['A2'] = "LÓGICA: (Ton. Recibidas - Ton. Remisionadas). Negativo = Pérdida."
    ws['A2'].font = Font(italic=True, color="6c757d", size=10)
    ws['A2'].alignment = center

    # Cabeceras de columnas
    headers = [
        ('MES', 15),
        ('VIAJES', 10),
        ('TON. REMISIONADAS', 20),
        ('TON. RECIBIDAS', 20),
        ('DIFERENCIA (NETA)', 20),
        ('% MERMA/GANANCIA', 18),
        ('PROM. CARGA', 15),
        ('ESTATUS', 15)
    ]

    for col_num, (header_title, width) in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num, value=header_title)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = thin_border
        ws.column_dimensions[get_column_letter(col_num)].width = width

    # ==========================
    # 2. CONSULTA DE DATOS
    # ==========================
    qs = RegistroLogistico.objects.filter(
        fecha_carga__year=year,
        status__in=['TERMINADO', 'AUDITADO']
    ).annotate(
        mes=TruncMonth('fecha_carga')
    ).values('mes').annotate(
        total_rem=Coalesce(Sum('toneladas_remisionadas'), 0.0, output_field=FloatField()),
        total_rec=Coalesce(Sum('toneladas_recibidas'), 0.0, output_field=FloatField()),
        total_viajes=Count('id')
    ).order_by('mes')

    data_by_month = {item['mes'].month: item for item in qs} if qs else {}

    # ==========================
    # 3. PROCESAMIENTO
    # ==========================
    row_idx = 4
    acc_viajes = 0
    acc_rem = 0.0
    acc_rec = 0.0

    for mes_num in range(1, 13):
        data = data_by_month.get(mes_num, {'total_rem': 0, 'total_rec': 0, 'total_viajes': 0})
        
        t_rem = float(data['total_rem'])
        t_rec = float(data['total_rec'])
        viajes = data['total_viajes']
        
        # --- NUEVA LÓGICA ---
        # Real (Recibido) - Documento (Remisionado)
        diferencia = t_rec - t_rem 
        
        merma_pct = (diferencia / t_rem) if t_rem > 0 else 0.0
        promedio_carga = (t_rec / viajes) if viajes > 0 else 0.0
        
        acc_viajes += viajes
        acc_rem += t_rem
        acc_rec += t_rec

        # A: Mes
        ws.cell(row=row_idx, column=1, value=meses_nombres[mes_num]).border = thin_border
        
        # B: Viajes
        c_viajes = ws.cell(row=row_idx, column=2, value=viajes)
        c_viajes.number_format = int_fmt
        c_viajes.alignment = center
        c_viajes.border = thin_border

        # C: Remisionadas
        ws.cell(row=row_idx, column=3, value=t_rem).number_format = number_fmt
        ws.cell(row=row_idx, column=3).border = thin_border

        # D: Recibidas
        ws.cell(row=row_idx, column=4, value=t_rec).number_format = number_fmt
        ws.cell(row=row_idx, column=4).border = thin_border

        # E: Diferencia (Con color)
        c_dif = ws.cell(row=row_idx, column=5, value=diferencia)
        c_dif.number_format = number_fmt
        c_dif.border = thin_border
        
        if diferencia < 0:
            c_dif.font = red_font   # Negativo en Rojo
        elif diferencia > 0:
            c_dif.font = green_font # Positivo en Verde

        # F: % Merma (Con color)
        c_pct = ws.cell(row=row_idx, column=6, value=merma_pct)
        c_pct.number_format = percent_fmt
        c_pct.border = thin_border
        
        if merma_pct < 0:
            c_pct.font = red_font
        elif merma_pct > 0:
            c_pct.font = green_font

        # G: Promedio Carga
        ws.cell(row=row_idx, column=7, value=promedio_carga).number_format = number_fmt
        ws.cell(row=row_idx, column=7).border = thin_border

        # H: Estatus Visual
        if viajes == 0:
            estatus = "-"
            est_color = "000000"
        elif merma_pct < -0.01: # Peor que -1%
            estatus = "REVISAR"
            est_color = "DC3545" # Rojo
        elif merma_pct < 0: # Entre 0% y -1%
            estatus = "REGULAR"
            est_color = "FFC107" # Amarillo
        else:
            estatus = "OK"
            est_color = "198754" # Verde

        c_stat = ws.cell(row=row_idx, column=8, value=estatus)
        c_stat.font = Font(bold=True, color=est_color)
        c_stat.alignment = center
        c_stat.border = thin_border

        row_idx += 1

    # ==========================
    # 4. FILA DE TOTALES
    # ==========================
    ws.cell(row=row_idx, column=1, value="TOTAL ANUAL").font = total_font
    ws.cell(row=row_idx, column=1).fill = total_fill
    ws.cell(row=row_idx, column=1).border = thin_border

    # Total Viajes
    c_t_viajes = ws.cell(row=row_idx, column=2, value=acc_viajes)
    c_t_viajes.font = total_font; c_t_viajes.fill = total_fill; c_t_viajes.border = thin_border
    c_t_viajes.alignment = center

    # Total Rem
    c_t_rem = ws.cell(row=row_idx, column=3, value=acc_rem)
    c_t_rem.font = total_font; c_t_rem.fill = total_fill; c_t_rem.border = thin_border
    c_t_rem.number_format = number_fmt

    # Total Rec
    c_t_rec = ws.cell(row=row_idx, column=4, value=acc_rec)
    c_t_rec.font = total_font; c_t_rec.fill = total_fill; c_t_rec.border = thin_border
    c_t_rec.number_format = number_fmt

    # Total Diferencia
    acc_dif = acc_rec - acc_rem
    c_t_dif = ws.cell(row=row_idx, column=5, value=acc_dif)
    c_t_dif.fill = total_fill; c_t_dif.border = thin_border
    c_t_dif.number_format = number_fmt
    if acc_dif < 0:
        c_t_dif.font = Font(bold=True, color="DC3545") # Rojo sobre amarillo
    else:
        c_t_dif.font = Font(bold=True, color="198754")

    # Total Merma %
    acc_merma_pct = (acc_dif / acc_rem) if acc_rem > 0 else 0.0
    c_t_merma = ws.cell(row=row_idx, column=6, value=acc_merma_pct)
    c_t_merma.fill = total_fill; c_t_merma.border = thin_border
    c_t_merma.number_format = percent_fmt
    if acc_merma_pct < 0:
        c_t_merma.font = Font(bold=True, color="DC3545")
    else:
        c_t_merma.font = Font(bold=True, color="198754")

    # Promedio Global
    acc_prom = (acc_rec / acc_viajes) if acc_viajes > 0 else 0.0
    c_t_prom = ws.cell(row=row_idx, column=7, value=acc_prom)
    c_t_prom.font = total_font; c_t_prom.fill = total_fill; c_t_prom.border = thin_border
    c_t_prom.number_format = number_fmt

    # Cierre fila totales
    ws.cell(row=row_idx, column=8, value="").fill = total_fill
    ws.cell(row=row_idx, column=8).border = thin_border

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Resumen_Mensual_Logistica_{year}.xlsx"'
    wb.save(response)
    return response

    # =========================================================================
    # FUNCIÓN GENERADORA (REUTILIZABLE PARA AMBAS HOJAS)
    # =========================================================================
def _generar_excel_jerarquico(year, titulo_hoja, campo_primario, campo_secundario, titulo_col_a, color_grupo, nombre_archivo):
    """
    Genera un archivo Excel con una sola hoja en formato Matriz Jerárquica.
    FILTRO: Solo estatus TERMINADO o AUDITADO.
    LÓGICA 3R: Real (Recibido) - Documento (Remisionado).
    """
    wb = Workbook()
    ws = wb.active
    ws.title = titulo_hoja
    
    # --- ESTILOS ---
    header_month_font = Font(bold=True, color="FFFFFF")
    header_month_fill = PatternFill(start_color="0d6efd", end_color="0d6efd", fill_type="solid") # Azul
    
    header_sub_font = Font(bold=True, size=9)
    header_sub_fill = PatternFill(start_color="E9ECEF", end_color="E9ECEF", fill_type="solid") # Gris Claro
    
    group_font = Font(bold=True, color="FFFFFF", size=12)
    # color_grupo viene como parámetro
    
    total_font = Font(bold=True)
    total_fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid") # Amarillo

    # Estilos de Resultado
    red_font = Font(color="DC3545", bold=True)   # Rojo para pérdidas
    green_font = Font(color="198754", bold=True) # Verde para ganancias

    center = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    number_fmt = '#,##0.000'
    percent_fmt = '0.00%'

    meses_nombres = {
        1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL', 5: 'MAYO', 6: 'JUNIO',
        7: 'JULIO', 8: 'AGOSTO', 9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
    }

    # =========================================================================
    # 1. ENCABEZADOS Y MENSAJE
    # =========================================================================
    
    # Fila 1: Encabezados Principales
    ws.merge_cells('A1:A2') 
    ws['A1'] = titulo_col_a
    ws['A1'].alignment = center
    ws['A1'].font = Font(bold=True)
    ws['A1'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

    col_idx = 2
    for mes in range(1, 13):
        start_col = get_column_letter(col_idx)
        end_col = get_column_letter(col_idx + 2)
        ws.merge_cells(f'{start_col}1:{end_col}1')
        
        cell = ws.cell(row=1, column=col_idx, value=meses_nombres[mes])
        cell.font = header_month_font
        cell.fill = header_month_fill
        cell.alignment = center
        cell.border = thin_border
        
        # Sub-encabezados: REM (Remisionado), REC (Recibido), % (Diferencia)
        sub_headers = ["REM", "REC", "%"]
        for i, sub in enumerate(sub_headers):
            c_sub = ws.cell(row=2, column=col_idx + i, value=sub)
            c_sub.font = header_sub_font
            c_sub.fill = header_sub_fill
            c_sub.alignment = center
            c_sub.border = thin_border
        col_idx += 3

    # Totales Anuales
    ws.merge_cells(f'{get_column_letter(col_idx)}1:{get_column_letter(col_idx + 2)}1')
    ws.cell(row=1, column=col_idx, value="TOTAL ANUAL").font = header_month_font
    ws.cell(row=1, column=col_idx).fill = header_month_fill
    ws.cell(row=1, column=col_idx).alignment = center
    
    ws.cell(row=2, column=col_idx, value="REM").font = header_sub_font
    ws.cell(row=2, column=col_idx+1, value="REC").font = header_sub_font
    ws.cell(row=2, column=col_idx+2, value="%").font = header_sub_font

    # --- INSERTAR FILA DE AVISO ---
    ws.insert_rows(1)
    last_col_letter = get_column_letter(col_idx + 2)
    ws.merge_cells(f'A1:{last_col_letter}1')
    ws['A1'] = "NOTA: Estatus TERMINADO/AUDITADO. Fórmula: (Recibido - Remisionado). Negativo = Pérdida."
    ws['A1'].font = Font(italic=True, color="6c757d", size=10)
    ws['A1'].alignment = center
    # -------------------------------------

    # =========================================================================
    # 2. DATOS
    # =========================================================================
    qs = RegistroLogistico.objects.filter(
        fecha_carga__year=year,
        status__in=['TERMINADO', 'AUDITADO']
    ).values(
        campo_primario, 
        campo_secundario, 
        'fecha_carga__month'
    ).annotate(
        sum_rem=Coalesce(Sum('toneladas_remisionadas'), 0.0, output_field=FloatField()),
        sum_rec=Coalesce(Sum('toneladas_recibidas'), 0.0, output_field=FloatField())
    ).order_by(campo_primario, campo_secundario)

    reporte_data = {}
    for item in qs:
        prim = item[campo_primario] or "SIN DEFINIR"
        sec = item[campo_secundario] or "SIN DEFINIR"
        mes = item['fecha_carga__month']
        
        if prim not in reporte_data:
            reporte_data[prim] = {}
        if sec not in reporte_data[prim]:
            reporte_data[prim][sec] = {m: {'rem': 0.0, 'rec': 0.0} for m in range(1, 13)}
        
        reporte_data[prim][sec][mes]['rem'] += item['sum_rem']
        reporte_data[prim][sec][mes]['rec'] += item['sum_rec']

    # =========================================================================
    # 3. ESCRIBIR FILAS
    # =========================================================================
    current_row = 4
    
    for grupo_principal, sub_items in reporte_data.items():
        # A. Encabezado Grupo (Ej: Proveedor A)
        last_col_letter = get_column_letter(40)
        
        cell_group = ws.cell(row=current_row, column=1, value=grupo_principal)
        cell_group.font = group_font
        cell_group.fill = color_grupo
        cell_group.alignment = left_align
        
        # Merge visual para el título del grupo
        # (Opcional: Si quieres mergear toda la fila, descomenta la sig línea)
        # ws.merge_cells(f'A{current_row}:{last_col_letter}{current_row}')
        
        current_row += 1

        total_grupo_meses = {m: {'rem': 0.0, 'rec': 0.0} for m in range(1, 13)}
        total_grupo_anual_rem = 0.0
        total_grupo_anual_rec = 0.0

        # B. Sub-items (Ej: Material 1, Material 2...)
        for sub_nombre, meses_data in sub_items.items():
            ws.cell(row=current_row, column=1, value=sub_nombre).border = thin_border
            anual_sub_rem = 0.0
            anual_sub_rec = 0.0
            col_ptr = 2

            # Loop Meses
            for mes in range(1, 13):
                val_rem = meses_data[mes]['rem']
                val_rec = meses_data[mes]['rec']
                
                # CÁLCULO 3R
                diferencia = val_rec - val_rem
                merma_pct = (diferencia / val_rem) if val_rem > 0 else 0.0

                anual_sub_rem += val_rem
                anual_sub_rec += val_rec
                total_grupo_meses[mes]['rem'] += val_rem
                total_grupo_meses[mes]['rec'] += val_rec

                # Celdas Mes
                c1 = ws.cell(row=current_row, column=col_ptr, value=val_rem)
                c1.number_format = number_fmt; c1.border = thin_border
                
                c2 = ws.cell(row=current_row, column=col_ptr + 1, value=val_rec)
                c2.number_format = number_fmt; c2.border = thin_border
                
                c3 = ws.cell(row=current_row, column=col_ptr + 2, value=merma_pct)
                c3.number_format = percent_fmt; c3.border = thin_border
                
                # Color Dinámico
                if merma_pct < 0:
                    c3.font = red_font
                elif merma_pct > 0:
                    c3.font = green_font

                col_ptr += 3

            # Totales Anuales Sub-Item
            ws.cell(row=current_row, column=col_ptr, value=anual_sub_rem).number_format = number_fmt
            ws.cell(row=current_row, column=col_ptr+1, value=anual_sub_rec).number_format = number_fmt
            
            dif_anual = anual_sub_rec - anual_sub_rem
            merma_anual = (dif_anual / anual_sub_rem) if anual_sub_rem > 0 else 0.0
            
            c_anual = ws.cell(row=current_row, column=col_ptr+2, value=merma_anual)
            c_anual.number_format = percent_fmt
            if merma_anual < 0: c_anual.font = red_font
            elif merma_anual > 0: c_anual.font = green_font

            total_grupo_anual_rem += anual_sub_rem
            total_grupo_anual_rec += anual_sub_rec
            current_row += 1

        # C. Total del Grupo
        ws.cell(row=current_row, column=1, value=f"TOTAL {grupo_principal}").font = total_font
        ws.cell(row=current_row, column=1).fill = total_fill
        ws.cell(row=current_row, column=1).border = thin_border
        
        col_ptr = 2
        for mes in range(1, 13):
            t_rem = total_grupo_meses[mes]['rem']
            t_rec = total_grupo_meses[mes]['rec']
            
            # CÁLCULO 3R GRUPAL
            t_dif = t_rec - t_rem
            t_merma = (t_dif / t_rem) if t_rem > 0 else 0.0
            
            # Escribir celdas totales
            c_trem = ws.cell(row=current_row, column=col_ptr, value=t_rem)
            c_trem.font = total_font; c_trem.fill = total_fill; c_trem.border = thin_border; c_trem.number_format = number_fmt
            
            c_trec = ws.cell(row=current_row, column=col_ptr+1, value=t_rec)
            c_trec.font = total_font; c_trec.fill = total_fill; c_trec.border = thin_border; c_trec.number_format = number_fmt
            
            c_tpct = ws.cell(row=current_row, column=col_ptr+2, value=t_merma)
            c_tpct.font = total_font; c_tpct.fill = total_fill; c_tpct.border = thin_border; c_tpct.number_format = percent_fmt
            
            if t_merma < 0: c_tpct.font = Font(color="DC3545", bold=True)
            elif t_merma > 0: c_tpct.font = Font(color="198754", bold=True)
            
            col_ptr += 3
        
        # Total Anual Grupo
        # REM
        c_gt_rem = ws.cell(row=current_row, column=col_ptr, value=total_grupo_anual_rem)
        c_gt_rem.font = total_font; c_gt_rem.fill = total_fill; c_gt_rem.number_format = number_fmt
        
        # REC
        c_gt_rec = ws.cell(row=current_row, column=col_ptr+1, value=total_grupo_anual_rec)
        c_gt_rec.font = total_font; c_gt_rec.fill = total_fill; c_gt_rec.number_format = number_fmt
        
        # %
        tot_dif_final = total_grupo_anual_rec - total_grupo_anual_rem
        tot_merma_final = (tot_dif_final / total_grupo_anual_rem) if total_grupo_anual_rem > 0 else 0.0
        
        c_gt_pct = ws.cell(row=current_row, column=col_ptr+2, value=tot_merma_final)
        c_gt_pct.font = total_font; c_gt_pct.fill = total_fill; c_gt_pct.number_format = percent_fmt
        
        if tot_merma_final < 0: c_gt_pct.font = Font(color="DC3545", bold=True)
        elif tot_merma_final > 0: c_gt_pct.font = Font(color="198754", bold=True)

        current_row += 2 

    # Ajustar Anchos
    ws.column_dimensions['A'].width = 35
    for i in range(2, 45): 
        ws.column_dimensions[get_column_letter(i)].width = 11

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}_{year}.xlsx"'
    wb.save(response)
    return response


# =========================================================================
# VISTAS INDIVIDUALES (LAS QUE LLAMA EL HTML)
# =========================================================================
    
def enviar_alerta_merma_logistica(registro):
    """
    Envía una alerta por correo si el Registro Logístico está TERMINADO y:
    - La merma es >= 1% (Falta material)
    - O la merma es <= -1% (Sobra material / Excedente)
    """
    # 1. Validar Estatus
    if registro.status != 'TERMINADO':
        return

    # 2. Validar Porcentaje
    pct = registro.merma_porcentaje
    
    if pct is None:
        return

    # 3. Lógica: Valor absoluto mayor o igual a 1 (Cubre 1.5% y -1.22%)
    if abs(pct) >= 1.0:
        
        # Determinar el tipo de diferencia para el asunto
        if pct > 0:
            tipo_alerta = "MERMA (Faltante)"
            color_alerta = "ROJO" # Solo informativo en texto
        else:
            tipo_alerta = "EXCEDENTE (Sobrante)"
            color_alerta = "AMARILLO"

        # Construir Asunto y Mensaje
        asunto = f"ALERTA LOGÍSTICA: {tipo_alerta} del {pct:.2f}% - Remisión {registro.remision}"
        
        mensaje = f"""
        ALERTA DE DIFERENCIA DE PESO ({tipo_alerta})
        --------------------------------------------------
        Se ha detectado una variación fuera del rango permitido (>= 1% o <= -1%).

        RESUMEN:
        - Remisión:      {registro.remision}
        - Fecha Carga:   {registro.fecha_carga}
        - Estatus:       {registro.get_status_display()}
        
        TRANSPORTE:
        - Transportista: {registro.transportista}
        - Chofer:        {registro.chofer}
        - Unidad:        {registro.placas_tractor or registro.tractor or 'N/A'} 
        - Tolva/Caja:    {registro.placas_tolva or registro.tolva or 'N/A'}

        MATERIAL:
        - Material:      {registro.material}
        
        PESOS (TONELADAS):
        - Remisionadas:  {registro.toneladas_remisionadas}
        - Recibidas:     {registro.toneladas_recibidas}
        - Diferencia:    {registro.merma_absoluta} Ton
        
        >>> PORCENTAJE: {pct:.2f}% <<<
        
        COMENTARIOS:
        {registro.comentario or 'Sin comentarios'}
        
        --------------------------------------------------
        Este es un mensaje automático del sistema.
        """

        # 4. Enviar Correo
        try:
            send_mail(
                asunto,
                mensaje,
                settings.DEFAULT_FROM_EMAIL,
                # Ajusta tus destinatarios aquí
                ['aux.contable@3rrecycling.com.mx', 'calidad@3rrecycling.com.mx', 'mmartinez@3rrecycling.com.mx'], 
                fail_silently=True,
            )
            print(f"📧 Correo de alerta ({tipo_alerta}) enviado para {registro.remision}")
        except Exception as e:
            print(f"❌ Error enviando correo logística: {e}")

            
@login_required
def export_balanza_pesos(request):
    """
    OPCIÓN 4: Balanza de Pesos (Entradas vs Salidas).
    Compara Peso TERNIUM vs Peso 3R.
    LÓGICA ACTUALIZADA: Real - Documento.
    - Negativo: Pérdida / Faltante.
    - Positivo: Ganancia / Sobrante.
    """
    try:
        year = int(request.GET.get('year', datetime.date.today().year))
    except ValueError:
        year = datetime.date.today().year

    wb = Workbook()
    ws = wb.active
    ws.title = f"Balanza Materiales {year}"

    # --- ESTILOS ---
    header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid") # Azul Oscuro
    header_font = Font(bold=True, color="FFFFFF")
    
    # Colores de fondo para distinguir tipos de movimiento
    row_entrada_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid") # Verde Claro (Ingreso)
    row_salida_fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid") # Azul Claro (Salida)
    
    # Fuentes para Diferencias
    alert_font = Font(color="DC3545", bold=True) # Rojo (Pérdida/Negativo)
    good_font = Font(color="198754", bold=True)  # Verde (Ganancia/Positivo)

    center = Alignment(horizontal='center', vertical='center')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    number_fmt = '#,##0.000'

    # 1. ENCABEZADOS
    headers = [
        "FECHA",
        "MOVIMIENTO",
        "FOLIO / REMISIÓN",
        "MATERIAL / CALIDAD", 
        "PESO TERNIUM (Ton)", 
        "PESO 3R (Ton)",       
        "DIFERENCIA (Ton)", # Real - Papel
        "% MERMA"
    ]

    ws.append(headers)
    for col_num, cell in enumerate(ws[1], 1):
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border
        ws.column_dimensions[get_column_letter(col_num)].width = 18

    # =========================================================
    # 2. OBTENER DATOS (ENTRADAS Y SALIDAS)
    # =========================================================
    data_rows = []

    # A) ENTRADAS (EntradaMaquila)
    # Contexto: Ternium envía (Papel), 3R recibe (Real)
    entradas = EntradaMaquila.objects.filter(fecha_ingreso__year=year)
    
    for e in entradas:
        p_ternium = float(e.peso_remision or 0) # Papel / Documento
        p_3r = float(e.peso_neto or 0)          # Real / Físico
        
        # LÓGICA: Real (3R) - Papel (Ternium)
        # Si llegaron 90 (3R) y decían 100 (Ternium) -> -10 (Pérdida)
        diff = p_3r - p_ternium
        pct = (diff / p_ternium) if p_ternium > 0 else 0

        # Nombre del material/calidad
        nombre_material = str(e.calidad) if e.calidad else "N/A"

        data_rows.append({
            'fecha': e.fecha_ingreso,
            'tipo': 'ENTRADA (Ingreso)',
            'folio': e.num_boleta_remision or f"ID-{e.pk}",
            'material': nombre_material,
            'peso_ternium': p_ternium, # Papel
            'peso_3r': p_3r,           # Real
            'diferencia': diff,
            'pct': pct,
            'is_entrada': True
        })

    # B) SALIDAS (RegistroLogistico)
    # Contexto: 3R envía (Papel), Ternium recibe (Real)
    salidas = RegistroLogistico.objects.filter(
        fecha_carga__year=year,
        status__in=['TERMINADO', 'AUDITADO']
    ).select_related('material')

    for s in salidas:
        p_3r = float(s.toneladas_remisionadas or 0) # Papel / Documento (Lo que envió 3R)
        p_ternium = float(s.toneladas_recibidas or 0) # Real / Físico (Lo que pesó Ternium)
        
        # LÓGICA: Real (Ternium) - Papel (3R)
        # Si Ternium recibió 90 y 3R envió 100 -> -10 (Pérdida)
        diff = p_ternium - p_3r 
        pct = (diff / p_3r) if p_3r > 0 else 0

        data_rows.append({
            'fecha': s.fecha_carga,
            'tipo': 'SALIDA (Envío)',
            'folio': s.remision,
            'material': s.material.nombre if s.material else "N/A",
            'peso_ternium': p_ternium, # Real
            'peso_3r': p_3r,           # Papel
            'diferencia': diff,
            'pct': pct,
            'is_entrada': False
        })

    # 3. ORDENAR POR FECHA
    data_rows.sort(key=lambda x: x['fecha'], reverse=True)

    # =========================================================
    # 3. ESCRIBIR EN EXCEL
    # =========================================================
    for row_data in data_rows:
        row = [
            row_data['fecha'],
            row_data['tipo'],
            row_data['folio'],
            row_data['material'],
            row_data['peso_ternium'],
            row_data['peso_3r'],
            row_data['diferencia'],
            row_data['pct']
        ]
        ws.append(row)
        
        current_row = ws.max_row
        
        # Color de fondo para diferenciar Entradas de Salidas
        fill_color = row_entrada_fill if row_data['is_entrada'] else row_salida_fill
        
        # Variables para estilos condicionales
        val_diff = row_data['diferencia']
        val_pct = row_data['pct']

        for i in range(1, 9):
            cell = ws.cell(row=current_row, column=i)
            cell.border = thin_border
            cell.alignment = center
            cell.fill = fill_color
            
            # Formatos numéricos para Pesos y Diferencia (Cols 5, 6, 7)
            if i in [5, 6, 7]: 
                cell.number_format = number_fmt
            
            # Estilo condicional para la Columna Diferencia (7) y Porcentaje (8)
            if i in [7, 8]:
                # Negativo = Pérdida -> Rojo
                if val_diff < 0:
                    cell.font = alert_font
                # Positivo = Ganancia -> Verde (Opcional, ayuda visual)
                elif val_diff > 0:
                    cell.font = good_font
            
            # Formato Porcentaje específico (Col 8)
            if i == 8: 
                cell.number_format = '0.00%'

    ws.column_dimensions['C'].width = 25 
    ws.column_dimensions['D'].width = 30 

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Balanza_Pesos_Ternium_3R_{year}.xlsx"'
    wb.save(response)
    return response

@login_required
@require_POST
def rechazar_registro_logistica(request, pk):
    registro = get_object_or_404(RegistroLogistico, pk=pk)
    
    # Validaciones de seguridad
    if registro.status == 'AUDITADO':
        messages.error(request, 'No se puede rechazar un registro auditado.')
    elif registro.status == 'CANCELADO':
        messages.error(request, 'El registro está cancelado.')
    else:
        registro.status = 'RECHAZADO'
        registro.save(update_fields=['status'])
        messages.warning(request, f'El registro {registro.remision} ha sido marcado como RECHAZADO.')
        
    return redirect('lista_registros_logistica')

@login_required
@require_POST
def restablecer_registro_logistica(request, pk):
    """
    Permite volver a poner en circulación un registro Rechazado.
    Lo regresa a PENDIENTE para que sea validado nuevamente.
    """
    registro = get_object_or_404(RegistroLogistico, pk=pk)
    
    if registro.status == 'RECHAZADO':
        registro.status = 'PENDIENTE' # Regresa a pendiente para revisión
        registro.save(update_fields=['status'])
        messages.success(request, f'El registro {registro.remision} ha sido RESTABLECIDO y está Pendiente.')
    else:
        messages.error(request, 'Solo se pueden restablecer registros con estatus Rechazado.')
        
    return redirect('lista_registros_logistica')


def _build_pivot_excel_logistica(year, tipo_reporte):
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    
    # --- CONFIGURACIÓN AVANZADA DE EXCEL (IDÉNTICO A TABLA DINÁMICA) ---
    ws.sheet_properties.outlinePr.summaryBelow = False  # Para que el botón [-] quede en el padre
    ws.freeze_panes = 'B5'  # Congelar encabezados y columna de nombres
    ws.sheet_view.showGridLines = False # Ocultar cuadrícula de fondo para aspecto limpio
    
    # --- COLORES EXACTOS DE LAS FOTOS ---
    blue_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    gray_fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
    
    # --- FUENTES (Calibri 11 es el estándar de Tablas Dinámicas) ---
    white_bold = Font(name='Calibri', size=11, color="FFFFFF", bold=True)
    black_bold = Font(name='Calibri', size=11, color="000000", bold=True)
    black_normal = Font(name='Calibri', size=11, color="000000")
    
    # --- ALINEACIONES (Incluye sangría para los subtipos) ---
    center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center")
    left_indent = Alignment(horizontal="left", vertical="center", indent=2) # Sangría visual
    
    # --- BORDES GRISES SUTILES ---
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    
    number_fmt = '#,##0.00'
    
    meses_nombres = {
        1: 'Ene', 2: 'Feb', 3: 'Mar', 4: 'Abr', 5: 'May', 6: 'Jun',
        7: 'Jul', 8: 'Ago', 9: 'Sep', 10: 'Oct', 11: 'Nov', 12: 'Dic'
    }
    
    # --- CONSULTA BASE ---
    qs = RegistroLogistico.objects.filter(
        fecha_carga__year=year,
        status__in=['TERMINADO', 'AUDITADO']
    )
    
    # --- CONSTRUCCIÓN DE MATRIZ DE DATOS ---
    matrix = {}
    
    if tipo_reporte == 'proveedor':
        ws.title = "Por Proveedor"
        datos = qs.values('transportista__nombre', 'material__nombre', 'fecha_carga__month').annotate(
            rem=Sum('toneladas_remisionadas'),
            rec=Sum('toneladas_recibidas')
        )
        for d in datos:
            prov = d['transportista__nombre'] or "SIN PROVEEDOR"
            mat = d['material__nombre'] or "SIN MATERIAL"
            mes = d['fecha_carga__month']
            rem = float(d['rem'] or 0)
            rec = float(d['rec'] or 0)
            
            if prov not in matrix: matrix[prov] = {}
            if mat not in matrix[prov]: matrix[prov][mat] = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
            
            matrix[prov][mat][mes]['rem'] += rem
            matrix[prov][mat][mes]['rec'] += rec
            
    elif tipo_reporte == 'material':
        ws.title = "Por Material"
        datos = qs.values('material__nombre', 'fecha_carga__month').annotate(
            rem=Sum('toneladas_remisionadas'),
            rec=Sum('toneladas_recibidas')
        )
        for d in datos:
            mat = d['material__nombre'] or "SIN MATERIAL"
            mes = d['fecha_carga__month']
            rem = float(d['rem'] or 0)
            rec = float(d['rec'] or 0)
            
            if mat not in matrix: matrix[mat] = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
            
            matrix[mat][mes]['rem'] += rem
            matrix[mat][mes]['rec'] += rec

    # --- ALTURA DE FILAS ---
    ws.row_dimensions[3].height = 20
    ws.row_dimensions[4].height = 45

    # --- ENCABEZADOS (Filas 3 y 4) ---
    c_val = ws.cell(row=3, column=1, value="Valores")
    c_val.font = black_bold
    c_val.fill = gray_fill
    c_val.alignment = center_wrap
    c_val.border = thin_border
    
    c_a4 = ws.cell(row=4, column=1, value="PROVEEDOR (TRANSPORTISTA)" if tipo_reporte == 'proveedor' else "Material")
    c_a4.font = black_bold
    c_a4.fill = gray_fill
    c_a4.alignment = left_align
    c_a4.border = thin_border
    
    col_idx = 2
    for m in range(1, 13):
        # Header Mes (Azul)
        ws.merge_cells(start_row=3, start_column=col_idx, end_row=3, end_column=col_idx+1)
        c_mes = ws.cell(row=3, column=col_idx, value=meses_nombres[m])
        c_mes.fill = blue_fill
        c_mes.font = white_bold
        c_mes.alignment = center_wrap
        c_mes.border = thin_border
        ws.cell(row=3, column=col_idx+1).border = thin_border
        
        # Sub-headers (Gris - Letra normal)
        c_rem = ws.cell(row=4, column=col_idx, value="Suma de TONELADAS REMISIONADAS")
        c_rem.fill = gray_fill
        c_rem.font = black_normal
        c_rem.alignment = center_wrap
        c_rem.border = thin_border
        
        c_rec = ws.cell(row=4, column=col_idx+1, value="Suma de TONELADAS RECIBIDAS TERNIUM")
        c_rec.fill = gray_fill
        c_rec.font = black_normal
        c_rec.alignment = center_wrap
        c_rec.border = thin_border
        
        ws.column_dimensions[get_column_letter(col_idx)].width = 16
        ws.column_dimensions[get_column_letter(col_idx+1)].width = 16
        
        col_idx += 2
        
    # Total General Header
    ws.merge_cells(start_row=3, start_column=col_idx, end_row=3, end_column=col_idx+1)
    c_gt = ws.cell(row=3, column=col_idx, value="Total general")
    c_gt.fill = blue_fill
    c_gt.font = white_bold
    c_gt.alignment = center_wrap
    c_gt.border = thin_border
    ws.cell(row=3, column=col_idx+1).border = thin_border
    
    c_rem_gt = ws.cell(row=4, column=col_idx, value="Suma de TONELADAS REMISIONADAS")
    c_rem_gt.fill = gray_fill
    c_rem_gt.font = black_normal
    c_rem_gt.alignment = center_wrap
    c_rem_gt.border = thin_border
    
    c_rec_gt = ws.cell(row=4, column=col_idx+1, value="Suma de TONELADAS RECIBIDAS TERNIUM")
    c_rec_gt.fill = gray_fill
    c_rec_gt.font = black_normal
    c_rec_gt.alignment = center_wrap
    c_rec_gt.border = thin_border
    
    ws.column_dimensions[get_column_letter(col_idx)].width = 16
    ws.column_dimensions[get_column_letter(col_idx+1)].width = 16
    
    # --- FUNCIÓN AUXILIAR PARA ESCRIBIR VALORES ---
    def escribir_valores(row_num, dict_meses, es_negrita):
        idx = 2
        t_rem = 0.0
        t_rec = 0.0
        for m in range(1, 13):
            v_rem = dict_meses[m]['rem']
            v_rec = dict_meses[m]['rec']
            t_rem += v_rem
            t_rec += v_rec
            
            # Si el valor es 0, dejamos la celda "None" (Vacía) para copiar el efecto de la foto
            c1 = ws.cell(row=row_num, column=idx, value=v_rem if v_rem > 0 else None)
            c2 = ws.cell(row=row_num, column=idx+1, value=v_rec if v_rec > 0 else None)
            
            c1.number_format = number_fmt
            c2.number_format = number_fmt
            c1.border = thin_border
            c2.border = thin_border
            c1.font = black_bold if es_negrita else black_normal
            c2.font = black_bold if es_negrita else black_normal
            idx += 2
            
        # Totales de la fila
        c_t1 = ws.cell(row=row_num, column=idx, value=t_rem if t_rem > 0 else None)
        c_t2 = ws.cell(row=row_num, column=idx+1, value=t_rec if t_rec > 0 else None)
        
        c_t1.number_format = number_fmt
        c_t2.number_format = number_fmt
        c_t1.border = thin_border
        c_t2.border = thin_border
        c_t1.font = black_bold if es_negrita else black_normal
        c_t2.font = black_bold if es_negrita else black_normal
            
    # --- ESCRITURA DE DATOS ---
    current_row = 5
    totales_globales = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
    
    if tipo_reporte == 'proveedor':
        for prov in sorted(matrix.keys()):
            totales_prov = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
            
            # Calcular totales del proveedor
            for mat in matrix[prov]:
                for m in range(1,13):
                    totales_prov[m]['rem'] += matrix[prov][mat][m]['rem']
                    totales_prov[m]['rec'] += matrix[prov][mat][m]['rec']
                    totales_globales[m]['rem'] += matrix[prov][mat][m]['rem']
                    totales_globales[m]['rec'] += matrix[prov][mat][m]['rec']
            
            # Fila Proveedor (Bold)
            c_p = ws.cell(row=current_row, column=1, value=prov)
            c_p.font = black_bold
            c_p.alignment = left_align
            c_p.border = thin_border
            escribir_valores(current_row, totales_prov, True)
            current_row += 1
            
            # Filas Materiales (Indentadas y con agrupación)
            for mat in sorted(matrix[prov].keys()):
                ws.row_dimensions[current_row].outline_level = 1  # Agrupación [- / +]
                c_m = ws.cell(row=current_row, column=1, value=mat)
                c_m.font = black_normal
                c_m.alignment = left_indent # Sangría alineada a la derecha
                c_m.border = thin_border
                escribir_valores(current_row, matrix[prov][mat], False)
                current_row += 1
    else:
        for mat in sorted(matrix.keys()):
            for m in range(1,13):
                totales_globales[m]['rem'] += matrix[mat][m]['rem']
                totales_globales[m]['rec'] += matrix[mat][m]['rec']
            
            c_m = ws.cell(row=current_row, column=1, value=mat)
            c_m.font = black_normal
            c_m.alignment = left_align
            c_m.border = thin_border
            escribir_valores(current_row, matrix[mat], False)
            current_row += 1
            
    # --- FILA DE TOTAL GENERAL FINAL ---
    c_tg = ws.cell(row=current_row, column=1, value="Total general")
    c_tg.font = black_bold
    c_tg.alignment = left_align
    c_tg.border = thin_border
    escribir_valores(current_row, totales_globales, True)
    
    ws.column_dimensions['A'].width = 35
    
    return wb


def _build_pivot_excel_logistica(year, tipo_reporte):
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    
    # --- CONFIGURACIÓN AVANZADA DE EXCEL ---
    ws.sheet_properties.outlinePr.summaryBelow = False  # Para que el botón [-] quede arriba
    ws.freeze_panes = 'B5'  # Congelar encabezados y columna de nombres
    
    # --- ESTILOS VISUALES ---
    blue_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    gray_fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
    
    white_font = Font(color="FFFFFF", bold=True, size=10)
    black_bold_font = Font(color="000000", bold=True, size=10)
    normal_font = Font(color="000000", size=10)
    
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center")
    
    # Borde sutil para parecer tabla dinámica
    thin_border = Border(
        left=Side(style='thin', color='BFBFBF'),
        right=Side(style='thin', color='BFBFBF'),
        top=Side(style='thin', color='BFBFBF'),
        bottom=Side(style='thin', color='BFBFBF')
    )
    
    number_fmt = '#,##0.00'
    
    meses_nombres = {
        1: 'Ene', 2: 'Feb', 3: 'Mar', 4: 'Abr', 5: 'May', 6: 'Jun',
        7: 'Jul', 8: 'Ago', 9: 'Sep', 10: 'Oct', 11: 'Nov', 12: 'Dic'
    }
    
    # --- CONSULTA BASE ---
    qs = RegistroLogistico.objects.filter(
        fecha_carga__year=year,
        status__in=['TERMINADO', 'AUDITADO']
    )
    
    # --- CONSTRUCCIÓN DE MATRIZ DE DATOS ---
    matrix = {}
    
    if tipo_reporte == 'proveedor':
        ws.title = "Por Proveedor"
        datos = qs.values('transportista__nombre', 'material__nombre', 'fecha_carga__month').annotate(
            rem=Sum('toneladas_remisionadas'),
            rec=Sum('toneladas_recibidas')
        )
        for d in datos:
            prov = d['transportista__nombre'] or "SIN PROVEEDOR"
            mat = d['material__nombre'] or "SIN MATERIAL"
            mes = d['fecha_carga__month']
            rem = float(d['rem'] or 0)
            rec = float(d['rec'] or 0)
            
            if prov not in matrix: matrix[prov] = {}
            if mat not in matrix[prov]: matrix[prov][mat] = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
            
            matrix[prov][mat][mes]['rem'] += rem
            matrix[prov][mat][mes]['rec'] += rec
            
    elif tipo_reporte == 'material':
        ws.title = "Por Material"
        datos = qs.values('material__nombre', 'fecha_carga__month').annotate(
            rem=Sum('toneladas_remisionadas'),
            rec=Sum('toneladas_recibidas')
        )
        for d in datos:
            mat = d['material__nombre'] or "SIN MATERIAL"
            mes = d['fecha_carga__month']
            rem = float(d['rem'] or 0)
            rec = float(d['rec'] or 0)
            
            if mat not in matrix: matrix[mat] = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
            
            matrix[mat][mes]['rem'] += rem
            matrix[mat][mes]['rec'] += rec

    # --- ALTURA DE FILAS PARA QUE EL TEXTO QUEPA PERFECTAMENTE ---
    ws.row_dimensions[3].height = 25
    ws.row_dimensions[4].height = 55

    # --- ENCABEZADOS (Filas 3 y 4) ---
    c_val = ws.cell(row=3, column=1, value="Valores")
    c_val.font = black_bold_font
    c_val.fill = gray_fill
    c_val.alignment = center_align
    c_val.border = thin_border
    
    c_a4 = ws.cell(row=4, column=1, value="PROVEEDOR (TRANSPORTISTA)" if tipo_reporte == 'proveedor' else "Material")
    c_a4.fill = gray_fill
    c_a4.font = black_bold_font
    c_a4.alignment = center_align
    c_a4.border = thin_border
    
    col_idx = 2
    for m in range(1, 13):
        # Header Mes (Azul)
        ws.merge_cells(start_row=3, start_column=col_idx, end_row=3, end_column=col_idx+1)
        c_mes = ws.cell(row=3, column=col_idx, value=meses_nombres[m])
        c_mes.fill = blue_fill
        c_mes.font = white_font
        c_mes.alignment = center_align
        ws.cell(row=3, column=col_idx+1).border = thin_border # Borde de la celda mergeada
        c_mes.border = thin_border
        
        # Sub-headers (Gris)
        c_rem = ws.cell(row=4, column=col_idx, value="Suma de TONELADAS REMISIONADAS")
        c_rem.fill = gray_fill
        c_rem.alignment = center_align
        c_rem.font = normal_font
        c_rem.border = thin_border
        
        c_rec = ws.cell(row=4, column=col_idx+1, value="Suma de TONELADAS RECIBIDAS TERNIUM")
        c_rec.fill = gray_fill
        c_rec.alignment = center_align
        c_rec.font = normal_font
        c_rec.border = thin_border
        
        ws.column_dimensions[get_column_letter(col_idx)].width = 16
        ws.column_dimensions[get_column_letter(col_idx+1)].width = 16
        
        col_idx += 2
        
    # Total General Header
    ws.merge_cells(start_row=3, start_column=col_idx, end_row=3, end_column=col_idx+1)
    c_gt = ws.cell(row=3, column=col_idx, value="Total general")
    c_gt.fill = blue_fill
    c_gt.font = white_font
    c_gt.alignment = center_align
    c_gt.border = thin_border
    ws.cell(row=3, column=col_idx+1).border = thin_border
    
    c_rem_gt = ws.cell(row=4, column=col_idx, value="Suma de TONELADAS REMISIONADAS")
    c_rem_gt.fill = gray_fill
    c_rem_gt.alignment = center_align
    c_rem_gt.font = normal_font
    c_rem_gt.border = thin_border
    
    c_rec_gt = ws.cell(row=4, column=col_idx+1, value="Suma de TONELADAS RECIBIDAS TERNIUM")
    c_rec_gt.fill = gray_fill
    c_rec_gt.alignment = center_align
    c_rec_gt.font = normal_font
    c_rec_gt.border = thin_border
    
    ws.column_dimensions[get_column_letter(col_idx)].width = 16
    ws.column_dimensions[get_column_letter(col_idx+1)].width = 16
    
    # --- FUNCIÓN AUXILIAR PARA ESCRIBIR VALORES CON FORMATO Y BORDES ---
    def escribir_valores(row_num, dict_meses, es_negrita):
        idx = 2
        t_rem = 0.0
        t_rec = 0.0
        for m in range(1, 13):
            v_rem = dict_meses[m]['rem']
            v_rec = dict_meses[m]['rec']
            t_rem += v_rem
            t_rec += v_rec
            
            c1 = ws.cell(row=row_num, column=idx, value=v_rem if v_rem > 0 else "")
            c2 = ws.cell(row=row_num, column=idx+1, value=v_rec if v_rec > 0 else "")
            
            c1.number_format = number_fmt
            c2.number_format = number_fmt
            c1.border = thin_border
            c2.border = thin_border
            
            if es_negrita:
                c1.font = black_bold_font; c2.font = black_bold_font
            else:
                c1.font = normal_font; c2.font = normal_font
            idx += 2
            
        # Totales de la fila
        c_t1 = ws.cell(row=row_num, column=idx, value=t_rem if t_rem > 0 else "")
        c_t2 = ws.cell(row=row_num, column=idx+1, value=t_rec if t_rec > 0 else "")
        
        c_t1.number_format = number_fmt
        c_t2.number_format = number_fmt
        c_t1.border = thin_border
        c_t2.border = thin_border
        
        if es_negrita:
            c_t1.font = black_bold_font; c_t2.font = black_bold_font
        else:
            c_t1.font = normal_font; c_t2.font = normal_font
            
    # --- ESCRITURA DE DATOS ---
    current_row = 5
    totales_globales = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
    
    if tipo_reporte == 'proveedor':
        for prov in sorted(matrix.keys()):
            totales_prov = {m: {'rem':0.0, 'rec':0.0} for m in range(1,13)}
            
            # Calcular totales del proveedor
            for mat in matrix[prov]:
                for m in range(1,13):
                    totales_prov[m]['rem'] += matrix[prov][mat][m]['rem']
                    totales_prov[m]['rec'] += matrix[prov][mat][m]['rec']
                    totales_globales[m]['rem'] += matrix[prov][mat][m]['rem']
                    totales_globales[m]['rec'] += matrix[prov][mat][m]['rec']
            
            # Fila Proveedor (Bold)
            c_p = ws.cell(row=current_row, column=1, value=f"{prov}")
            c_p.font = black_bold_font
            c_p.alignment = left_align
            c_p.border = thin_border
            escribir_valores(current_row, totales_prov, True)
            current_row += 1
            
            # Filas Materiales (Indentadas y con agrupación)
            for mat in sorted(matrix[prov].keys()):
                ws.row_dimensions[current_row].outline_level = 1  # <-- Agrupación mágica de Excel
                c_m = ws.cell(row=current_row, column=1, value=f"   {mat}")
                c_m.font = normal_font
                c_m.alignment = left_align
                c_m.border = thin_border
                escribir_valores(current_row, matrix[prov][mat], False)
                current_row += 1
    else:
        for mat in sorted(matrix.keys()):
            for m in range(1,13):
                totales_globales[m]['rem'] += matrix[mat][m]['rem']
                totales_globales[m]['rec'] += matrix[mat][m]['rec']
            
            c_m = ws.cell(row=current_row, column=1, value=mat)
            c_m.font = normal_font
            c_m.alignment = left_align
            c_m.border = thin_border
            escribir_valores(current_row, matrix[mat], False)
            current_row += 1
            
    # --- FILA DE TOTAL GENERAL FINAL ---
    c_tg = ws.cell(row=current_row, column=1, value="Total general")
    c_tg.font = black_bold_font
    c_tg.alignment = left_align
    c_tg.border = thin_border
    escribir_valores(current_row, totales_globales, True)
    
    ws.column_dimensions['A'].width = 35
    
    return wb

# =========================================================================
# VISTAS INDIVIDUALES ACTUALIZADAS
# =========================================================================

@login_required
def export_logistica_reporte_proveedor(request):
    try:
        year = int(request.GET.get('year', datetime.date.today().year))
    except ValueError:
        year = datetime.date.today().year

    wb = _build_pivot_excel_logistica(year, 'proveedor')
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Logistica_Por_Proveedor_{year}.xlsx"'
    wb.save(response)
    
    return response

@login_required
def export_logistica_reporte_material(request):
    try:
        year = int(request.GET.get('year', datetime.date.today().year))
    except ValueError:
        year = datetime.date.today().year

    wb = _build_pivot_excel_logistica(year, 'material')
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Logistica_Por_Material_{year}.xlsx"'
    wb.save(response)
    
    return response

import os
import io
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

@login_required
def descargar_reporte_destruccion(request, remision_id):
    try:
        remision = get_object_or_404(Remision, pk=remision_id)
        
        origen_nombre = remision.origen.nombre if remision.origen else "S/N"
        operador_nombre = remision.operador.nombre if remision.operador else (remision.operador_manual or "S/N")
        
        unidad_nombre = remision.unidad.internal_id if remision.unidad else (remision.unidad_manual or "S/N")
        placas_unidad = getattr(remision.unidad, 'license_plate', getattr(remision.unidad, 'placas', "S/N")) if remision.unidad else (getattr(remision, 'placas_unidad_manual', "S/N") or "S/N")
        
        contenedor_nombre = remision.contenedor.nombre if remision.contenedor else (remision.contenedor_manual or "S/N")
        placas_cont = getattr(remision.contenedor, 'placas', "S/N") if remision.contenedor else (getattr(remision, 'placas_contenedor_manual', "S/N") or "S/N")

        template_path = os.path.join(settings.BASE_DIR, 'templates_word', 'plantilla_destruccion.docx')
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Mm
        doc = DocxTemplate(template_path)

        def get_inline_image(image_field, width_mm=55):
            if not image_field or not image_field.name:
                return ""
            try:
                s3_client = boto3.client('s3', aws_access_key_id=settings.AWS_ACCESS_KEY_ID, aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY, region_name=settings.AWS_S3_REGION_NAME)
                s3_key = f"{settings.AWS_MEDIA_LOCATION}/{image_field.name}"
                file_stream = io.BytesIO()
                s3_client.download_fileobj(settings.AWS_STORAGE_BUCKET_NAME, s3_key, file_stream)
                file_stream.seek(0)
                return InlineImage(doc, file_stream, width=Mm(width_mm))
            except Exception as e:
                print(f"Error cargando imagen para Word: {e}")
                return ""

        # Verificamos si existe un Material 2 seleccionado
        tiene_mat2 = bool(remision.destruccion_material_2)

        # --- DICCIONARIO DE DATOS PARA EL ARCHIVO DE WORD ---
        context = {
            'fecha': remision.fecha_destruccion.strftime('%d/%m/%Y') if remision.fecha_destruccion else (remision.fecha.strftime('%d/%m/%Y') if remision.fecha else ""),
            'folio': remision.remision,
            'origen': origen_nombre,
            'operador': operador_nombre,
            
            # --- FILA 1 ---
            'unidad': unidad_nombre,
            'placas_unidad': placas_unidad,
            'caja': contenedor_nombre,
            'placas_caja': placas_cont,
            'hora_inicio': remision.hora_entrada.strftime('%H:%M') if remision.hora_entrada else "",
            'hora_fin': remision.hora_salida.strftime('%H:%M') if remision.hora_salida else "",
            'material1': remision.destruccion_material_1 or "",
            'peso': f"{remision.destruccion_peso_1}" if remision.destruccion_peso_1 else "",
            
            # --- FILA 2 (Solo se llenan si hay material 2) ---
            'unidad2': unidad_nombre if tiene_mat2 else "",
            'placas_unidad2': placas_unidad if tiene_mat2 else "",
            'caja2': contenedor_nombre if tiene_mat2 else "",
            'placas_caja2': placas_cont if tiene_mat2 else "",
            'hora_inicio2': remision.hora_entrada.strftime('%H:%M') if (remision.hora_entrada and tiene_mat2) else "",
            'hora_fin2': remision.hora_salida.strftime('%H:%M') if (remision.hora_salida and tiene_mat2) else "",
            'material2': remision.destruccion_material_2 or "",
            'peso2': f"{remision.destruccion_peso_2}" if remision.destruccion_peso_2 else "",
            
            'suma': {'peso': remision.peso_bascula if remision.peso_bascula else 0}, 
            'Comentarios': remision.comentarios_destruccion or "",

            'foto1': get_inline_image(remision.foto_ingreso, 55),
            'foto1_2': get_inline_image(remision.foto_ingreso_2, 55),
            'foto2': get_inline_image(remision.foto_vertido, 55),
            'foto2_1': get_inline_image(remision.foto_vertido_2, 55),
            'foto3': get_inline_image(remision.foto_destruccion, 55),
            'foto3_1': get_inline_image(remision.foto_destruccion_2, 55),
        }

        doc.render(context)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Reporte_Destruccion_{remision.remision}.docx"'
        return response

    except Exception as e:
        messages.error(request, f"Ocurrió un error al generar el Word: {str(e)}")
        return redirect('remision_lista')

from .models import ConfiguracionManifiesto
from .forms import ConfiguracionManifiestoForm

# --- NUEVA VISTA: PANTALLA DE CONFIGURACIÓN ---
@login_required
def asignar_manifiesto(request):
    configuraciones = ConfiguracionManifiesto.objects.all().order_by('-creado_en')
    
    if request.method == 'POST':
        # Eliminar una configuración existente
        if 'eliminar' in request.POST:
            config_id = request.POST.get('config_id')
            ConfiguracionManifiesto.objects.filter(id=config_id).delete()
            messages.success(request, "Configuración eliminada correctamente.")
            return redirect('asignar_manifiesto')
            
        # Agregar una nueva configuración
        else:
            form = ConfiguracionManifiestoForm(request.POST)
            if form.is_valid():
                form.save()
                messages.success(request, "Configuración de manifiesto agregada exitosamente.")
                return redirect('asignar_manifiesto')
            else:
                messages.error(request, "Error: Esta combinación de Origen y Material ya existe.")
    else:
        form = ConfiguracionManifiestoForm()

    # Pre-filtrar solo Lugares que son ORIGEN
    form.fields['origen'].queryset = Lugar.objects.filter(tipo__in=['ORIGEN', 'AMBOS']).order_by('nombre')

    return render(request, 'ternium/asignar_manifiesto.html', {
        'form': form, 
        'configuraciones': configuraciones
    })


from .models import ControlManifiestoTrane
from .forms import ControlManifiestoTraneForm
def is_cordinador_trane(user):
    return user.is_authenticated and (user.is_superuser or user.groups.filter(name='Cordinadores_trane').exists())
@login_required
@user_passes_test(is_cordinador_trane, login_url='home')
def lista_control_trane(request):
    # 1. Leer los parámetros de fecha desde la URL (si el usuario usó el filtro)
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')

    # 2. Consultar todos los registros por defecto (ordenados por fecha descendente)
    registros = ControlManifiestoTrane.objects.all().order_by('-fecha_captura', '-id')

    # 3. Aplicar los filtros si el usuario seleccionó fechas
    if fecha_inicio:
        registros = registros.filter(fecha_captura__gte=fecha_inicio)
    if fecha_fin:
        registros = registros.filter(fecha_captura__lte=fecha_fin)

    # 4. Enviar los datos y las variables de fecha a la plantilla
    return render(request, 'ternium/control_trane_lista.html', {
        'registros': registros,
        'fecha_inicio': fecha_inicio,
        'fecha_fin': fecha_fin,
    })

@login_required
@user_passes_test(is_cordinador_trane, login_url='home')
def crear_control_trane(request):
    valores_manuales = {
        'operador': '', 'unidad': '', 'placas_unidad': '',
        'contenedor': '', 'placas_contenedor': ''
    }
    
    if request.method == 'POST':
        # Capturar textos manuales
        valores_manuales['operador'] = request.POST.get('operador_texto', '').strip().upper()
        valores_manuales['unidad'] = request.POST.get('unidad_texto', '').strip().upper()
        valores_manuales['placas_unidad'] = request.POST.get('placas_unidad_texto', '').strip().upper()
        valores_manuales['contenedor'] = request.POST.get('contenedor_texto', '').strip().upper()
        valores_manuales['placas_contenedor'] = request.POST.get('placas_contenedor_texto', '').strip().upper()

        form = ControlManifiestoTraneForm(request.POST, request.FILES)
        if form.is_valid():
            obj = form.save(commit=False)
            
            # Asignar campos manuales y limpiar FKs si se usó texto manual
            if valores_manuales['operador']: obj.operador_manual = valores_manuales['operador']; obj.operador = None
            if valores_manuales['unidad']: obj.unidad_manual = valores_manuales['unidad']; obj.placas_unidad_manual = valores_manuales['placas_unidad']; obj.unidad = None
            if valores_manuales['contenedor']: obj.contenedor_manual = valores_manuales['contenedor']; obj.placas_contenedor_manual = valores_manuales['placas_contenedor']; obj.contenedor = None
            
            # Subir Manifiesto a S3
            if 'manifiesto' in request.FILES:
                archivo = request.FILES['manifiesto']
                s3_path = f"trane/manifiestos/{timezone.now().timestamp()}_{archivo.name.replace(' ', '_')}"
                ruta = _subir_archivo_a_s3(archivo, s3_path)
                if ruta: obj.manifiesto = ruta
                
            # Subir Documento Trane a S3
            if 'documento_trane' in request.FILES:
                archivo = request.FILES['documento_trane']
                s3_path = f"trane/documentos/{timezone.now().timestamp()}_{archivo.name.replace(' ', '_')}"
                ruta = _subir_archivo_a_s3(archivo, s3_path)
                if ruta: obj.documento_trane = ruta

            obj.save()
            messages.success(request, 'Registro creado exitosamente.')
            return redirect('lista_control_trane')
    else:
        form = ControlManifiestoTraneForm()
    
    return render(request, 'ternium/control_trane_form.html', {
        'form': form, 'titulo': 'Nuevo Registro Documental TRANE', 
        'valores_manuales': valores_manuales, 'is_editing': False
    })

@login_required
@user_passes_test(is_cordinador_trane, login_url='home')
def editar_control_trane(request, pk):
    obj_original = get_object_or_404(ControlManifiestoTrane, pk=pk)
    
    valores_manuales = {
        'operador': obj_original.operador_manual or '',
        'unidad': obj_original.unidad_manual or '',
        'placas_unidad': obj_original.placas_unidad_manual or '',
        'contenedor': obj_original.contenedor_manual or '',
        'placas_contenedor': obj_original.placas_contenedor_manual or ''
    }

    if request.method == 'POST':
        valores_manuales['operador'] = request.POST.get('operador_texto', '').strip().upper()
        valores_manuales['unidad'] = request.POST.get('unidad_texto', '').strip().upper()
        valores_manuales['placas_unidad'] = request.POST.get('placas_unidad_texto', '').strip().upper()
        valores_manuales['contenedor'] = request.POST.get('contenedor_texto', '').strip().upper()
        valores_manuales['placas_contenedor'] = request.POST.get('placas_contenedor_texto', '').strip().upper()

        form = ControlManifiestoTraneForm(request.POST, request.FILES, instance=obj_original)
        if form.is_valid():
            obj = form.save(commit=False)
            
            # Asignar campos manuales
            if valores_manuales['operador']: obj.operador_manual = valores_manuales['operador']; obj.operador = None
            else: obj.operador_manual = None
            
            if valores_manuales['unidad']: obj.unidad_manual = valores_manuales['unidad']; obj.placas_unidad_manual = valores_manuales['placas_unidad']; obj.unidad = None
            else: obj.unidad_manual = None; obj.placas_unidad_manual = None
            
            if valores_manuales['contenedor']: obj.contenedor_manual = valores_manuales['contenedor']; obj.placas_contenedor_manual = valores_manuales['placas_contenedor']; obj.contenedor = None
            else: obj.contenedor_manual = None; obj.placas_contenedor_manual = None

            # Actualizar Manifiesto en S3
            if 'manifiesto' in request.FILES:
                if obj_original.manifiesto and hasattr(obj_original.manifiesto, 'name'):
                    _eliminar_archivo_de_s3(obj_original.manifiesto.name)
                archivo = request.FILES['manifiesto']
                s3_path = f"trane/manifiestos/{timezone.now().timestamp()}_{archivo.name.replace(' ', '_')}"
                ruta = _subir_archivo_a_s3(archivo, s3_path)
                if ruta: obj.manifiesto = ruta
                
            # Actualizar Documento Trane en S3
            if 'documento_trane' in request.FILES:
                if obj_original.documento_trane and hasattr(obj_original.documento_trane, 'name'):
                    _eliminar_archivo_de_s3(obj_original.documento_trane.name)
                archivo = request.FILES['documento_trane']
                s3_path = f"trane/documentos/{timezone.now().timestamp()}_{archivo.name.replace(' ', '_')}"
                ruta = _subir_archivo_a_s3(archivo, s3_path)
                if ruta: obj.documento_trane = ruta

            obj.save()
            messages.success(request, 'Registro actualizado exitosamente.')
            return redirect('lista_control_trane')
    else:
        form = ControlManifiestoTraneForm(instance=obj_original)
    
    return render(request, 'ternium/control_trane_form.html', {
        'form': form, 'titulo': 'Editar Registro TRANE', 
        'valores_manuales': valores_manuales, 'is_editing': True, 'registro': obj_original
    })

@login_required
@require_POST
@user_passes_test(is_cordinador_trane, login_url='home')
def eliminar_control_trane(request, pk):
    obj = get_object_or_404(ControlManifiestoTrane, pk=pk)
    # Eliminar archivos de S3
    if obj.manifiesto and hasattr(obj.manifiesto, 'name'):
        _eliminar_archivo_de_s3(obj.manifiesto.name)
    if obj.documento_trane and hasattr(obj.documento_trane, 'name'):
        _eliminar_archivo_de_s3(obj.documento_trane.name)
        
    obj.delete()
    messages.success(request, 'Registro eliminado exitosamente.')
    return redirect('lista_control_trane')

@login_required
@user_passes_test(is_cordinador_trane, login_url='home')
def generar_trane_desde_remision(request, remision_id):
    remision = get_object_or_404(Remision, id=remision_id)
    
    # Doble validación de seguridad: si por error llega aquí una ya vinculada, lo bloquea
    if hasattr(remision, 'controlmanifiestotrane'):
        messages.warning(request, 'Esta operación ya fue vinculada anteriormente.')
        return redirect('lista_control_trane')

    material_obj = None
    peso_total_carga = 0 # <--- NUEVO: Variable para sumar el peso

    detalle_primero = remision.detalles.first() # Ajusta 'detalles' según como se llame tu related_name en el formset
    if detalle_primero:
        material_obj = detalle_primero.material
        # <--- NUEVO: Sumar el peso de la carga de todos los detalles
        for d in remision.detalles.all():
            if d.peso_ld:
                peso_total_carga += float(d.peso_ld)

    # Se copian los datos y se asigna el folio de la remisión original
    nuevo_control = ControlManifiestoTrane.objects.create(
        remision_vinculada=remision, # AQUÍ SE CREA EL VÍNCULO PARA QUE DESAPAREZCA DE LA LISTA
        fecha_captura=remision.fecha,
        operador=remision.operador,
        operador_manual=remision.operador_manual if hasattr(remision, 'operador_manual') else None,
        linea_transporte=remision.linea_transporte,
        unidad=remision.unidad,
        unidad_manual=remision.unidad_manual if hasattr(remision, 'unidad_manual') else None,
        placas_unidad_manual=remision.placas_unidad_manual if hasattr(remision, 'placas_unidad_manual') else None,
        contenedor=remision.contenedor,
        contenedor_manual=remision.contenedor_manual if hasattr(remision, 'contenedor_manual') else None,
        placas_contenedor_manual=remision.placas_contenedor_manual if hasattr(remision, 'placas_contenedor_manual') else None,
        origen=remision.origen,
        destino=remision.destino,
        material=material_obj,
        cantidad_kg=peso_total_carga, # <--- NUEVO: Se asigna el peso total sumado
        folio=remision.remision, # Copiamos el folio original (ej. FOL-0123)
        manifiesto=remision.manifiesto if hasattr(remision, 'manifiesto') and remision.manifiesto else None
    )

    messages.success(request, 'Información generada exitosamente. Adjunta los documentos faltantes.')
    return redirect('editar_control_trane', pk=nuevo_control.id)

from django.db import transaction
from .models import DetalleRemision # Asegúrate de tener esto importado arriba

@login_required
@user_passes_test(is_cordinador_trane, login_url='home')
def importar_trane_a_remision(request):
    if request.method == 'POST':
        trane_id = request.POST.get('trane_id')
        empresa_id = request.POST.get('empresa_id')

        trane = get_object_or_404(ControlManifiestoTrane, id=trane_id)
        empresa = get_object_or_404(Empresa, id=empresa_id)

        try:
            with transaction.atomic():
                # 1. Calculamos el Folio
                siguiente_folio = ""
                if empresa.prefijo:
                    siguiente_folio = calcular_siguiente_folio(empresa.prefijo)

                # 2. Crear la remisión principal en blanco
                nueva_remision = Remision(
                    remision=siguiente_folio, 
                    empresa=empresa,
                    fecha=trane.fecha_captura or timezone.now().date(),
                    linea_transporte=trane.linea_transporte,
                    operador=trane.operador,
                    operador_manual=trane.operador_manual,
                    unidad=trane.unidad,
                    unidad_manual=trane.unidad_manual,
                    placas_unidad_manual=trane.placas_unidad_manual,
                    contenedor=trane.contenedor,
                    contenedor_manual=trane.contenedor_manual,
                    placas_contenedor_manual=trane.placas_contenedor_manual,
                    origen=trane.origen,
                    destino=trane.destino,
                    manifiesto=trane.manifiesto,
                    status='CREADO'
                )
                nueva_remision.save()

                # 3. MÉTODO DIRECTO PARA PASAR MATERIAL Y CARGA
                detalles_creados = False
                peso_capturado = 0
                nombre_material = ""

                if trane.material:
                    peso_capturado = trane.cantidad_kg if trane.cantidad_kg else 0
                    nombre_material = trane.material.nombre
                    
                    # Usamos objects.create de forma limpia (solo lo indispensable)
                    DetalleRemision.objects.create(
                        remision=nueva_remision,
                        material=trane.material,
                        peso_ld=peso_capturado,
                        peso_dlv=peso_capturado
                    )
                    detalles_creados = True

                # 4. Vincular el TRANE para que desaparezca del modal
                trane.remision_vinculada = nueva_remision
                trane.save()

            # 5. MENSAJES DE DIAGNÓSTICO PARA SABER QUÉ PASÓ
            if detalles_creados:
                messages.success(request, f'Remisión {siguiente_folio} generada. Se copiaron con éxito: {peso_capturado} Kg de {nombre_material}.')
            else:
                messages.warning(request, f'Remisión {siguiente_folio} generada. ATENCIÓN: El registro de TRANE original no tenía Material asignado.')
                
            return redirect('remision_editar', pk=nueva_remision.id)

        except Exception as e:
            messages.error(request, f'Error crítico al generar remisión: {str(e)}')
            return redirect('remision_lista')
        

# ============================================================
# AGREGAR ESTO A TU views.py (Django)
# ============================================================
# URL requerida en urls.py:
#   path('api/dashboard-operacion/', views.api_dashboard_operacion, name='api_dashboard_operacion'),
# ============================================================

from django.http import JsonResponse
from django.db.models import Count, Sum, F, Q, FloatField
from django.db.models.functions import TruncMonth, Coalesce
from django.contrib.auth.decorators import login_required


# ============================================================
# AGREGAR ESTA FUNCIÓN A TU views.py (Django)
# ============================================================
# URL ya registrada: path('api/dashboard-operacion/', views.api_dashboard_operacion, ...)
# ============================================================

from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from django.db.models import Sum, Count, FloatField, Q 
from django.db.models.functions import Coalesce, TruncMonth

# (IMPORTANTE: MANTÉN TUS IMPORTACIONES DE MODELOS HASTA ARRIBA. Ej: from tu_app.models import Empresa...)

#@login_required
def api_dashboard_operacion(request):
    operacion = request.GET.get('operacion', 'MTY')
    year = int(request.GET.get('year', timezone.now().year))
    meses_str = request.GET.get('meses', '')
    origen_str = request.GET.get('origen', 'TODOS')
    material_str = request.GET.get('material', 'TODOS')

    empresa = Empresa.objects.filter(prefijo__iexact=operacion).first()

    if not empresa:
        return JsonResponse({'error': f'Operación "{operacion}" no encontrada.'}, status=404)

    # =========================
    # STATUS COUNTS Y BASE SIN FILTRO DE MATERIAL
    # =========================
    all_qs = Remision.objects.filter(empresa=empresa, fecha__year=year)

    if meses_str:
        meses_list2 = [int(m) for m in meses_str.split(',') if m.isdigit()]
        if meses_list2:
            all_qs = all_qs.filter(fecha__month__in=meses_list2)

    sc = dict(
        all_qs.values('status')
        .annotate(c=Count('id'))
        .values_list('status', 'c')
    )

    # =========================
    # FILTRO PRINCIPAL (Para las gráficas)
    # =========================
    qs = all_qs.filter(status__in=['TERMINADO', 'AUDITADO'])

    if origen_str and origen_str != 'TODOS':
        qs = qs.filter(origen__nombre__iexact=origen_str)

    if material_str and material_str != 'TODOS':
        qs = qs.filter(detalles__material__nombre__iexact=material_str)

    # =========================
    # KPIs
    # =========================
    agg = qs.aggregate(
        total_carga=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        total_descarga=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField()),
        total_remisiones=Count('id', distinct=True),
    )

    tc = agg['total_carga']
    td = agg['total_descarga']

    merma_pct = round(((tc - td) / tc) * 100, 1) if tc > 0 else 0

    top_o = qs.values('origen__nombre').annotate(cnt=Count('id')).order_by('-cnt').first()
    top_d = qs.values('destino__nombre').annotate(cnt=Count('id')).order_by('-cnt').first()

    # =========================
    # RENDIMIENTO
    # =========================
    now = timezone.now()
    cur = qs.filter(fecha__month=now.month).aggregate(
        t=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField())
    )['t']

    prev_month = now.month - 1 if now.month > 1 else 12
    prev_year = year if now.month > 1 else year - 1

    prev = Remision.objects.filter(
        empresa=empresa,
        status__in=['TERMINADO', 'AUDITADO'],
        fecha__year=prev_year,
        fecha__month=prev_month
    ).aggregate(
        t=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField())
    )['t']

    if prev and prev > 0:
        cambio = round(((cur - prev) / prev) * 100, 1)
        rend = f"+{cambio}%" if cambio >= 0 else f"{cambio}%"
    else:
        rend = "N/A"

    MESES = ['Ene','Feb','Mar','Abr','May','Jun','Jul','Ago','Sep','Oct','Nov','Dic']
    COLORS = ['#0ea5e9','#10b981','#f59e0b','#6366f1','#8b5cf6','#ef4444','#ec4899','#14b8a6']

    # =========================
    # TIMELINE
    # =========================
    timeline = qs.annotate(
        mes=TruncMonth('fecha')
    ).values('mes').annotate(
        carga=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        descarga=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField()),
    ).order_by('mes')

    # =========================
    # RANK GENERICO 
    # =========================
    def rank(name_field, limit=10):
        return [
            {'name': r[name_field], 'value': r['value']}
            for r in qs.values(name_field)
            .exclude(**{name_field: None})
            .exclude(**{name_field: ''})
            .annotate(value=Count('id', distinct=True))
            .order_by('-value')[:limit]
        ]
    
    def rank_filtros(name_field, limit=15):
        return [
            {'name': r[name_field], 'value': r['value']}
            for r in all_qs.filter(status__in=['TERMINADO', 'AUDITADO']).values(name_field)
            .exclude(**{name_field: None})
            .exclude(**{name_field: ''})
            .annotate(value=Count('id', distinct=True))
            .order_by('-value')[:limit]
        ]

    # =========================
    # OPERADORES CON MERMA Y DESGLOSE
    # =========================
    operadores_merma = []
    op_data = (
        qs.filter(operador__isnull=False)
        .values('operador__nombre')
        .annotate(
            viajes=Count('id', distinct=True),
            total_carga_op=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
            total_descarga_op=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField()),
        )
        .order_by('-viajes')[:10]
    )
    
    for op_item in op_data:
        op_name = op_item['operador__nombre'] or 'Sin Nombre'
        c_total = op_item['total_carga_op']
        d_total = op_item['total_descarga_op']
        merma_total = round(((c_total - d_total) / c_total) * 100, 1) if c_total > 0 else 0
        
        detalles_list = []
        if op_name != 'Sin Nombre':
            rems_op = qs.filter(operador__nombre=op_name).annotate(
                c=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
                d=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField())
            ).values('remision', 'fecha', 'origen__nombre', 'destino__nombre', 'c', 'd', 'comentario').order_by('-fecha')[:30]
            
            for r in rems_op:
                c = r['c']
                d = r['d']
                dif = d - c 
                m_pct = round((abs(dif) / c * 100), 1) if c > 0 and dif < 0 else 0
                detalles_list.append({
                    'remision': r['remision'],
                    'fecha': r['fecha'].strftime('%Y-%m-%d') if r['fecha'] else '',
                    'origen': r['origen__nombre'] or '-',
                    'destino': r['destino__nombre'] or '-',
                    'carga': round(c, 2),
                    'descarga': round(d, 2),
                    'diferencia': round(dif, 2),
                    'merma': m_pct,
                    'comentario': r.get('comentario', '') or ''
                })

        operadores_merma.append({
            'name': op_name,
            'value': op_item['viajes'],
            'carga': round(c_total, 2),
            'descarga': round(d_total, 2),
            'merma': merma_total,
            'detalles': detalles_list
        })

    # =========================
    # MATERIALES CON DESGLOSE
    # =========================
    mat_data = qs.values('detalles__material__nombre').annotate(
        carga=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField()),
        descarga=Coalesce(Sum('detalles__peso_dlv'), 0.0, output_field=FloatField()),
    ).order_by('-carga')
    
    mat_top8_raw = list(mat_data[:8])
    materiales_final = []

    for i, m in enumerate(mat_top8_raw):
        mat_name = m['detalles__material__nombre'] or 'Sin Material'
        
        detalles_mat = []
        if mat_name != 'Sin Material':
            rems_mat = qs.filter(detalles__material__nombre=mat_name).annotate(
                c=Coalesce(Sum('detalles__peso_ld', filter=Q(detalles__material__nombre=mat_name)), 0.0, output_field=FloatField()),
                d=Coalesce(Sum('detalles__peso_dlv', filter=Q(detalles__material__nombre=mat_name)), 0.0, output_field=FloatField())
            ).values('remision', 'origen__nombre', 'destino__nombre', 'c', 'd').order_by('-id')[:30]

            for r in rems_mat:
                c = r['c']
                d = r['d']
                dif = d - c
                m_pct = round((abs(dif) / c * 100), 1) if c > 0 and dif < 0 else 0
                detalles_mat.append({
                    'remision': r['remision'],
                    'origen': r['origen__nombre'] or '-',
                    'destino': r['destino__nombre'] or '-',
                    'material': mat_name,
                    'carga': round(c, 2),
                    'descarga': round(d, 2),
                    'diferencia': round(dif, 2),
                    'merma': m_pct,
                })

        materiales_final.append({
            'name': mat_name,
            'carga': round(m['carga'], 2),
            'descarga': round(m['descarga'], 2),
            'color': COLORS[i % len(COLORS)],
            'detalles': detalles_mat
        })

    # =========================
    # LÍNEAS DE TRANSPORTE (Todas)
    # ¡AQUÍ ESTABA EL ERROR 500! YA ESTÁ CORREGIDO (linea_transporte__nombre)
    # =========================
    lineas_transporte = [
        {
            'name': r['linea_transporte__nombre'] or 'Sin Línea',
            'value': r['value']
        }
        for r in qs.filter(linea_transporte__isnull=False)
        .values('linea_transporte__nombre')
        .annotate(value=Count('id', distinct=True))
        .exclude(linea_transporte__nombre='')
        .order_by('-value')
    ]

    # =========================
    # REMISIONES PENDIENTES/CANCELADAS
    # =========================
    def lista_remisiones_por_status(status_val, limit=50):
        return list(
            all_qs.filter(status=status_val)
            .values('id', 'remision', 'fecha', 'origen__nombre', 'destino__nombre')
            .order_by('-fecha')[:limit]
        )

    pendientes_list = lista_remisiones_por_status('PENDIENTE')
    cancelados_list = lista_remisiones_por_status('CANCELADO')

    for item in pendientes_list + cancelados_list:
        if item.get('fecha'):
            item['fecha'] = item['fecha'].isoformat()

    # =========================
    # RESPONSE FINAL
    # =========================
    return JsonResponse({
        'kpis': {
            'totalKilos': round(td, 2),
            'totalCarga': round(tc, 2),
            'totalDescarga': round(td, 2),
            'remisiones': agg['total_remisiones'],
            'topOrigen': top_o['origen__nombre'] if top_o else 'N/A',
            'topDestino': top_d['destino__nombre'] if top_d else 'N/A',
            'rendimientoMensual': rend,
            'mermaPercent': merma_pct,
            'operadores': qs.values('operador').exclude(operador=None).distinct().count(),
            'unidades': qs.values('unidad').exclude(unidad=None).distinct().count(),
            'terminados': sc.get('TERMINADO', 0),
            'auditados': sc.get('AUDITADO', 0),
            'cancelados': sc.get('CANCELADO', 0),
            'pendientes': sc.get('PENDIENTE', 0),
        },
        'kilosPorMes': [
            {
                'mes': MESES[e['mes'].month - 1],
                'carga': round(e['carga'], 2),
                'descarga': round(e['descarga'], 2)
            }
            for e in timeline if e['mes']
        ],
        'materiales': materiales_final,
        'operadores': operadores_merma,
        'unidades': lineas_transporte, 
        'destinos': rank('destino__nombre'),
        'origenesChart': rank('origen__nombre'),
        'origenes': rank_filtros('origen__nombre'), 
        'materialesFiltro': rank_filtros('detalles__material__nombre'),
        'pendientesList': pendientes_list,
        'canceladosList': cancelados_list,
    })
    
"""
=============================================================================
API VIEW PARA REMISIONES - AGREGAR AL FINAL DE views.py (ternium/views.py)
=============================================================================
CORRECCIÓN: No usar @login_required porque redirige a /login/ (HTML).
Se valida manualmente request.user.is_authenticated → JSON 401.

URL a agregar en ternium/urls.py:
    path('api/remisiones-lista/', views.api_remisiones_lista, name='api_remisiones_lista'),
"""

# ===================== PEGAR ESTO AL FINAL DE views.py =====================

from django.views.decorators.csrf import csrf_exempt

@csrf_exempt
def api_remisiones_lista(request):
    """
    API JSON para remisiones. Misma lógica de filtros y permisos
    que RemisionListView pero devuelve JSON puro.
    """
    import datetime as dt

    # --- AUTH CHECK (devuelve JSON 401, NO redirige a login HTML) ---
    if not request.user.is_authenticated:
        return JsonResponse(
            {'error': 'NO_AUTORIZADO', 'detail': 'Sesión no válida.'},
            status=401
        )

    if request.method != 'GET':
        return JsonResponse({'error': 'Método no permitido'}, status=405)

    # 1. QUERYSET BASE
    queryset = Remision.objects.select_related(
        'empresa', 'origen', 'destino', 'operador'
    ).prefetch_related(
        'detalles__material',
        'evidencias',
        'facturas',
    ).order_by('-pk')

    # 2. PERMISOS POR EMPRESA
    if not request.user.is_superuser:
        perfil = getattr(request.user, 'ternium_profile', None)
        if perfil:
            mis_empresas = perfil.empresas_autorizadas.all()
            queryset = queryset.filter(empresa__in=mis_empresas)
        else:
            queryset = queryset.none()

    # 3. PARÁMETROS
    params = request.GET.copy()
    q_remision = params.get('q_remision', '').strip()
    q_prefijo = params.get('q_prefijo', '').strip()
    q_material = params.get('q_material', '').strip()
    q_status = params.get('q_status', '').strip()
    q_origen = params.get('q_origen', '').strip()
    q_destino = params.get('q_destino', '').strip()
    q_operador = params.get('q_operador', '').strip()
    q_fecha_desde = params.get('q_fecha_desde', '').strip()
    q_fecha_hasta = params.get('q_fecha_hasta', '').strip()
    q_folio_ld = params.get('q_folio_ld', '').strip()
    q_folio_dlv = params.get('q_folio_dlv', '').strip()
    q_destruccion = params.get('q_destruccion', '').strip()

    # 4. FILTRADO
    if q_remision:
        queryset = queryset.filter(remision__icontains=q_remision)
    else:
        filtros_activos = any(
            k.startswith('q_') and v for k, v in params.items()
        )
        if not filtros_activos:
            today = timezone.now().date()
            month_ago = today - dt.timedelta(days=30)
            q_fecha_desde = month_ago.strftime('%Y-%m-%d')
            q_fecha_hasta = today.strftime('%Y-%m-%d')
        if q_fecha_desde:
            queryset = queryset.filter(fecha__gte=q_fecha_desde)
        if q_fecha_hasta:
            queryset = queryset.filter(fecha__lte=q_fecha_hasta)

    if q_prefijo:
        queryset = queryset.filter(empresa__prefijo__icontains=q_prefijo)
    if q_material:
        queryset = queryset.filter(detalles__material_id=q_material)
    if q_origen:
        queryset = queryset.filter(origen_id=q_origen)
    if q_destino:
        queryset = queryset.filter(destino_id=q_destino)
    if q_status:
        queryset = queryset.filter(status=q_status)
    if q_operador:
        queryset = queryset.filter(
            Q(operador__nombre__icontains=q_operador) |
            Q(operador_manual__icontains=q_operador)
        )
    if q_folio_ld:
        queryset = queryset.filter(folio_ld__icontains=q_folio_ld)
    if q_folio_dlv:
        queryset = queryset.filter(folio_dlv__icontains=q_folio_dlv)

    if q_destruccion:
        from .models import ConfiguracionManifiesto
        completo_q = Q(
            fecha_destruccion__isnull=False,
            destruccion_material_1__isnull=False,
            foto_ingreso__isnull=False,
            foto_vertido__isnull=False,
            foto_destruccion__isnull=False
        ) & ~Q(destruccion_material_1='') & ~Q(foto_ingreso='') & ~Q(foto_vertido='') & ~Q(foto_destruccion='')
        if q_destruccion == 'COMPLETO':
            queryset = queryset.filter(completo_q)
        elif q_destruccion == 'PENDIENTE':
            configuraciones = ConfiguracionManifiesto.objects.all()
            q_requiere = Q()
            for conf in configuraciones:
                q_requiere |= Q(origen=conf.origen, detalles__material=conf.material)
            if configuraciones.exists():
                queryset = queryset.filter(q_requiere).exclude(completo_q)
            else:
                queryset = queryset.none()

    if q_material or q_destruccion:
        queryset = queryset.distinct()

    # 5. PAGINACIÓN
    page_number = int(params.get('page', 1))
    page_size = int(params.get('page_size', 25))
    paginator = Paginator(queryset, page_size)
    page_obj = paginator.get_page(page_number)

    # 6. CONTADORES
    total_pendientes = queryset.filter(status='PENDIENTE').count()
    total_cancelados = queryset.filter(status='CANCELADO').count()

    # 7. SERIALIZAR
    remisiones_data = []
    for rem in page_obj.object_list:
        detalles = []
        for det in rem.detalles.all():
            detalles.append({
                'material': det.material.nombre if det.material else '-',
                'peso_ld': float(det.peso_ld or 0),
                'peso_dlv': float(det.peso_dlv or 0),
                'bultos': det.bultos if hasattr(det, 'bultos') and det.bultos else None,
            })

        evidencias_urls = []
        if rem.evidencia_documento and rem.evidencia_documento.name:
            evidencias_urls.append(rem.evidencia_documento.url)
        if rem.boleta_salida_medline and rem.boleta_salida_medline.name:
            evidencias_urls.append(rem.boleta_salida_medline.url)
        if rem.manifiesto and rem.manifiesto.name:
            evidencias_urls.append(rem.manifiesto.url)
        for campo in ['foto_ingreso', 'foto_ingreso_2', 'foto_vertido', 'foto_vertido_2',
                       'foto_destruccion', 'foto_destruccion_2']:
            archivo = getattr(rem, campo, None)
            if archivo and archivo.name:
                evidencias_urls.append(archivo.url)
        for ev in rem.evidencias.all():
            if ev.archivo and ev.archivo.name:
                evidencias_urls.append(ev.archivo.url)

        facturas_data = [{'id': fac.id} for fac in rem.facturas.all()]

        total_ld = float(rem.total_peso_ld or 0)
        total_dlv = float(rem.total_peso_dlv or 0)
        diff = total_dlv - total_ld
        porcentaje_merma = float(rem.porcentaje_merma) if rem.porcentaje_merma else 0

        remisiones_data.append({
            'id': rem.pk,
            'remision': rem.remision,
            'fecha': rem.fecha.strftime('%d/%m/%Y') if rem.fecha else '',
            'fecha_iso': rem.fecha.isoformat() if rem.fecha else '',
            'status': rem.status,
            'status_display': rem.get_status_display(),
            'origen': rem.origen.nombre if rem.origen else '-',
            'destino': rem.destino.nombre if rem.destino else '-',
            'folio_ld': rem.folio_ld or '-',
            'folio_dlv': rem.folio_dlv or '-',
            'total_peso_ld': total_ld,
            'total_peso_dlv': total_dlv,
            'diff': round(diff, 2),
            'porcentaje_merma': round(porcentaje_merma, 1),
            'detalles': detalles,
            'evidencias_urls': evidencias_urls,
            'facturas': facturas_data,
            'permite_manifiesto': rem.permite_manifiesto_destruccion,
            'destruccion_completa': rem.destruccion_fiscal_completa,
        })

    # 8. DATOS PARA FILTROS
    prefijos = list(Empresa.objects.values_list('prefijo', flat=True).distinct().order_by('prefijo'))
    materiales_list = list(Material.objects.values('id', 'nombre').order_by('nombre'))
    origenes_list = list(Lugar.objects.filter(tipo__in=['ORIGEN', 'AMBOS']).values('id', 'nombre').order_by('nombre'))
    destinos_list = list(Lugar.objects.filter(tipo__in=['DESTINO', 'AMBOS']).values('id', 'nombre').order_by('nombre'))
    estatus_choices = [{'value': v, 'display': d} for v, d in Remision.STATUS_CHOICES]

    # 9. RESPUESTA
    return JsonResponse({
        'remisiones': remisiones_data,
        'pagination': {
            'current_page': page_obj.number,
            'total_pages': paginator.num_pages,
            'total_count': paginator.count,
            'has_previous': page_obj.has_previous(),
            'has_next': page_obj.has_next(),
            'previous_page': page_obj.previous_page_number() if page_obj.has_previous() else None,
            'next_page': page_obj.next_page_number() if page_obj.has_next() else None,
        },
        'counters': {
            'total_pendientes': total_pendientes,
            'total_cancelados': total_cancelados,
        },
        'filters': {
            'prefijos': prefijos,
            'materiales': materiales_list,
            'origenes': origenes_list,
            'destinos': destinos_list,
            'estatus_choices': estatus_choices,
        },
        'permissions': {
            'can_add': request.user.has_perm('ternium.add_remision'),
            'can_change': request.user.has_perm('ternium.change_remision'),
            'can_view_plastico': request.user.has_perm('ternium.view_plastico'),
            'can_view_tarimas': request.user.has_perm('ternium.view_controltarima'),
        },
    })
    
@require_POST
def exportar_zip_medline(request):
    """
    Exporta en un archivo ZIP TODAS las Boletas de Salida de remisiones MEDLINE.
    Los archivos se nombran: {remision}_{folio_ld}_{folio_dlv}.ext
    """
    # Filtro base: Origen MEDLINE con archivo registrado
    queryset = Remision.objects.filter(
        origen__nombre__icontains='MEDLINE'
    ).exclude(boleta_salida_medline__exact='').exclude(boleta_salida_medline__isnull=True)

    if not queryset.exists():
        messages.warning(request, "No se encontraron boletas de salida MEDLINE registradas.")
        return redirect('remision_lista')

    buffer = BytesIO()
    archivos_agregados = 0

    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for remision in queryset:
            archivo = remision.boleta_salida_medline

            if archivo and hasattr(archivo, 'name') and archivo.name:
                try:
                    ext = archivo.name.split('.')[-1]
                    remision_part = (remision.remision or 'SIN_REMISION').replace('/', '-')
                    folio_ld_part = (remision.folio_ld or 'SIN_LD').replace('/', '-')
                    folio_dlv_part = (remision.folio_dlv or 'SIN_DLV').replace('/', '-')
                    nuevo_nombre = f"{remision_part}_{folio_ld_part}_{folio_dlv_part}.{ext}"

                    file_data = None
                    try:
                        with archivo.open('rb') as f:
                            file_data = f.read()
                    except Exception:
                        try:
                            req = urllib.request.Request(
                                archivo.url,
                                headers={'User-Agent': 'Mozilla/5.0'}
                            )
                            with urllib.request.urlopen(req) as response:
                                file_data = response.read()
                        except Exception as e_url:
                            print(f"Error leyendo archivo remision {remision.id}: {e_url}")

                    if file_data:
                        zip_file.writestr(nuevo_nombre, file_data)
                        archivos_agregados += 1

                except Exception as e:
                    print(f"Error procesando remisión {remision.id}: {e}")
                    continue

    if archivos_agregados == 0:
        messages.error(request, "No fue posible leer los archivos físicos en S3.")
        return redirect('remision_lista')

    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="Boletas_Salida_MEDLINE.zip"'
    return response


# ============================================================
# API VIEW: Dashboard Trane para Next.js
# ============================================================
# AGREGAR a views.py y registrar en urls.py:
#   path('api/dashboard-trane/', views.api_dashboard_trane, name='api_dashboard_trane'),
# ============================================================
# También agregar a api.ts el método getDashboardTrane
# ============================================================

import json
import datetime as dt
from django.http import JsonResponse
from django.db.models import Sum, Count, FloatField
from django.db.models.functions import Coalesce, TruncMonth
from django.utils import timezone

def api_dashboard_trane(request):
    """
    API completa para Dashboard Trane en React/Next.js.
    Devuelve KPIs, gráficas, tablas (fase1, fase2, bitácora) y filtros.
    """

    # --- 0. FILTROS ---
    filtro_material_id = request.GET.get('material')
    fecha_inicio = request.GET.get('fecha_inicio')
    fecha_fin = request.GET.get('fecha_fin')

    # MODIFICACIÓN: Filtro automático para "lo que va de este año"
    if not fecha_inicio and not fecha_fin and not filtro_material_id:
        hoy = timezone.now().date()
        fecha_fin = hoy.isoformat()
        # Se establece el 1 de enero del año actual
        fecha_inicio = hoy.replace(month=1, day=1).isoformat()

    # Materiales disponibles para el slicer
    materiales_disponibles = list(
        Material.objects.filter(
            detalleremision__remision__origen__nombre__iexact='TRANE'
        ).distinct().order_by('nombre').values('id', 'nombre')
    )

    # --- 1. CONSULTA PRINCIPAL ---
    queryset = Remision.objects.filter(
        origen__nombre__iexact='TRANE',
        status__in=['TERMINADO', 'AUDITADO']
    )
    if filtro_material_id:
        queryset = queryset.filter(detalles__material_id=filtro_material_id)
    if fecha_inicio:
        queryset = queryset.filter(fecha__gte=fecha_inicio)
    if fecha_fin:
        queryset = queryset.filter(fecha__lte=fecha_fin)

    # --- 2. KPIs ---
    agg = queryset.aggregate(
        total_kg=Coalesce(Sum('detalles__peso_ld'), 0.0, output_field=FloatField())
    )
    total_tons = round(agg['total_kg'] / 1000, 2)
    total_viajes = queryset.count()

    # --- 3. GRÁFICAS ---
    # Mensual
    mensual_qs = queryset.annotate(mes=TruncMonth('fecha')).values('mes').annotate(
        total_kg=Sum('detalles__peso_ld')
    ).order_by('mes')
    chart_mensual = {
        'labels': [m['mes'].strftime("%b %Y") for m in mensual_qs if m['mes']],
        'data': [round((m['total_kg'] or 0) / 1000, 2) for m in mensual_qs if m['mes']],
    }

    # Destinos
    destinos_qs = queryset.values('destino__nombre').annotate(
        total_kg=Sum('detalles__peso_ld')
    ).order_by('-total_kg')
    chart_destinos = {
        'labels': [d['destino__nombre'] or 'N/A' for d in destinos_qs],
        'data': [round((d['total_kg'] or 0) / 1000, 2) for d in destinos_qs],
    }

    # Materiales
    materiales_qs = queryset.values('detalles__material__nombre').annotate(
        total_kg=Sum('detalles__peso_ld')
    ).order_by('-total_kg')
    chart_materiales = {
        'labels': [m['detalles__material__nombre'] or 'N/A' for m in materiales_qs],
        'data': [round((m['total_kg'] or 0) / 1000, 2) for m in materiales_qs],
    }

    # --- 4. MOVIMIENTOS RAW (para breakdown y bitácora) ---
    raw_movimientos = queryset.select_related(
        'origen', 'destino', 'linea_transporte', 'operador', 'unidad'
    ).prefetch_related('detalles__material', 'evidencias').order_by('-fecha')[:500]

    breakdown_map = {}
    bitacora = []

    for mov in raw_movimientos:
        detalles_cached = mov.detalles.all()
        detalle = detalles_cached[0] if detalles_cached else None
        mat_nom = detalle.material.nombre if (detalle and detalle.material) else "Sin Clasificar"
        dest_nom = mov.destino.nombre if mov.destino else "Sin Destino"
        key = f"{mat_nom}_{dest_nom}"

        peso_ton = round((mov.total_peso_ld or 0) / 1000, 3)
        transporte = mov.linea_transporte.nombre if mov.linea_transporte else 'S/T'
        operador = mov.operador.nombre if mov.operador else (mov.operador_manual or 'S/O')
        folio = mov.folio_ld.strip() if mov.folio_ld and mov.folio_ld.strip() else mov.remision
        trazabilidad = mov.trazabilidad_notas or ""

        # Archivos
        archivos = []
        if mov.evidencia_documento:
            archivos.append({'url': mov.evidencia_documento.url, 'type': 'evidencia'})
        if mov.manifiesto:
            archivos.append({'url': mov.manifiesto.url, 'type': 'manifiesto'})
        for ex in mov.evidencias.all():
            if ex.archivo:
                archivos.append({'url': ex.archivo.url, 'type': 'extra'})

        if key not in breakdown_map:
            breakdown_map[key] = []
        breakdown_map[key].append({
            'remision': folio, 'fecha': mov.fecha.strftime("%d/%m/%Y"),
            'peso': peso_ton, 'archivos': archivos,
            'origen': mov.origen.nombre if mov.origen else 'N/A',
            'destino': dest_nom, 'operador': operador,
            'trazabilidad': trazabilidad,
        })

        bitacora.append({
            'remision': folio, 'fecha': mov.fecha.strftime("%Y-%m-%d"),
            'status': mov.get_status_display(),
            'origen': mov.origen.nombre if mov.origen else 'N/A',
            'destino': dest_nom, 'transporte': transporte,
            'operador': operador, 'material': mat_nom,
            'peso': peso_ton, 'trazabilidad': trazabilidad,
            'archivos': archivos,
        })

    # --- 5. TABLA FASE 1 ---
    trazabilidad_qs = queryset.values(
        'detalles__material__nombre', 'destino__nombre', 'destino__es_patio'
    ).annotate(
        kilos=Sum('detalles__peso_ld'), viajes=Count('id')
    ).order_by('detalles__material__nombre')

    fase1 = []
    for item in trazabilidad_qs:
        mat_nom = item['detalles__material__nombre'] or "Sin Clasificar"
        dest_nom = item['destino__nombre'] or "Sin Destino"
        tipo = "PATIO" if item['destino__es_patio'] else "RECICLADOR"
        key = f"{mat_nom}_{dest_nom}"
        fase1.append({
            'material': mat_nom, 'destino': dest_nom,
            'tipo_destino': tipo,
            'toneladas': round((item['kilos'] or 0) / 1000, 3),
            'viajes': item['viajes'],
            'breakdown': breakdown_map.get(key, []),
        })

    # --- 6. TABLA FASE 2 ---
    envios_a_patio = queryset.filter(destino__es_patio=True)
    patios_ids = envios_a_patio.values_list('destino_id', flat=True).distinct()
    mat_ids = [filtro_material_id] if filtro_material_id else list(
        envios_a_patio.values_list('detalles__material_id', flat=True).distinct()
    )

    fase2 = []
    if patios_ids:
        salidas_patio = Remision.objects.filter(
            origen_id__in=patios_ids, detalles__material_id__in=mat_ids,
            status__in=['TERMINADO', 'AUDITADO']
        ).exclude(destino__es_patio=True).select_related(
            'origen', 'destino', 'operador'
        ).prefetch_related('detalles__material', 'evidencias')

        if fecha_inicio:
            salidas_patio = salidas_patio.filter(fecha__gte=fecha_inicio)
        if fecha_fin:
            salidas_patio = salidas_patio.filter(fecha__lte=fecha_fin)

        bd_f2 = {}
        for mov in salidas_patio:
            det = mov.detalles.first()
            mn = det.material.nombre if det and det.material else "S/C"
            on = mov.origen.nombre if mov.origen else "S/O"
            dn = mov.destino.nombre if mov.destino else "S/D"
            k2 = f"{on}_{mn}_{dn}"
            f2_folio = mov.folio_ld.strip() if mov.folio_ld and mov.folio_ld.strip() else mov.remision
            if k2 not in bd_f2:
                bd_f2[k2] = []
            arch2 = []
            if mov.evidencia_documento:
                arch2.append({'url': mov.evidencia_documento.url, 'type': 'evidencia'})
            if mov.manifiesto:
                arch2.append({'url': mov.manifiesto.url, 'type': 'manifiesto'})
            for ex in mov.evidencias.all():
                if ex.archivo:
                    arch2.append({'url': ex.archivo.url, 'type': 'extra'})
            bd_f2[k2].append({
                'remision': f2_folio, 'fecha': mov.fecha.strftime("%d/%m/%Y"),
                'peso': round((mov.total_peso_ld or 0) / 1000, 3), 'archivos': arch2,
                'origen': on, 'destino': dn,
                'operador': mov.operador.nombre if mov.operador else 'S/O',
                'trazabilidad': mov.trazabilidad_notas or "",
            })

        agrupado = salidas_patio.values(
            'origen__nombre', 'detalles__material__nombre', 'destino__nombre'
        ).annotate(kilos=Sum('detalles__peso_ld'), viajes=Count('id')).order_by('detalles__material__nombre')

        for s in agrupado:
            k2 = f"{s['origen__nombre']}_{s['detalles__material__nombre']}_{s['destino__nombre']}"
            fase2.append({
                'patio_origen': s['origen__nombre'],
                'material': s['detalles__material__nombre'],
                'cliente_final': s['destino__nombre'],
                'toneladas': round((s['kilos'] or 0) / 1000, 3),
                'breakdown': bd_f2.get(k2, []),
            })

    # --- 7. SALIDAS (ControlManifiestoTrane) ---
    salidas_qs = ControlManifiestoTrane.objects.select_related(
        'linea_transporte', 'operador', 'destino', 'material', 'remision_vinculada'
    )
    if fecha_inicio:
        salidas_qs = salidas_qs.filter(fecha_captura__gte=fecha_inicio)
    if fecha_fin:
        salidas_qs = salidas_qs.filter(fecha_captura__lte=fecha_fin)
    salidas_qs = salidas_qs.order_by('-fecha_captura', '-id')[:100]

    salidas_list = []
    for r in salidas_qs:
        docs = []
        if r.manifiesto:
            docs.append({'url': r.manifiesto.url, 'type': 'manifiesto'})
        if r.documento_trane:
            docs.append({'url': r.documento_trane.url, 'type': 'doc_trane'})
        salidas_list.append({
            'folio': r.folio or '-',
            'remision_vinculada': r.remision_vinculada.remision if r.remision_vinculada else None,
            'fecha': r.fecha_captura.strftime("%Y-%m-%d") if r.fecha_captura else '-',
            'transporte': r.linea_transporte.nombre if r.linea_transporte else '-',
            'operador': r.operador.nombre if r.operador else (r.operador_manual or '-'),
            'destino': r.destino.nombre if r.destino else '-',
            'material': r.material.nombre if r.material else '-',
            'peso_kg': round(float(r.cantidad_kg or 0), 4),
            'documentos': docs,
        })

    # --- RESPONSE ---
    return JsonResponse({
        'kpis': {'totalTons': total_tons, 'totalViajes': total_viajes},
        'charts': {
            'mensual': chart_mensual,
            'destinos': chart_destinos,
            'materiales': chart_materiales,
        },
        'fase1': fase1,
        'fase2': fase2,
        'bitacora': bitacora,
        'salidas': salidas_list,
        'filtros': {
            'materiales': [{'id': m['id'], 'nombre': m['nombre']} for m in materiales_disponibles],
            'materialActual': int(filtro_material_id) if filtro_material_id else None,
            'fechaInicio': fecha_inicio,
            'fechaFin': fecha_fin,
        },
    })